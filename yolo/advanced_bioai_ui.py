
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BioAI Research Platform v2.0 Advanced - 高级版本
包含批量选择、提示词模板和自动工作流功能的科技风Web界面
"""

import os
import json
import base64
import random
import datetime
import re
import unicodedata
from datetime import datetime as dt
from typing import Dict, Any, List
from langchain_milvus import Milvus
from docx import Document as DocxDocument
from typing import List, Any
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownTextSplitter

embeddings = OpenAIEmbeddings(model="/data/bge-m3/",api_key="sk-hzw-20250804",base_url="http://220.185.228.102:30199/v1")
vectorstore_path = "/root/net-disk/test1/长鑫1103/vec.db"

# 延迟初始化向量数据库以避免启动时的锁定问题
vectorstore = None

def init_vectorstore():
    """延迟初始化向量数据库"""
    global vectorstore
    if vectorstore is None:
        try:
            vectorstore = Milvus(
                embedding_function=embeddings,
                connection_args={
                    "uri": vectorstore_path,
                },
                index_params={"index_type": "FLAT", "metric_type": "COSINE"},
                auto_id=True,
                drop_old=False,
            )
            print("✅ 向量数据库初始化成功")
        except Exception as e:
            print(f"❌ 知识库初始化失败: {e}")
            vectorstore = None
    return vectorstore

# ===== 支持中文的安全文件名处理函数 =====
def safe_filename_with_chinese(filename: str) -> str:
    """
    安全处理文件名，保留中文字符同时移除危险字符

    Args:
        filename: 原始文件名

    Returns:
        安全的文件名
    """
    if not filename:
        return "file"

    # 移除路径分隔符
    filename = os.path.basename(filename)

    # 保留中文、英文、数字、下划线、连字符、点
    # 允许的字符：中文字符、英文字母、数字、下划线、连字符、点
    safe_name = re.sub(r'[^\w\-\.\u4e00-\u9fff]', '_', filename, flags=re.UNICODE)

    # 移除连续的下划线和点
    safe_name = re.sub(r'_+', '_', safe_name)
    safe_name = re.sub(r'\.+', '.', safe_name)

    # 处理扩展名前的下划线（如 "文件名_.pdf" -> "文件名.pdf"）
    safe_name = re.sub(r'_+\.', '.', safe_name)

    # 移除开头和结尾的下划线和点
    safe_name = safe_name.strip('_.')

    # 限制长度（保留扩展名）
    if len(safe_name) > 200:
        name_part, ext = os.path.splitext(safe_name)
        safe_name = name_part[:180] + ext

    return safe_name if safe_name else "file"

# ===== 在文件顶部或合适位置新增 =====
def _get_collection_name_from_vectorstore(vs) -> str:
    # 兼容不同版本的 langchain_milvus
    return (
        getattr(vs, "collection_name", None)
        or getattr(vs, "_collection_name", None)
        or "LangChainCollection"   # 默认名（不同版本可能不同，按需替换）
    )

def file_id_exists(vectorstore, file_id: str) -> bool:
    """
    在 Milvus / Milvus-Lite 中查询是否已有给定 file_id 的数据。
    兼容两种常见写法：
    1) 动态字段：file_id == "xxx"
    2) JSON 元数据：metadata["file_id"] == "xxx"
    任一查询命中即认为存在。
    """
    from pymilvus import connections, Collection
    import contextlib

    # 连接到与 vectorstore 相同的 uri（Milvus-Lite 时即本地 vec.db）
    uri = None
    vs = init_vectorstore()
    if vs and hasattr(vs, "connection_args"):
        uri = vs.connection_args.get("uri")
    # 兜底
    if not uri:
        uri = "vec.db"

    with contextlib.suppress(Exception):
        connections.disconnect("default")
    connections.connect(alias="default", uri=uri)

    col_name = _get_collection_name_from_vectorstore(vectorstore)

    try:
        col = Collection(col_name)
    except Exception:
        # 集合都不存在，当然也不可能有该 file_id
        return False

    # Lite 里也需要 load 一下
    with contextlib.suppress(Exception):
        col.load()

    # 先尝试动态字段（enable_dynamic_field 时，metadata 会被摊平为独立列）
    expr_candidates = [
        f'file_id == "{file_id}"',
        f'metadata["file_id"] == "{file_id}"',  # 如果是 JSON 列
    ]

    for expr in expr_candidates:
        try:
            # 只取 1 条判断是否存在；output_fields 取现成字段，避免主键名不一致
            res = col.query(expr=expr, output_fields=["file_id"], limit=1)
            if len(res) > 0:
                return True
        except Exception:
            continue

    return False

def write_wider_window(split_docs: List[Any], original_documents: Any, offset: int = 200):
    original_text = original_documents.page_content
    count = 0
    for doc in split_docs:
        count += 1
        doc_text = doc.page_content
        start_index = original_text.index(doc_text)
        end_index = start_index + len(doc_text) - 1
        wider_text = original_text[
            max(0, start_index - offset) : min(len(original_text), end_index + offset)
        ]
        doc.metadata["start_index"] = start_index
        doc.metadata["end_index"] = end_index
        doc.metadata["wider_text"] = wider_text

def load_magic_pdf_results(filename: str) -> dict:
    """加载Magic PDF生成的完整结果文件"""
    try:
        # 构建输出目录路径
        base_name = os.path.splitext(filename)[0]
        output_dir = f"/root/net-disk/test1/长鑫1103/uploads/{base_name}"

        results = {}

        # 尝试加载Markdown文件
        md_file = os.path.join(output_dir, f"{base_name}.md")
        if os.path.exists(md_file):
            with open(md_file, 'r', encoding='utf-8') as f:
                results['markdown'] = f.read()
                print(f"✅ Loaded Magic PDF Markdown: {len(results['markdown'])} characters")

        # 尝试加载content_list.json
        content_list_file = os.path.join(output_dir, f"{base_name}_content_list.json")
        if os.path.exists(content_list_file):
            with open(content_list_file, 'r', encoding='utf-8') as f:
                results['content_list'] = json.load(f)
                print(f"✅ Loaded Magic PDF content list")

        # 尝试加载middle.json
        middle_file = os.path.join(output_dir, f"{base_name}_middle.json")
        if os.path.exists(middle_file):
            with open(middle_file, 'r', encoding='utf-8') as f:
                results['middle_json'] = json.load(f)
                print(f"✅ Loaded Magic PDF middle JSON")

        return results

    except Exception as e:
        print(f"⚠️ Failed to load Magic PDF results for {filename}: {e}")
        return {}

def process_content_to_vector(content: str, vectorstore, original_filename: str, file_id: str):
    """将内容字符串处理并添加到向量存储，并返回向量ID列表"""
    from langchain_core.documents import Document

    original_doc = Document(page_content=content)

    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100, separators=[""])

    doc_splits = splitter.split_documents([original_doc])
    write_wider_window(doc_splits, original_doc, offset=300)

    for split_doc in doc_splits:
        split_doc.metadata["source"] = original_filename
        split_doc.metadata["file_id"] = file_id
        split_doc.page_content = f"{split_doc.page_content}"

    if doc_splits:
        vs = init_vectorstore()
        if vs:
            return vs.add_documents(doc_splits)
    return []

def analyze_content_characteristics(content: str, metadata: list) -> dict:
    """分析内容特征，用于优化AI提示词"""
    characteristics = {
        'content_type': 'general',
        'has_molecular_data': False,
        'has_tables': False,
        'has_patents': False,
        'has_bioactivity': False,
        'has_smiles': False,
        'document_types': set(),
        'key_terms': set(),
        'complexity_level': 'medium'
    }

    content_lower = content.lower()

    # 检测文档类型
    for doc in metadata:
        filename = doc.get('filename', '').lower()
        if 'patent' in filename or 'wo' in filename or 'us' in filename:
            characteristics['has_patents'] = True
            characteristics['document_types'].add('patent')
        elif 'paper' in filename or 'article' in filename:
            characteristics['document_types'].add('research_paper')
        elif 'compound' in filename or 'molecule' in filename:
            characteristics['document_types'].add('molecular_data')

    # 检测内容特征
    molecular_keywords = ['smiles', 'molecular weight', 'mw', 'clogp', 'tpsa', 'hbd', 'hba', 'ic50', 'ec50', 'ki']
    table_keywords = ['table', '表格', 'compound', 'example', '化合物']
    bioactivity_keywords = ['ic50', 'ec50', 'ki', 'bioactivity', '生物活性', 'inhibition', 'binding']

    for keyword in molecular_keywords:
        if keyword in content_lower:
            characteristics['has_molecular_data'] = True
            characteristics['key_terms'].add(keyword)

    for keyword in table_keywords:
        if keyword in content_lower:
            characteristics['has_tables'] = True

    for keyword in bioactivity_keywords:
        if keyword in content_lower:
            characteristics['has_bioactivity'] = True

    # 检测SMILES结构
    import re
    smiles_pattern = r'[A-Za-z0-9@+\-\[\]()=#$:/\\.]+'
    if re.search(r'smiles.*?[:=]\s*[A-Za-z0-9@+\-\[\]()=#$:/\\.]+', content_lower):
        characteristics['has_smiles'] = True

    # 确定内容类型
    if characteristics['has_molecular_data'] and characteristics['has_bioactivity']:
        characteristics['content_type'] = 'pharmaceutical_research'
    elif characteristics['has_patents']:
        characteristics['content_type'] = 'patent_analysis'
    elif characteristics['has_molecular_data']:
        characteristics['content_type'] = 'molecular_data'
    elif characteristics['has_tables']:
        characteristics['content_type'] = 'structured_data'

    # 评估复杂度
    if len(characteristics['key_terms']) > 5 and characteristics['has_molecular_data']:
        characteristics['complexity_level'] = 'high'
    elif len(characteristics['key_terms']) < 2:
        characteristics['complexity_level'] = 'low'

    return characteristics

def build_intelligent_prompt(query: str, combined_content: str, document_metadata: list,
                           template_info: dict = None, retrieval_stats: dict = None) -> str:
    """基于知识库内容和查询特征智能构建AI提示词"""

    # 分析内容特征
    content_chars = analyze_content_characteristics(combined_content, document_metadata)

    # 分析查询意图
    query_lower = query.lower()
    query_intent = 'general'

    if any(word in query_lower for word in ['提取', 'extract', '找出', '列出']):
        query_intent = 'extraction'
    elif any(word in query_lower for word in ['分析', 'analyze', '比较', 'compare']):
        query_intent = 'analysis'
    elif any(word in query_lower for word in ['总结', 'summarize', '概述', 'overview']):
        query_intent = 'summarization'
    elif any(word in query_lower for word in ['计算', 'calculate', '预测', 'predict']):
        query_intent = 'computation'

    # 构建系统提示词
    system_prompt = build_system_prompt(content_chars, query_intent)

    # 构建用户提示词
    user_prompt = build_user_prompt(query, combined_content, content_chars, query_intent,
                                  document_metadata, retrieval_stats)

    return user_prompt

def smart_content_preview(content: str, max_chars: int = 8000) -> str:
    """智能内容预览，保留关键信息"""
    if len(content) <= max_chars:
        return content

    # 按段落分割
    paragraphs = content.split('\n\n')

    # 优先保留的内容类型
    priority_keywords = [
        'abstract', '摘要', 'summary', 'conclusion', '结论',
        'IC50', 'EC50', 'Ki', 'activity', '活性',
        'compound', '化合物', 'molecule', '分子',
        'table', '表格', 'figure', '图',
        'example', '实例', 'claim', '权利要求'
    ]

    # 分类段落
    high_priority = []
    medium_priority = []
    low_priority = []

    for para in paragraphs:
        para_lower = para.lower()
        if any(keyword in para_lower for keyword in priority_keywords):
            if len(para) < 1000:  # 避免单个段落过长
                high_priority.append(para)
            else:
                # 截断长段落但保留开头
                high_priority.append(para[:800] + "...[内容截断]")
        elif len(para) > 50:  # 过滤掉太短的段落
            if len(para) < 500:
                medium_priority.append(para)
            else:
                low_priority.append(para[:300] + "...[内容截断]")

    # 组合内容
    result = []
    current_length = 0

    # 添加高优先级内容
    for para in high_priority:
        if current_length + len(para) < max_chars * 0.7:  # 为其他内容预留空间
            result.append(para)
            current_length += len(para)

    # 添加中等优先级内容
    for para in medium_priority:
        if current_length + len(para) < max_chars * 0.9:
            result.append(para)
            current_length += len(para)
        else:
            break

    # 如果还有空间，添加低优先级内容
    for para in low_priority:
        if current_length + len(para) < max_chars:
            result.append(para)
            current_length += len(para)
        else:
            break

    final_content = '\n\n'.join(result)

    # 添加截断说明
    if len(content) > len(final_content):
        truncation_info = f"\n\n[📄 内容智能截断: 原文档 {len(content):,} 字符，已优化为 {len(final_content):,} 字符，保留关键信息]"
        final_content += truncation_info

    return final_content

def smart_json_content_extraction(content: str, filename: str = "") -> str:
    """专门针对JSON文件的智能内容提取"""
    try:
        import json

        # 检查是否是JSON文件
        if not (filename.lower().endswith('.json') or content.strip().startswith('{')):
            return content

        print(f"🔍 检测到JSON文件，进行智能结构化提取...")

        # 尝试解析JSON
        try:
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')

            json_data = json.loads(content)

            # 构建结构化摘要
            summary_parts = []
            summary_parts.append("=== JSON文件结构化分析 ===")

            # 分析JSON结构
            def analyze_json_structure(data, prefix="", max_depth=3, current_depth=0):
                if current_depth >= max_depth:
                    return ["... (结构层次过深，已截断)"]

                analysis = []
                if isinstance(data, dict):
                    analysis.append(f"{prefix}📋 对象包含 {len(data)} 个字段:")
                    for key, value in list(data.items())[:10]:  # 限制显示前10个字段
                        if isinstance(value, (dict, list)):
                            analysis.append(f"{prefix}  • {key}: {type(value).__name__} ({len(value)} 项)")
                            if current_depth < 2:  # 只深入2层
                                analysis.extend(analyze_json_structure(value, prefix + "    ", max_depth, current_depth + 1))
                        else:
                            value_str = str(value)[:100] + "..." if len(str(value)) > 100 else str(value)
                            analysis.append(f"{prefix}  • {key}: {value_str}")

                    if len(data) > 10:
                        analysis.append(f"{prefix}  ... 还有 {len(data) - 10} 个字段")

                elif isinstance(data, list):
                    analysis.append(f"{prefix}📝 数组包含 {len(data)} 个元素")
                    if data and len(data) > 0:
                        first_item = data[0]
                        analysis.append(f"{prefix}  • 元素类型: {type(first_item).__name__}")
                        if isinstance(first_item, dict) and current_depth < 2:
                            analysis.append(f"{prefix}  • 示例元素结构:")
                            analysis.extend(analyze_json_structure(first_item, prefix + "    ", max_depth, current_depth + 1))

                return analysis

            # 添加结构分析
            structure_analysis = analyze_json_structure(json_data)
            summary_parts.extend(structure_analysis)

            # 提取关键数据
            summary_parts.append("\n=== 关键数据提取 ===")

            def extract_key_data(data, path=""):
                key_data = []
                if isinstance(data, dict):
                    for key, value in data.items():
                        current_path = f"{path}.{key}" if path else key

                        # 查找可能的重要字段
                        important_keywords = [
                            'name', 'title', 'description', 'content', 'text', 'summary',
                            'id', 'type', 'category', 'status', 'result', 'data',
                            'compound', 'molecule', 'chemical', 'formula', 'smiles',
                            'activity', 'property', 'value', 'concentration', 'dose',
                            'clinical', 'trial', 'study', 'experiment', 'test'
                        ]

                        if any(keyword in key.lower() for keyword in important_keywords):
                            if isinstance(value, (str, int, float, bool)):
                                value_str = str(value)[:200] + "..." if len(str(value)) > 200 else str(value)
                                key_data.append(f"🔑 {current_path}: {value_str}")
                            elif isinstance(value, list) and len(value) > 0:
                                key_data.append(f"🔑 {current_path}: 数组 ({len(value)} 项)")
                                if isinstance(value[0], (str, int, float)):
                                    sample = value[:3]
                                    key_data.append(f"   示例: {sample}")

                        # 递归处理嵌套结构（限制深度）
                        if isinstance(value, dict) and len(path.split('.')) < 3:
                            key_data.extend(extract_key_data(value, current_path))
                        elif isinstance(value, list) and len(value) > 0 and isinstance(value[0], dict) and len(path.split('.')) < 2:
                            key_data.extend(extract_key_data(value[0], f"{current_path}[0]"))

                return key_data

            key_data = extract_key_data(json_data)
            if key_data:
                summary_parts.extend(key_data[:20])  # 限制显示前20个关键数据
            else:
                summary_parts.append("未找到明显的关键数据字段")

            # 添加统计信息
            summary_parts.append(f"\n=== 文件统计 ===")
            summary_parts.append(f"📊 原始文件大小: {len(content):,} 字符")
            summary_parts.append(f"📊 JSON对象类型: {type(json_data).__name__}")

            if isinstance(json_data, dict):
                summary_parts.append(f"📊 顶级字段数量: {len(json_data)}")
            elif isinstance(json_data, list):
                summary_parts.append(f"📊 数组元素数量: {len(json_data)}")

            # 如果内容仍然很长，添加原始内容的关键部分
            if len(content) > 50000:  # 对于超大JSON文件
                summary_parts.append(f"\n=== 原始内容片段 ===")
                summary_parts.append(content[:2000] + "\n... (内容过长，已截断) ...")
                summary_parts.append(content[-1000:])
            else:
                summary_parts.append(f"\n=== 完整JSON内容 ===")
                summary_parts.append(content)

            result = '\n'.join(summary_parts)
            print(f"✅ JSON智能提取完成: 原始 {len(content):,} 字符 -> 结构化 {len(result):,} 字符")
            return result

        except json.JSONDecodeError as e:
            print(f"⚠️ JSON解析失败，使用原始内容: {e}")
            return content

    except Exception as e:
        print(f"❌ JSON智能提取失败: {e}")
        return content

def smart_vector_retrieval(vectorstore, query: str, doc_id: str, max_chunks: int = 10, is_json: bool = False) -> str:
    """智能向量检索，获取最相关的内容片段"""
    try:
        # 构建文档过滤表达式
        expr = f'file_id == "{doc_id}"'

        # 根据文件类型调整检索参数
        if is_json:
            # JSON文件需要更多样化的检索
            search_kwargs = {
                "k": max_chunks + 4,  # JSON文件获取更多片段
                "expr": expr,
                "lambda_mult": 0.5,  # 更注重多样性
                "fetch_k": max_chunks * 3  # 增加候选集
            }
            print(f"🔍 JSON文件向量检索模式: 获取 {max_chunks + 4} 个多样化片段")
        else:
            # 普通文件的检索参数
            search_kwargs = {
                "k": max_chunks,
                "expr": expr,
                "lambda_mult": 0.7,  # 平衡相关性和多样性
                "fetch_k": max_chunks * 2  # 标准候选集
            }

        # 创建检索器
        vs = init_vectorstore()
        if not vs:
            return "向量数据库未初始化，无法进行检索"

        retriever = vs.as_retriever(
            search_type="mmr",  # 使用MMR确保多样性
            search_kwargs=search_kwargs
        )

        # 执行检索
        docs = retriever.invoke(query)

        if docs:
            # 合并检索到的内容
            retrieved_content = []
            total_length = 0
            max_total_length = 15000 if is_json else 12000  # JSON文件允许更多内容

            for i, doc in enumerate(docs):
                if doc.page_content.strip():
                    content_piece = doc.page_content.strip()

                    # 控制总长度
                    if total_length + len(content_piece) > max_total_length:
                        # 截断最后一个片段
                        remaining_length = max_total_length - total_length
                        if remaining_length > 100:  # 至少保留100字符
                            content_piece = content_piece[:remaining_length] + "..."
                            retrieved_content.append(f"[片段 {i+1}] {content_piece}")
                        break

                    retrieved_content.append(f"[片段 {i+1}] {content_piece}")
                    total_length += len(content_piece)

            result = '\n\n'.join(retrieved_content)

            if is_json:
                # 为JSON文件添加结构化说明
                result = f"=== JSON文件向量检索结果 ===\n检索到 {len(retrieved_content)} 个相关数据片段，总长度 {len(result)} 字符\n\n{result}"

            print(f"✅ 向量检索成功: 获取 {len(docs)} 个相关片段，总长度 {len(result)} 字符")
            return result
        else:
            print("⚠️ 向量检索未找到相关内容")
            return ""

    except Exception as e:
        print(f"❌ 向量检索失败: {e}")
        return ""

def _get_segment_file_as_document(doc_id):
    """将segment file转换为文档格式以供Information Extraction使用"""
    try:
        from bson import ObjectId

        # 转换为ObjectId
        if isinstance(doc_id, str):
            try:
                object_id = ObjectId(doc_id)
            except:
                print(f"❌ 无效的segment file ID格式: {doc_id}")
                return None
        else:
            object_id = doc_id

        # 查找segment file
        segment_doc = db_manager.db.segment_files.find_one({"_id": object_id})
        if not segment_doc:
            return None

        # 转换为文档格式
        content = segment_doc.get('content', '')

        # 如果是BSON Binary，转换为字符串
        if hasattr(content, 'binary'):
            content = content.binary.decode('utf-8', errors='ignore')
        elif isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')

        # 构建虚拟文档对象
        virtual_doc = {
            'id': str(segment_doc['_id']),
            'filename': segment_doc.get('virtual_filename', 'Unknown'),
            'file_type': segment_doc.get('file_type', 'Unknown'),
            'text_content': content,  # 使用text_content字段以保持兼容性
            'content': content,       # 同时保留content字段
            'is_segment_file': True,
            'parent_document_id': segment_doc.get('document_id'),
            'upload_time': segment_doc.get('created_time', '').isoformat() if segment_doc.get('created_time') else None
        }

        print(f"✅ 成功转换segment file: {virtual_doc['filename']}")
        return virtual_doc

    except Exception as e:
        print(f"❌ 转换segment file失败: {e}")
        return None

def build_system_prompt(content_chars: dict, query_intent: str) -> str:
    """构建系统提示词"""
    base_prompt = "你是一个专业的生物医学文档分析助手，"

    if content_chars['content_type'] == 'pharmaceutical_research':
        base_prompt += "专精于药物研发、分子设计和生物活性分析。"
    elif content_chars['content_type'] == 'patent_analysis':
        base_prompt += "专精于专利文档分析和技术信息提取。"
    elif content_chars['content_type'] == 'molecular_data':
        base_prompt += "专精于分子结构分析和理化性质解读。"
    else:
        base_prompt += "擅长从科学文献中提取关键信息。"

    if query_intent == 'extraction':
        base_prompt += "你的任务是精确提取和整理文档中的具体数据。"
    elif query_intent == 'analysis':
        base_prompt += "你的任务是深入分析数据间的关系和趋势。"
    elif query_intent == 'summarization':
        base_prompt += "你的任务是提供清晰简洁的内容总结。"
    elif query_intent == 'computation':
        base_prompt += "你的任务是进行数据计算和预测分析。"

    return base_prompt

def build_user_prompt(query: str, content: str, content_chars: dict, query_intent: str,
                     metadata: list, stats: dict) -> str:
    """构建用户提示词"""

    # 基础结构
    prompt = f"""基于以下知识库内容回答用户查询。

**用户查询**: {query}

**知识库统计**:
- 文档数量: {stats.get('document_count', 0)}
- 内容片段: {stats.get('total_chunks', 0)}
- 总字符数: {stats.get('content_length', 0):,}
- 文档类型: {', '.join(content_chars.get('document_types', ['通用']))}

**知识库内容**:
{content}

---

**分析要求**:
"""

    # 根据内容特征添加特定要求
    if content_chars['has_molecular_data']:
        prompt += """
1. **分子信息提取**: 识别并列出所有分子名称、SMILES结构、分子式
2. **理化性质**: 提取分子量(MW)、脂水分配系数(cLogP)、极性表面积(TPSA)等数据
3. **数据验证**: 检查数据的一致性和合理性"""

    if content_chars['has_bioactivity']:
        prompt += """
4. **生物活性数据**: 提取IC50、EC50、Ki等生物活性指标
5. **活性关系**: 分析结构-活性关系(SAR)"""

    if content_chars['has_tables']:
        prompt += """
6. **表格数据**: 系统性整理表格中的结构化信息
7. **数据对比**: 比较不同化合物间的性质差异"""

    if content_chars['has_patents']:
        prompt += """
8. **专利信息**: 提取专利号、申请人、技术领域等关键信息
9. **技术要点**: 总结专利的核心技术和创新点"""

    # 根据查询意图添加特定指导
    if query_intent == 'extraction':
        prompt += """

**输出格式要求**:
- 使用结构化列表或表格形式
- 标注数据来源（文档名称、页码等）
- 对于数值数据，保留原始单位和精度
- 如发现数据缺失或不一致，明确标注"""

    elif query_intent == 'analysis':
        prompt += """

**分析深度要求**:
- 提供数据趋势和模式分析
- 识别异常值和特殊情况
- 给出可能的科学解释
- 提出进一步研究建议"""

    elif query_intent == 'summarization':
        prompt += """

**总结要求**:
- 突出最重要的发现和结论
- 使用分层结构组织信息
- 提供量化的统计信息
- 保持客观和准确性"""

    # 添加输出格式指导
    prompt += """

**回答格式**:
1. **直接回答**: 首先直接回答用户的具体问题
2. **详细分析**: 提供支撑性的详细信息和数据
3. **关键发现**: 突出最重要的发现和洞察
4. **数据总结**: 提供结构化的数据汇总
5. **建议**: 基于分析结果提出相关建议

请确保回答准确、全面且易于理解。使用中文回答，保持专业性和可读性。"""

    return prompt

def optimize_retrieval_parameters(query: str, existing_content: str = "") -> dict:
    """根据查询特征优化检索参数"""
    query_lower = query.lower()

    # 默认参数
    params = {
        'search_type': 'similarity',
        'k': 20,
        'additional_kwargs': {}
    }

    # 分析查询类型和复杂度
    is_specific_query = any(term in query_lower for term in [
        'smiles', 'ic50', 'ec50', 'molecular weight', 'compound', '化合物', '分子量'
    ])

    is_broad_query = any(term in query_lower for term in [
        '总结', 'summarize', '概述', 'overview', '所有', 'all', '全部'
    ])

    is_comparison_query = any(term in query_lower for term in [
        '比较', 'compare', '对比', 'versus', 'vs', '差异', 'difference'
    ])

    # 根据查询类型调整参数
    if is_specific_query:
        # 精确查询：减少检索数量，提高精度
        params['k'] = 15
        params['search_type'] = 'similarity'

    elif is_broad_query:
        # 广泛查询：增加检索数量，使用MMR避免重复
        params['k'] = 30
        params['search_type'] = 'mmr'
        params['additional_kwargs'] = {
            'lambda_mult': 0.7,  # 平衡相关性和多样性
            'fetch_k': 50       # 增加候选集
        }

    elif is_comparison_query:
        # 比较查询：中等数量，使用MMR确保多样性
        params['k'] = 25
        params['search_type'] = 'mmr'
        params['additional_kwargs'] = {
            'lambda_mult': 0.6,  # 更注重多样性
            'fetch_k': 40
        }

    # 根据现有内容长度调整
    if len(existing_content) > 10000:
        # 已有大量内容，减少新检索的数量
        params['k'] = max(10, params['k'] - 5)
    elif len(existing_content) < 1000:
        # 内容较少，增加检索数量
        params['k'] = min(40, params['k'] + 10)

    return params

def enhance_query_with_context(original_query: str, document_metadata: list) -> str:
    """基于文档上下文增强查询"""
    enhanced_query = original_query

    # 分析文档类型
    doc_types = set()
    for doc in document_metadata:
        filename = doc.get('filename', '').lower()
        if 'patent' in filename or 'wo' in filename:
            doc_types.add('patent')
        elif 'compound' in filename or 'molecule' in filename:
            doc_types.add('molecular')
        elif 'bioactivity' in filename or 'activity' in filename:
            doc_types.add('bioactivity')

    # 根据文档类型添加上下文关键词
    context_keywords = []
    if 'patent' in doc_types:
        context_keywords.extend(['patent', 'invention', 'claim'])
    if 'molecular' in doc_types:
        context_keywords.extend(['molecular', 'compound', 'structure'])
    if 'bioactivity' in doc_types:
        context_keywords.extend(['activity', 'bioassay', 'inhibition'])

    # 如果查询较短且文档有明确类型，添加上下文
    if len(original_query.split()) < 5 and context_keywords:
        enhanced_query = f"{original_query} {' '.join(context_keywords[:2])}"

    return enhanced_query



from typing import List
from pymilvus import Collection, connections
import contextlib

def _get_collection_name(vs) -> str:
    return (
        getattr(vs, "collection_name", None)
        or getattr(vs, "_collection_name", None)
        or "LangChainCollection"
    )

def _build_file_ids_expr(vectorstore, file_ids: List[str]) -> str:
    """
    返回适用于 Milvus 的 expr：
      - 如果有独立列:    file_id in ["a","b"]
      - 如果是 JSON 列:  metadata["file_id"] in ["a","b"]
    """
    if not file_ids:
        return ""  # 不加过滤

    # 转义并拼接
    def _q(s: str) -> str:
        return '"' + s.replace('\\', '\\\\').replace('"', '\\"') + '"'
    ids_part = ", ".join(_q(x) for x in file_ids)

    # 连接到当前向量库的 uri（Milvus-Lite 就是你的 vec.db）
    uri = getattr(vectorstore, "connection_args", {}).get("uri", "vec.db")
    with contextlib.suppress(Exception):
        connections.disconnect("default")
    connections.connect(alias="default", uri=uri)

    col_name = _get_collection_name(vectorstore)
    col = Collection(col_name)
    field_names = {f.name for f in col.schema.fields}

    if "file_id" in field_names:
        return f"file_id in [{ids_part}]"
    elif "metadata" in field_names:
        # metadata 是 JSON 列
        return f'metadata["file_id"] in [{ids_part}]'
    else:
        # 如果没有两者，给出一个永不命中的表达式，避免报错
        return 'file_id == "__NO_SUCH_FIELD__"'

try:
    from flask import Flask, render_template_string, request, jsonify, session, redirect, url_for, send_file
    from database_manager import db_manager
    from user_manager import user_manager, login_required, admin_required, track_usage
    from prompt_templates import prompt_manager
    from workflow_manager import create_workflow_manager
    from pdf_processor import pdf_processor

    # 导入新增的模块
    from smiles_extractor import get_smiles_extractor
    from chemical_segmentation import get_segmentation_processor
    from patent_analyzer import get_patent_analyzer
    from pdf_ocr_processor import get_pdf_ocr_processor

    # 导入Ce替换函数
    import sys
    parent_dir = os.path.join(os.path.dirname(__file__), '..')
    if parent_dir not in sys.path:
        sys.path.insert(0, parent_dir)
    from replace import replace_ce_with_allyl

    # 使用现有的化学结构处理器
    print("🔄 Using existing chemical structure processor")
    from chemical_structure_processor import create_chemical_processor
    chemical_processor = create_chemical_processor(db_manager)

    app = Flask(__name__)

    # 配置session密钥
    app.secret_key = 'bioai_secret_key_2024_secure_random_string'
    app.config['SESSION_TYPE'] = 'filesystem'
    app.config['PERMANENT_SESSION_LIFETIME'] = 3600  # 1小时

    # 配置上传文件夹
    UPLOAD_FOLDER = 'uploads'
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

    # 配置静态文件夹用于视频
    STATIC_FOLDER = 'static'
    if not os.path.exists(STATIC_FOLDER):
        os.makedirs(STATIC_FOLDER)

    VIDEO_FOLDER = os.path.join(STATIC_FOLDER, 'video')
    if not os.path.exists(VIDEO_FOLDER):
        os.makedirs(VIDEO_FOLDER)

    # 创建工作流管理器实例
    workflow_manager = create_workflow_manager(db_manager)
    app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max file size (支持大文件)

    def get_display_file_type(segment_file_type):
        """根据片段文件类型返回显示的文件类型"""
        type_mapping = {
            'layout_pdf': 'WPS PDF 文档',
            'model_pdf': 'WPS PDF 文档',
            'spans_pdf': 'WPS PDF 文档',
            'markdown': 'Markdown 源文件',
            'content_list_json': 'JSON 文件',
            'middle_json': 'JSON 文件',
            'content_list': 'JSON 文件',
            'middle': 'JSON 文件'
        }
        return type_mapping.get(segment_file_type, segment_file_type.upper())

    def generate_document_processing_results(filename, segments):
        """生成文档处理结果的三种格式"""
        from datetime import datetime

        # 1. 生成 Markdown 内容
        markdown_content = f"""# {filename} - Processing Results

## Document Information
- **Filename**: {filename}
- **Processing Date**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- **Total Segments**: {len(segments)}

## Extracted Content

"""

        for i, segment in enumerate(segments):
            segment_type = segment.get('type', 'text').upper()
            page_num = segment.get('page', 1)
            content = segment.get('text', '').strip()

            markdown_content += f"""### Segment {i+1} - {segment_type} (Page {page_num})

```
{content}
```

"""

        # 2. 生成 content_list.json
        content_list = {
            "document_info": {
                "filename": filename,
                "processing_date": datetime.now().isoformat(),
                "total_segments": len(segments)
            },
            "content_list": []
        }

        for i, segment in enumerate(segments):
            content_item = {
                "segment_id": i + 1,
                "type": segment.get('type', 'text'),
                "page": segment.get('page', 1),
                "content": segment.get('text', ''),
                "bbox": segment.get('bbox', []),
                "confidence": segment.get('confidence', 1.0)
            }
            content_list["content_list"].append(content_item)

        # 3. 生成 middle.json (中间处理数据)
        middle_json = {
            "metadata": {
                "filename": filename,
                "processing_timestamp": datetime.now().isoformat(),
                "processor_version": "长鑫 v2.0",
                "total_segments": len(segments)
            },
            "processing_pipeline": [
                "document_upload",
                "pdf_parsing",
                "content_extraction",
                "segment_analysis",
                "result_generation"
            ],
            "segments_metadata": [],
            "statistics": {
                "total_characters": sum(len(seg.get('text', '')) for seg in segments),
                "avg_confidence": sum(seg.get('confidence', 1.0) for seg in segments) / len(segments) if segments else 0,
                "page_distribution": {}
            }
        }

        # 统计页面分布
        page_counts = {}
        for segment in segments:
            page = str(segment.get('page', 1))  # 转换为字符串键
            page_counts[page] = page_counts.get(page, 0) + 1
        middle_json["statistics"]["page_distribution"] = page_counts

        # 添加段落元数据
        for i, segment in enumerate(segments):
            segment_meta = {
                "segment_index": i,
                "type": segment.get('type', 'text'),
                "page": segment.get('page', 1),
                "character_count": len(segment.get('text', '')),
                "confidence": segment.get('confidence', 1.0),
                "bbox": segment.get('bbox', []),
                "processing_notes": f"Extracted from page {segment.get('page', 1)}"
            }
            middle_json["segments_metadata"].append(segment_meta)

        return {
            "markdown": markdown_content,
            "content_list": content_list,
            "middle_json": middle_json
        }

    # HTML模板 - 第一部分
    HTML_TEMPLATE_PART1 = '''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>🔬 长鑫 Research Platform v2.0 Advanced</title>
        <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css" rel="stylesheet">
        <link href="https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700;900&family=Rajdhani:wght@300;400;500;600;700&display=swap" rel="stylesheet">
        <!-- ExcelJS库用于XLSX文件生成（支持图片嵌入） -->
        <script src="https://cdnjs.cloudflare.com/ajax/libs/exceljs/4.3.0/exceljs.min.js"></script>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            :root {
                /* 白色系科技风配色方案 */
                --primary: #0066FF;
                --secondary: #4A90E2;
                --accent: #007AFF;
                --success: #00C851;
                --warning: #FF8F00;
                --danger: #FF4444;

                /* 白色科技背景 */
                --background: #FFFFFF;
                --surface: #F8FAFB;
                --card: rgba(255, 255, 255, 0.95);
                --glass: rgba(0, 102, 255, 0.03);

                /* 白色主题文字颜色 */
                --text-primary: #1A1A1A;
                --text-secondary: #4A5568;
                --text-muted: #9CA3AF;

                /* 边框和效果 */
                --border: rgba(0, 102, 255, 0.15);
                --border-hover: rgba(0, 102, 255, 0.3);
                --glow: 0 0 20px rgba(0, 102, 255, 0.15);
                --glow-strong: 0 0 30px rgba(0, 102, 255, 0.25);

                /* 渐变效果 */
                --gradient-primary: linear-gradient(135deg, #0066FF 0%, #4A90E2 100%);
                --gradient-accent: linear-gradient(135deg, #007AFF 0%, #0066FF 100%);
                --gradient-surface: linear-gradient(135deg, #F8FAFB 0%, #FFFFFF 100%);

                /* 白色主题特有 */
                --shadow-light: 0 2px 8px rgba(0, 0, 0, 0.08);
                --shadow-medium: 0 4px 16px rgba(0, 0, 0, 0.12);
                --shadow-strong: 0 8px 32px rgba(0, 0, 0, 0.16);
            }

            body {
                font-family: 'Inter', 'SF Pro Display', -apple-system, BlinkMacSystemFont, sans-serif;
                background: var(--background);
                background-image:
                    radial-gradient(circle at 20% 80%, rgba(0, 102, 255, 0.04) 0%, transparent 50%),
                    radial-gradient(circle at 80% 20%, rgba(0, 122, 255, 0.04) 0%, transparent 50%),
                    radial-gradient(circle at 40% 40%, rgba(74, 144, 226, 0.03) 0%, transparent 50%),
                    linear-gradient(180deg, rgba(248, 250, 251, 0.8) 0%, rgba(255, 255, 255, 1) 100%);
                min-height: 100vh;
                color: var(--text-primary);
                overflow-x: hidden;
                position: relative;
                line-height: 1.6;
                letter-spacing: -0.01em;
            }

            /* 科技感文字样式 */
            .text-secondary {
                color: var(--text-secondary) !important;
                font-weight: 500 !important;
            }

            .text-muted {
                color: var(--text-muted) !important;
                font-weight: 400 !important;
            }

            /* 白色系科技感动态背景网格 */
            body::before {
                content: '';
                position: fixed;
                top: 0;
                left: 0;
                width: 100%;
                height: 100%;
                background-image:
                    linear-gradient(rgba(0, 102, 255, 0.08) 1px, transparent 1px),
                    linear-gradient(90deg, rgba(0, 102, 255, 0.08) 1px, transparent 1px);
                background-size: 60px 60px;
                animation: gridMove 40s linear infinite;
                pointer-events: none;
                z-index: -1;
                opacity: 0.6;
            }

            /* 白色主题科技感扫描线效果 */
            body::after {
                content: '';
                position: fixed;
                top: 0;
                left: -100%;
                width: 100%;
                height: 1px;
                background: linear-gradient(90deg, transparent, var(--primary), transparent);
                animation: scanLine 6s ease-in-out infinite;
                pointer-events: none;
                z-index: 1000;
                opacity: 0.7;
            }

            @keyframes gridMove {
                0% { transform: translate(0, 0); }
                100% { transform: translate(50px, 50px); }
            }

            @keyframes scanLine {
                0%, 100% { left: -100%; opacity: 0; }
                50% { left: 100%; opacity: 1; }
            }

            @keyframes fadeIn {
                0% {
                    opacity: 0;
                    transform: scale(0.8);
                }
                100% {
                    opacity: 1;
                    transform: scale(1);
                }
            }

            /* 白色系科技感导航栏 */
            .navbar {
                display: flex;
                justify-content: space-between;
                align-items: center;
                padding: 16px 32px;
                background: var(--card);
                backdrop-filter: blur(20px);
                border-bottom: 1px solid var(--border);
                margin-bottom: 24px;
                position: relative;
                overflow: hidden;
                box-shadow: var(--shadow-light);
            }

            .navbar::before {
                content: '';
                position: absolute;
                bottom: 0;
                left: -100%;
                width: 100%;
                height: 2px;
                background: var(--gradient-primary);
                animation: navGlow 4s ease-in-out infinite;
            }

            @keyframes navGlow {
                0%, 100% { left: -100%; }
                50% { left: 100%; }
            }

            .logo {
                display: flex;
                align-items: center;
                gap: 20px;
            }

            .logo i {
                font-size: 2.2em;
                color: var(--primary);
                filter: drop-shadow(0 2px 4px rgba(0, 102, 255, 0.3));
                animation: logoRotate 12s linear infinite;
            }

            @keyframes logoRotate {
                0% { transform: rotate(0deg); }
                100% { transform: rotate(360deg); }
            }

            .logo h1 {
                margin: 0;
                font-size: 2.0em;
                font-weight: 700;
                line-height: 1.2;
            }

            .brand {
                font-family: 'Inter', 'Helvetica Neue', 'Arial', sans-serif;
                font-weight: 900;
                font-size: 2.0em;
                color: var(--primary);
                letter-spacing: 0.05em;
                text-shadow: 0 2px 4px rgba(0, 102, 255, 0.3);
                display: inline-block;
                margin-right: 16px;
                vertical-align: middle;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            }

            .brand:hover {
                text-shadow: 0 4px 8px rgba(0, 102, 255, 0.5);
                transform: scale(1.02);
            }

            .subtitle {
                font-family: 'Inter', sans-serif;
                font-weight: 900;
                color: var(--text-secondary);
                font-size: 1.1em;
                opacity: 0.9;
                vertical-align: middle;
                display: inline-block;
                margin-left: 8px;
            }

            .status {
                display: flex;
                align-items: center;
            }

            .status i {
                color: var(--success);
                font-size: 0.8em;
                margin-right: 8px;
                animation: blink 2s infinite;
            }

            @keyframes blink {
                0%, 50% { opacity: 1; }
                51%, 100% { opacity: 0.3; }
            }

            /* 白色系科技感卡片样式 */
            .glass-card {
                background: var(--card);
                backdrop-filter: blur(20px);
                border: 1px solid var(--border);
                border-radius: 12px;
                padding: 24px;
                margin: 16px 0;
                box-shadow: var(--shadow-light);
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                position: relative;
                overflow: hidden;
            }

            .glass-card::before {
                content: '';
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                height: 2px;
                background: var(--gradient-primary);
                opacity: 0.8;
            }

            .glass-card:hover {
                transform: translateY(-2px);
                box-shadow: var(--shadow-medium);
                border-color: var(--border-hover);
            }

            .glass-card:hover::before {
                opacity: 1;
                box-shadow: var(--glow);
            }

            /* 用户信息栏样式 */
            .user-info-bar {
                background: rgba(0, 20, 40, 0.9);
                border-bottom: 1px solid var(--border);
                padding: 15px 40px;
                display: flex;
                justify-content: space-between;
                align-items: center;
                backdrop-filter: blur(10px);
                position: sticky;
                top: 0;
                z-index: 1000;
            }

            .user-info-left {
                display: flex;
                align-items: center;
                gap: 20px;
            }

            .welcome-text {
                font-size: 1.2em;
                font-weight: 600;
                color: var(--accent);
                text-shadow: 0 0 10px rgba(0, 212, 255, 0.5);
            }

            .user-details {
                color: var(--text-secondary);
                font-size: 0.9em;
                padding: 5px 12px;
                background: rgba(0, 212, 255, 0.1);
                border-radius: 15px;
                border: 1px solid rgba(0, 212, 255, 0.3);
            }

            .user-info-right {
                display: flex;
                gap: 10px;
            }

            .user-btn {
                padding: 8px 16px;
                background: rgba(0, 212, 255, 0.1);
                border: 1px solid rgba(0, 212, 255, 0.3);
                border-radius: 8px;
                color: var(--text-primary);
                cursor: pointer;
                transition: all 0.3s ease;
                font-size: 0.9em;
                display: flex;
                align-items: center;
                gap: 6px;
            }

            .user-btn:hover {
                background: rgba(0, 212, 255, 0.2);
                border-color: var(--accent);
                transform: translateY(-2px);
                box-shadow: 0 4px 12px rgba(0, 212, 255, 0.3);
            }

            /* 科技感标签页 */
            .tabs {
                display: flex;
                margin-bottom: 32px;
                padding: 0 32px;
                gap: 8px;
                flex-wrap: wrap;
            }

            .tab {
                padding: 12px 24px;
                background: var(--surface);
                border: 1px solid var(--border);
                border-radius: 8px;
                color: var(--text-secondary);
                font-family: 'Inter', sans-serif;
                font-weight: 500;
                cursor: pointer;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                text-decoration: none;
                font-size: 0.875rem;
                position: relative;
                overflow: hidden;
            }

            .tab::before {
                content: '';
                position: absolute;
                top: 0;
                left: -100%;
                width: 100%;
                height: 100%;
                background: var(--gradient-primary);
                opacity: 0;
                transition: all 0.3s ease;
                z-index: -1;
            }

            .tab.active {
                background: var(--gradient-primary);
                color: var(--text-primary);
                border-color: var(--primary);
                box-shadow: 0 4px 20px rgba(0, 212, 255, 0.3);
                transform: translateY(-1px);
            }

            .tab.active::before {
                left: 0;
                opacity: 0.1;
            }

            .tab:hover:not(.active) {
                color: var(--primary);
                border-color: var(--border-hover);
                background: rgba(0, 212, 255, 0.05);
                transform: translateY(-1px);
            }

            /* 白色系科技感按钮 */
            .neon-button {
                background: var(--gradient-primary);
                border: 1px solid var(--border);
                border-radius: 8px;
                padding: 12px 24px;
                color: white;
                font-family: 'Inter', sans-serif;
                font-weight: 500;
                font-size: 0.875rem;
                cursor: pointer;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                box-shadow: var(--shadow-light);
                margin: 4px;
                position: relative;
                overflow: hidden;
            }

            .neon-button::before {
                content: '';
                position: absolute;
                top: 0;
                left: -100%;
                width: 100%;
                height: 100%;
                background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.3), transparent);
                transition: left 0.5s ease;
            }

            .neon-button:hover {
                transform: translateY(-2px);
                box-shadow: var(--shadow-medium);
                border-color: var(--primary);
            }

            .neon-button:hover::before {
                left: 100%;
            }

            .neon-button:active {
                transform: translateY(0);
                box-shadow: var(--shadow-light);
            }

            /* 选中状态 - 灰色填充 */
            .neon-button.selected {
                background: linear-gradient(45deg, #A9A9A9, #808080) !important;
                color: white;
                opacity: 0.7;
                cursor: default;
                box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.2);
            }

            .neon-button.selected:hover {
                transform: none;
                box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.2);
                border-color: #808080;
            }

            .auto-button {
                background: linear-gradient(45deg, var(--success), var(--warning));
                box-shadow: 0 0 20px rgba(0, 99, 207, 0.4);
                animation: autoGlow 2s infinite alternate;
            }

            @keyframes autoGlow {
                0% { box-shadow: 0 0 20px rgba(0, 99, 207, 0.4); }
                100% { box-shadow: 0 0 30px rgba(0, 99, 207, 0.8); }
            }

            /* 小按钮样式 */
            .small-button {
                padding: 8px 16px;
                font-size: 0.9em;
                border: none;
                border-radius: 6px;
                cursor: pointer;
                transition: all 0.3s ease;
                color: black;
                font-weight: 500;
                text-decoration: none;
                display: inline-flex;
                align-items: center;
                justify-content: center;
                min-width: 80px;
            }

            /* 美化的删除按钮样式 - 蓝色主题 */
            .delete-btn {
                background: linear-gradient(135deg, #0063CF, #004A9F) !important;
                border: 1px solid rgba(0, 99, 207, 0.3) !important;
                color: white !important;
                transition: all 0.3s ease;
                min-width: auto !important;
                padding: 8px 12px !important;
                border-radius: 8px;
                font-size: 0.85em;
                position: relative;
                overflow: hidden;
                box-shadow: 0 2px 8px rgba(0, 99, 207, 0.2);
            }

            .delete-btn::before {
                content: '';
                position: absolute;
                top: 0;
                left: -100%;
                width: 100%;
                height: 100%;
                background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.2), transparent);
                transition: left 0.5s;
            }

            .delete-btn:hover::before {
                left: 100%;
            }

            .delete-btn:hover {
                background: linear-gradient(135deg, #004A9F, #003875) !important;
                border-color: rgba(0, 99, 207, 0.5) !important;
                transform: scale(1.05) translateY(-1px);
                box-shadow: 0 6px 20px rgba(0, 99, 207, 0.4);
            }

            .delete-btn:active {
                transform: scale(0.98) translateY(0px);
                box-shadow: 0 2px 8px rgba(16, 185, 129, 0.3);
            }

            .delete-btn i {
                font-size: 1.1em;
                transition: all 0.3s ease;
            }

            .delete-btn:hover i {
                transform: rotate(10deg) scale(1.1);
            }

            /* 内容区域 */
            .content {
                padding: 0 40px;
            }

            .section-header {
                display: flex;
                align-items: center;
                margin-bottom: 25px;
            }

            .section-header i {
                font-size: 1.5em;
                margin-right: 15px;
            }

            .section-header h3 {
                color: var(--text-primary);
                font-family: 'Orbitron', monospace;
                font-weight: 700;
                margin: 0;
            }

            /* 白色系科技感输入框 */
            .input-field {
                width: 100%;
                padding: 12px 16px;
                border-radius: 8px;
                border: 1px solid var(--border);
                background: var(--background);
                color: var(--text-primary);
                font-family: 'Inter', sans-serif;
                font-size: 0.875rem;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                margin-bottom: 16px;
                outline: none;
                box-shadow: var(--shadow-light);
            }

            .input-field:focus {
                border-color: var(--primary);
                box-shadow: 0 0 0 3px rgba(0, 102, 255, 0.1), var(--shadow-medium);
                background: rgba(0, 102, 255, 0.02);
            }

            .input-field::placeholder {
                color: var(--text-muted);
                font-weight: 400;
            }

            /* 白底黑字下拉框样式 */
            .white-select-field {
                width: 100%;
                padding: 15px;
                border-radius: 15px;
                border: 1px solid #ccc;
                background: white;
                color: black;
                font-family: 'Rajdhani', sans-serif;
                font-size: 1em;
                margin-bottom: 20px;
            }

            .white-select-field option {
                background: white;
                color: black;
            }

            /* 多选框样式 */
            .multi-select {
                width: 100%;
                min-height: 120px;
                padding: 15px;
                border-radius: 15px;
                border: 1px solid var(--border);
                background: var(--glass);
                color: var(--text-primary);
                font-family: 'Rajdhani', sans-serif;
                font-size: 1em;
                backdrop-filter: blur(10px);
                margin-bottom: 20px;
                overflow-y: auto;
            }

            .multi-select option {
                background: var(--surface);
                color: var(--text-primary);
                padding: 8px;
                margin: 2px 0;
                border-radius: 5px;
            }

            .multi-select option:checked {
                background: linear-gradient(45deg, var(--primary), var(--accent));
                color: var(--text-primary);
            }

            /* 下拉框样式 */
            .select-field {
                width: 100%;
                padding: 15px;
                border-radius: 15px;
                border: 1px solid var(--border);
                background: var(--glass);
                color: var(--text-primary);
                font-family: 'Rajdhani', sans-serif;
                font-size: 1em;
                backdrop-filter: blur(10px);
                margin-bottom: 20px;
                cursor: pointer;
            }

            .select-field option {
                background: var(--surface);
                color: var(--text-primary);
            }

            /* 模板卡片 */
            .template-grid {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
                gap: 15px;
                margin: 20px 0;
            }

            .template-card {
                background: linear-gradient(135deg, var(--glass), rgba(255, 255, 255, 0.02));
                border: 1px solid var(--border);
                border-radius: 15px;
                padding: 20px;
                cursor: pointer;
                transition: all 0.3s ease;
                position: relative;
            }

            .template-card:hover {
                transform: translateY(-3px);
                border-color: var(--primary);
                box-shadow: 0 0 25px rgba(0, 212, 255, 0.3);
            }

            .template-card.selected {
                border-color: var(--success);
                background: linear-gradient(135deg, rgba(0, 99, 207, 0.1), rgba(0, 99, 207, 0.05));
            }

            .template-card h4 {
                color: var(--primary);
                font-family: 'Orbitron', monospace;
                margin-bottom: 10px;
                font-size: 1.1em;
            }

            .template-card p {
                color: var(--text-secondary);
                font-size: 0.9em;
                line-height: 1.4;
                margin-bottom: 10px;
            }

            .template-tags {
                display: flex;
                flex-wrap: wrap;
                gap: 5px;
                margin-top: 10px;
            }

            .template-tag {
                background: var(--accent);
                color: var(--text-primary);
                padding: 3px 8px;
                border-radius: 12px;
                font-size: 0.8em;
                font-family: 'Orbitron', monospace;
            }

            /* 工作流卡片 */
            .workflow-grid {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(350px, 1fr));
                gap: 20px;
                margin: 20px 0;
            }

            .workflow-card {
                background: linear-gradient(135deg, var(--glass), rgba(255, 255, 255, 0.02));
                border: 1px solid var(--border);
                border-radius: 15px;
                padding: 25px;
                cursor: pointer;
                transition: all 0.3s ease;
                position: relative;
            }

            .workflow-card:hover {
                transform: translateY(-3px);
                border-color: var(--success);
                box-shadow: 0 0 25px rgba(0, 99, 207, 0.3);
            }

            .workflow-card.selected {
                border-color: var(--warning);
                background: linear-gradient(135deg, rgba(255, 215, 0, 0.1), rgba(255, 215, 0, 0.05));
            }

            .workflow-card h4 {
                color: var(--success);
                font-family: 'Orbitron', monospace;
                margin-bottom: 10px;
                font-size: 1.2em;
            }

            .workflow-card p {
                color: var(--text-secondary);
                font-size: 0.9em;
                line-height: 1.4;
                margin-bottom: 15px;
            }

            .workflow-info {
                display: flex;
                justify-content: space-between;
                align-items: center;
                margin-top: 15px;
                padding-top: 15px;
                border-top: 1px solid var(--border);
            }

            .workflow-time {
                color: var(--warning);
                font-family: 'Orbitron', monospace;
                font-size: 0.9em;
            }

            .workflow-steps {
                color: var(--accent);
                font-size: 0.9em;
            }

            /* 进度条 */
            .progress-container {
                width: 100%;
                height: 20px;
                background: var(--surface);
                border-radius: 10px;
                overflow: hidden;
                margin: 15px 0;
                position: relative;
            }

            .progress-bar {
                height: 100%;
                background: linear-gradient(45deg, var(--primary), var(--success));
                border-radius: 10px;
                transition: width 0.3s ease;
                position: relative;
            }

            .progress-bar::after {
                content: '';
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                bottom: 0;
                background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.3), transparent);
                animation: progressShine 2s infinite;
            }

            @keyframes progressShine {
                0% { transform: translateX(-100%); }
                100% { transform: translateX(100%); }
            }

            /* 下载成功消息动画 */
            @keyframes slideInRight {
                0% {
                    transform: translateX(100%);
                    opacity: 0;
                }
                100% {
                    transform: translateX(0);
                    opacity: 1;
                }
            }

            @keyframes slideOutRight {
                0% {
                    transform: translateX(0);
                    opacity: 1;
                }
                100% {
                    transform: translateX(100%);
                    opacity: 0;
                }
            }

            .progress-text {
                text-align: center;
                margin-top: 10px;
                color: var(--text-secondary);
                font-family: 'Orbitron', monospace;
            }

            /* 白色系科技感数据表格 */
            .data-table {
                width: 100%;
                border-collapse: collapse;
                margin: 16px 0;
                background: var(--background);
                border-radius: 8px;
                overflow: hidden;
                border: 1px solid var(--border);
                box-shadow: var(--shadow-light);
            }

            .data-table th,
            .data-table td {
                padding: 12px 16px;
                text-align: left;
                border-bottom: 1px solid var(--border);
                font-size: 0.875rem;
            }

            .data-table th {
                background: var(--gradient-primary);
                color: white;
                font-family: 'Inter', sans-serif;
                font-weight: 600;
                font-size: 0.8rem;
                letter-spacing: 0.05em;
                position: sticky;
                top: 0;
                z-index: 10;
            }

            .data-table td {
                color: var(--text-secondary);
                font-family: 'Inter', sans-serif;
                font-weight: 400;
            }

            .data-table tr:hover {
                background: rgba(0, 102, 255, 0.04);
            }

            .data-table tr:nth-child(even) {
                background: rgba(0, 102, 255, 0.02);
            }

            /* 白色系科技感统计卡片 */
            .stats-grid {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
                gap: 16px;
                margin: 16px 0;
            }

            .stat-card {
                background: var(--card);
                border: 1px solid var(--border);
                border-radius: 12px;
                padding: 24px;
                text-align: center;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                position: relative;
                overflow: hidden;
                box-shadow: var(--shadow-light);
            }

            .stat-card::before {
                content: '';
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                height: 3px;
                background: var(--gradient-primary);
                opacity: 0.8;
            }

            .stat-card:hover {
                transform: translateY(-4px);
                border-color: var(--border-hover);
                box-shadow: var(--shadow-strong);
            }

            .stat-card:hover::before {
                opacity: 1;
                box-shadow: var(--glow);
            }

            .stat-card i {
                font-size: 2.5em;
                margin-bottom: 16px;
                display: block;
                color: var(--primary);
                filter: drop-shadow(0 2px 4px rgba(0, 102, 255, 0.3));
            }

            .stat-card .stat-value {
                font-size: 2.5em;
                font-weight: 700;
                color: var(--primary);
                font-family: 'JetBrains Mono', monospace;
                margin-bottom: 8px;
            }

            .stat-card .stat-label {
                color: var(--text-secondary);
                font-size: 0.875rem;
                font-weight: 500;
                text-transform: uppercase;
                letter-spacing: 0.05em;
            }

            /* 白色系科技感结果显示 */
            .result-box {
                padding: 20px;
                background: var(--card);
                border-radius: 8px;
                border: 1px solid var(--border);
                margin: 16px 0;
                font-family: 'Inter', sans-serif;
                line-height: 1.6;
                position: relative;
                box-shadow: var(--shadow-light);
            }

            .success-box {
                background: var(--card);
                border: 1px solid var(--success);
                box-shadow: 0 0 20px rgba(0, 200, 81, 0.1);
            }

            .hidden {
                display: none;
            }

            /* 白色系科技感进度条 */
            .progress-bar {
                width: 100%;
                height: 6px;
                background: var(--surface);
                border-radius: 3px;
                overflow: hidden;
                border: 1px solid var(--border);
                position: relative;
                box-shadow: inset 0 1px 3px rgba(0, 0, 0, 0.1);
            }

            .progress-fill {
                height: 100%;
                background: var(--gradient-primary);
                width: 0%;
                transition: width 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                border-radius: 3px;
                position: relative;
                overflow: hidden;
            }

            .progress-fill::after {
                content: '';
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                bottom: 0;
                background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.3), transparent);
                animation: progressShine 2s ease-in-out infinite;
            }

            @keyframes progressShine {
                0% { transform: translateX(-100%); }
                100% { transform: translateX(100%); }
            }

            /* 科技感加载动画 */
            .loading-spinner {
                display: inline-block;
                width: 20px;
                height: 20px;
                border: 2px solid var(--border);
                border-radius: 50%;
                border-top-color: var(--primary);
                animation: spin 1s ease-in-out infinite;
            }

            @keyframes spin {
                to { transform: rotate(360deg); }
            }

            /* 白色系科技感滚动条 */
            ::-webkit-scrollbar {
                width: 8px;
                height: 8px;
            }

            ::-webkit-scrollbar-track {
                background: var(--surface);
                border-radius: 4px;
                border: 1px solid var(--border);
            }

            ::-webkit-scrollbar-thumb {
                background: var(--gradient-primary);
                border-radius: 4px;
                border: 1px solid rgba(255, 255, 255, 0.2);
            }

            ::-webkit-scrollbar-thumb:hover {
                background: var(--primary);
                box-shadow: 0 0 8px rgba(0, 102, 255, 0.3);
            }

            /* 响应式设计 */
            @media (max-width: 768px) {
                .navbar {
                    padding: 15px 20px;
                    flex-direction: column;
                    gap: 15px;
                }

                .logo h1 {
                    font-size: 1.8em;
                }

                .content {
                    padding: 0 20px;
                }

                .tabs {
                    padding: 0 20px;
                }

                .tab {
                    font-size: 0.8em;
                    padding: 12px 20px;
                }

                .template-grid,
                .workflow-grid,
                .stats-grid {
                    grid-template-columns: 1fr;
                }
            }
        </style>
    </head>
    <body>
        <!-- 顶部导航栏 -->
        <div class="navbar">
            <div class="logo">
                <div style="display: inline-block; margin-right: 20px; vertical-align: middle;">
                    <svg xmlns="http://www.w3.org/2000/svg" xmlns:xlink="http://www.w3.org/1999/xlink" width="105px" height="30px" viewBox="0 0 141 39" version="1.1">
                        <defs>
                            <polygon id="path-1" points="0 0 24.9151012 0 24.9151012 37.0113062 0 37.0113062"></polygon>
                            <polygon id="path-3" points="0 0 31.0344534 0 31.0344534 35.5884897 0 35.5884897"></polygon>
                        </defs>
                        <g stroke="none" stroke-width="1" fill="none" fill-rule="evenodd">
                            <g transform="translate(-153.000000, -31.000000)">
                                <g transform="translate(153.177057, 31.500000)">
                                    <path d="M49.8203267,9.42089936 C51.8844629,6.86497909 54.4639345,4.87349288 57.5587416,3.4436465 C60.4324055,2.10321541 63.4250237,1.40865017 66.5381928,1.36072198 C66.7178218,1.35835406 66.8611257,1.51483083 66.8447595,1.6940606 L66.1765399,9.33547581 C66.1322314,9.85041212 65.7198832,10.2607645 65.2037494,10.2954927 C62.8765566,10.4579571 60.7166184,11.2591022 58.7227371,12.6997263 C56.2769893,14.432946 54.9425458,16.5996703 54.7150158,19.2010965 C54.4866875,21.802922 55.4431118,23.9720412 57.584688,25.7052609 C59.416504,27.1806133 61.5540884,27.9689848 63.9962437,28.0703753 C64.2900368,28.0823506 64.5147726,28.3366254 64.4888262,28.6296201 L63.8309851,36.1468918 C63.7894708,36.6251039 63.3811144,36.9903495 62.902503,36.9715882 C59.9961067,36.8598191 57.2992775,36.1704431 54.8072253,34.9018636 C51.9619029,33.4752106 49.7305122,31.4825269 48.1142509,28.9278042 C46.0097981,31.4825269 43.4403059,33.4652313 40.4045767,34.8739213 C37.485806,36.2291219 34.47762,36.9180987 31.3788211,36.9707899 C31.1768384,36.9739833 31.0179665,36.7995436 31.0359294,36.5991575 L31.7041491,28.9529522 C31.7476592,28.4623656 32.1396495,28.0703753 32.6310344,28.0392397 C34.9733959,27.8891497 37.1469061,27.1203378 39.1507667,25.7052609 C41.6320412,23.9720412 42.9860442,21.802922 43.2143726,19.2010965 C43.4427009,16.5996703 42.4659186,14.432946 40.2880175,12.6997263 C38.5424234,11.2575055 36.5182048,10.4563604 34.2153617,10.2954927 C33.735952,10.2623612 33.37869,9.8376385 33.4210026,9.35822881 L34.0888231,1.71122514 C34.106786,1.51163743 34.2740405,1.35835406 34.4744266,1.36154747 C37.5732254,1.40984769 40.4648522,2.11558985 43.1532987,3.47318548 C45.9407407,4.88187557 48.1637487,6.86497909 49.8203267,9.42089936" fill="#0066FF"></path>
                                    <g transform="translate(115.084899, 0.000000)">
                                        <mask id="mask-2" fill="#fff">
                                            <use xlink:href="#path-1"></use>
                                        </mask>
                                        <path d="M15.7243301,15.6538241 C15.7335111,15.5552277 15.8149429,15.4793844 15.9135393,15.4793844 L16.4376566,15.4793844 L23.5365924,15.4793844 C24.002031,15.4793844 24.3904287,15.1237191 24.4303462,14.6590789 L24.9113526,9.08778746 C24.9616487,8.50099958 24.4994035,7.99883689 23.911019,7.99883689 L17.0907076,7.99883689 L16.6168864,7.99883689 C16.5043189,7.99883689 16.4165003,7.90183727 16.4260805,7.7908665 L16.6847462,4.89524794 L17.1557732,0.208529244 C17.1665509,0.0967601237 17.0791315,-0.000239505257 16.966564,-0.000239505257 L15.4261461,-0.000239505257 L9.60257575,-0.000239505257 L9.1111908,-0.000239505257 C8.95950414,-0.000239505257 8.83216718,0.114723018 8.81939356,0.267208031 L8.15795988,7.82359888 C8.14997637,7.92299356 8.06694788,7.99883689 7.96835155,7.99883689 L7.41070348,7.99883689 L1.56877109,7.99883689 C1.02629168,7.99883689 0.575622626,8.41358016 0.526923224,8.95166864 L0.436310402,10.0011008 L0.433915349,9.99351651 L0.00160836054,14.9788184 C-0.020346288,15.2410767 0.185229057,15.46701 0.449084015,15.46701 L2.2345957,15.46701 L7.48814351,15.46701 L7.48814351,15.4793844 L6.50736948,26.6814444 C6.1984077,30.2293156 6.70456215,32.8622768 8.01545425,34.5108713 C9.33432986,36.1706427 11.5984529,37.0113062 14.7443544,37.0113062 L23.4256216,37.0113062 C23.9393604,37.0113062 24.3696715,36.6181184 24.41398,36.1047788 L24.8965831,30.102777 C24.9412907,29.5934292 24.5393211,29.1547354 24.0259815,29.1547354 L18.345316,29.1547354 C18.0299674,29.1547354 17.7517421,29.1491469 17.4998624,29.1407643 L17.8347706,29.1407643 C15.1706738,29.1407643 14.7922555,27.2462777 14.7914571,26.1708991 C14.7914571,26.1708991 14.7914571,26.1605205 14.7922555,26.1549321 C14.7982431,26.0571341 14.8062266,25.9565419 14.8154077,25.8507604 L15.7243301,15.6538241 Z" fill="#0066FF" mask="url(#mask-2)"></path>
                                    </g>
                                    <g transform="translate(0.000000, 1.362461)">
                                        <mask id="mask-4" fill="#fff">
                                            <use xlink:href="#path-3"></use>
                                        </mask>
                                        <path d="M6.17061343,5.21922311 C9.79752137,1.73721585 13.933378,-0.00318901688 18.5833725,0 L30.6528408,0.00838707053 C30.876379,0.00998377224 31.0532137,0.200789627 31.0328558,0.424727042 L30.3502658,8.15076745 C30.3103483,8.60103733 29.9347242,8.94712243 29.4820592,8.94512655 L17.7942027,8.93674387 C15.5823717,8.93674387 13.6204244,9.80495042 11.9107561,11.546952 C10.2018861,13.287756 9.23907495,15.3778386 9.02391939,17.8128087 C8.80876384,20.250573 9.40313605,22.3278819 10.8062377,24.0535173 C12.2097385,25.7787535 14.0172048,26.64217 16.2314309,26.6441658 L27.5819842,26.6509518 C28.1923234,26.6525485 28.6701364,27.1742708 28.6162477,27.7834125 L27.9544149,35.2791287 C27.938847,35.4539675 27.7915513,35.5884897 27.6163133,35.5884897 L15.4414627,35.580107 C10.7906698,35.5769136 6.96137996,33.8412988 3.94840383,30.374061 C0.937024399,26.9052266 -0.351114706,22.7166788 0.0815914575,17.8056235 C0.515495148,12.8977616 2.54450385,8.70123036 6.17061343,5.21922311" fill="#0066FF" mask="url(#mask-4)"></path>
                                    </g>
                                    <path d="M28.1295331,26.5672803 L21.6469242,18.8723755 C21.4796697,18.6723887 21.5056161,18.3758013 21.7040063,18.2077485 L29.4021044,11.7095717 C29.6016921,11.541918 29.899477,11.5678644 30.067929,11.7666538 L36.5493404,19.4623568 C36.7169941,19.660747 36.6926444,19.9581327 36.4930567,20.1253872 L28.7945594,26.623564 C28.59577,26.7924152 28.2983844,26.7676663 28.1295331,26.5672803" fill="#0066FF"></path>
                                    <path d="M109.179549,5.54131337 C107.031586,2.89238523 104.307613,1.5056498 101.011621,1.37392191 L101.031181,1.36074912 L82.3078576,1.36074912 L82.3250221,1.37272438 C79.0038826,1.48529185 76.0208446,2.87282564 73.3818959,5.54131337 C70.6255895,8.32755786 69.0715996,11.6846232 68.7235186,15.6077193 L66.9292251,35.7644817 C66.9080688,36.0498921 66.8885092,36.201978 66.8853158,36.3768168 L66.8853158,36.386397 L66.8469949,36.8338727 C66.8370155,36.9448435 66.9260317,37.0418431 67.0370024,37.0418431 L68.0225666,37.0418431 L73.6525368,37.0418431 L74.5458914,37.0418431 C74.6444877,37.0418431 74.727117,36.9652014 74.7351006,36.8666051 L75.2871602,30.3029635 L76.5840811,15.86878 C76.7217967,14.3295596 77.3161689,13.0274493 78.3687945,11.9636468 C79.4210209,10.8998443 80.6229381,10.3681426 81.9769412,10.3681426 C83.3636766,10.3681426 84.4869563,10.8998443 85.3503727,11.9636468 C86.2125916,13.0274493 86.5754421,14.3295596 86.4385249,15.86878 L85.4349979,27.0013836 L85.4130433,26.9926017 L84.5512235,36.8338727 C84.5408449,36.9448435 84.6286635,37.0418431 84.741231,37.0418431 L85.749149,37.0418431 L91.3779217,37.0418431 L92.2489225,37.0418431 C92.3483171,37.0418431 92.4305473,36.9652014 92.4381316,36.8666051 L92.70518,33.7670079 L94.3106635,15.86878 C94.4475807,14.3295596 95.0427513,13.0274493 96.0941794,11.9636468 C97.1452083,10.8998443 98.3499197,10.3681426 99.702326,10.3681426 C101.08986,10.3681426 102.211942,10.8998443 103.076157,11.9636468 C103.939174,13.0274493 104.302024,14.3295596 104.164309,15.86878 L103.160782,27.004577 L103.130445,26.9926017 L102.268625,36.8338727 C102.258246,36.9448435 102.346464,37.0418431 102.457834,37.0418431 L103.474534,37.0418431 L109.104504,37.0418431 L109.966324,37.0418431 C110.06492,37.0418431 110.147949,36.9667981 110.15713,36.8666051 L110.304425,35.198451 L112.047625,15.6077193 C112.397302,11.6846232 111.44008,8.32755786 109.179549,5.54131337" fill="#0066FF"></path>
                                </g>
                            </g>
                        </g>
                    </svg>
                </div>
                <h1>
                    <span class="subtitle">Research Platform Advanced</span>
                </h1>
            </div>
            <div class="status">
                <i class="fas fa-circle"></i>
                <span style="font-family: 'Orbitron', monospace;">AI Workflow Ready</span>
            </div>
        </div>
    '''

    # HTML模板 - 第二部分（标签页和内容）
    HTML_TEMPLATE_PART2 = '''
        <!-- 用户信息栏 -->
        <div class="user-info-bar">
            <div class="user-info-left">
                <span class="welcome-text">🔬 长鑫 Research Platform</span>
                <span class="user-details" id="userDetails">加载中...</span>
            </div>
            <div class="user-info-right">
                <button class="user-btn" id="adminBtn" onclick="window.location.href='/admin'" style="display: none;">
                    <i class="fas fa-cog"></i> 管理面板
                </button>
                <button class="user-btn" onclick="logout()">
                    <i class="fas fa-sign-out-alt"></i> 退出登录
                </button>
            </div>
        </div>

        <!-- 标签页 -->
        <div class="tabs">
            <a href="#document" class="tab active" onclick="showTab('document')">📄 Document Processing</a>
            <a href="#database" class="tab" onclick="showTab('database')">🗄️ Database View</a>


            <a href="#molecule" class="tab" onclick="showTab('molecule')">🔬 Molecule Processing</a>


        </div>

        <!-- 主要内容 -->
        <div class="content">
            <!-- 文档处理 -->
            <div id="document-tab" class="tab-content">
                <div class="glass-card">
                    <div class="section-header">
                        <i class="fas fa-file-upload" style="color: var(--primary);"></i>
                        <h3>Document Upload & Processing</h3>
                    </div>

                    <!-- 文件上传区域 -->
                    <input type="file" id="file-upload" class="input-field" accept=".pdf,.docx,.txt" multiple>
                    <div style="margin-bottom: 20px;">
                        <label style="color: var(--text-secondary); font-size: 0.9em;">
                            <i class="fas fa-info-circle" style="margin-right: 5px;"></i>
                            Support multiple file selection for batch processing (PDF, DOCX, TXT)
                        </label>
                    </div>

                    <!-- PDF 处理选项 -->
                    <div id="pdf-options" class="hidden" style="margin-bottom: 20px; padding: 15px; border: 1px solid var(--border); border-radius: 8px; background: var(--glass);">
                        <h4 style="color: var(--accent); font-family: 'Orbitron', monospace; margin-bottom: 15px;">
                            <i class="fas fa-file-pdf" style="margin-right: 10px;"></i>
                            PDF Processing Options
                        </h4>
                        <div style="display: flex; gap: 20px; flex-wrap: wrap;">
                            <label style="color: var(--text-secondary); display: flex; align-items: center;">
                                <input type="checkbox" id="enable-ocr" checked style="margin-right: 8px;">
                                <i class="fas fa-eye" style="margin-right: 5px;"></i>
                                Enable OCR
                            </label>
                            <label style="color: var(--text-secondary); display: flex; align-items: center;">
                                <input type="checkbox" id="enable-formula" checked style="margin-right: 8px;">
                                <i class="fas fa-square-root-alt" style="margin-right: 5px;"></i>
                                Formula Recognition
                            </label>
                            <label style="color: var(--text-secondary); display: flex; align-items: center;">
                                <input type="checkbox" id="enable-table" style="margin-right: 8px;">
                                <i class="fas fa-table" style="margin-right: 5px;"></i>
                                Table Recognition
                            </label>
                        </div>
                        <div style="margin-top: 10px;">
                            <label style="color: var(--text-secondary); font-size: 0.9em;">
                                <i class="fas fa-magic" style="margin-right: 5px;"></i>
                                PDF files will be automatically segmented and analyzed using AI-powered processing
                            </label>
                        </div>
                    </div>

                    <!-- 处理按钮 -->
                    <button class="neon-button" onclick="processDocument()">
                        <i class="fas fa-play" style="margin-right: 10px;"></i>
                        Process & Store Documents
                    </button>

                    <!-- 处理进度 -->
                    <div id="processing-progress" class="hidden" style="margin: 20px 0;">
                        <h4 style="color: var(--primary); margin-bottom: 10px;">
                            <i class="fas fa-spinner fa-spin" style="margin-right: 10px;"></i>
                            Processing Documents...
                        </h4>
                        <div class="progress-bar">
                            <div id="progress-fill" class="progress-fill"></div>
                        </div>
                        <div id="progress-text" style="color: var(--text-secondary); margin-top: 10px;">Initializing...</div>
                    </div>

                    <!-- 处理结果 -->
                    <div id="document-result" class="result-box hidden"></div>
                </div>
            </div>

            <!-- 数据库视图 -->
            <div id="database-tab" class="tab-content hidden">
                <div class="glass-card">
                    <div class="section-header">
                        <i class="fas fa-database" style="color: var(--success);"></i>
                        <h3>Database Overview</h3>
                    </div>

                    <!-- 数据库连接信息 -->
                    <div class="info-box" id="database-connection-info" style="margin-bottom: 20px;">
                        <!-- 连接信息将通过JavaScript动态加载 -->
                    </div>

                    <!-- 数据库统计 -->
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px;">
                        <h4 style="color: var(--primary); margin: 0; font-family: 'Orbitron', monospace;">数据库统计</h4>
                        <div style="display: flex; gap: 10px;">
                            <button onclick="useMagicPdfContent()"
                                    style="background: linear-gradient(45deg, var(--accent), var(--secondary)); color: white; border: none; padding: 8px 16px; border-radius: 4px; cursor: pointer; font-size: 0.9em;">
                                <i class="fas fa-magic" style="margin-right: 5px;"></i>
                                使用Magic PDF内容
                            </button>
                            <button onclick="clearDatabase()"
                                    style="background: #e74c3c; color: white; border: none; padding: 8px 16px; border-radius: 4px; cursor: pointer; font-size: 0.9em;">
                                <i class="fas fa-trash" style="margin-right: 5px;"></i>
                                清空数据库
                            </button>
                        </div>
                    </div>

                    <!-- 文件筛选器 -->
                    <div style="margin-bottom: 20px; padding: 15px; background: rgba(0, 99, 207, 0.05); border: 1px solid rgba(0, 99, 207, 0.2); border-radius: 10px;">
                        <div style="margin-bottom: 15px;">
                            <label style="color: var(--text-secondary); margin-bottom: 5px; display: block;">
                                <i class="fas fa-filter" style="margin-right: 5px;"></i>
                                筛选文件 (Filter by Filename):
                            </label>
                            <select id="file-filter-select" class="white-select-field" onchange="onFileFilterChange()" style="margin-bottom: 10px;">
                                <option value="">🗂️ 所有文件 (总体统计)</option>
                            </select>
                            <div style="display: flex; gap: 10px;">
                                <button class="neon-button" onclick="loadFileFilterOptions()" style="font-size: 0.8em; padding: 5px 10px; background: linear-gradient(45deg, var(--primary), var(--secondary));">
                                    <i class="fas fa-refresh" style="margin-right: 5px;"></i>
                                    刷新文件列表
                                </button>
                                <button class="neon-button" onclick="resetFileFilter()" style="font-size: 0.8em; padding: 5px 10px; background: linear-gradient(45deg, var(--warning), #FF8F00);">
                                    <i class="fas fa-undo" style="margin-right: 5px;"></i>
                                    重置筛选
                                </button>
                            </div>
                        </div>
                        <div id="filter-status" style="margin-top: 8px; font-size: 0.85em; color: var(--text-secondary);">
                            当前显示: 所有文件的统计数据
                        </div>
                    </div>

                    <div class="stats-grid" id="database-stats">
                        <!-- 统计数据将通过JavaScript动态加载 -->
                    </div>

                    <!-- 文档表格 -->
                    <h4 style="color: var(--primary); margin: 30px 0 15px 0; font-family: 'Orbitron', monospace;">
                        <i class="fas fa-file-alt" style="margin-right: 10px;"></i>
                        Stored Documents
                    </h4>
                    <div style="overflow-x: auto;">
                        <table class="data-table" id="documents-table">
                            <thead>
                                <tr>
                                    <th>ID</th>
                                    <th>Filename</th>
                                    <th>Type</th>
                                    <th>Upload Time</th>
                                    <th>Status</th>
                                    <th>Size</th>
                                    <th>Pages</th>
                                    <th>Actions</th>
                                </tr>
                            </thead>
                            <tbody id="documents-tbody">
                                <!-- 数据将通过JavaScript动态加载 -->
                            </tbody>
                        </table>
                    </div>

                    <!-- 文档片段详情 -->
                    <div id="document-segments-section" class="hidden" style="margin-top: 30px;">
                        <h4 style="color: var(--accent); margin: 20px 0 15px 0; font-family: 'Orbitron', monospace;">
                            <i class="fas fa-puzzle-piece" style="margin-right: 10px;"></i>
                            Document Segments
                        </h4>
                        <div id="document-segments-container">
                            <!-- 文档片段将通过JavaScript动态加载 -->
                        </div>
                    </div>

                    <!-- 处理结果详情 -->
                    <div id="processing-results-section" class="hidden" style="margin-top: 30px;">
                        <h4 style="color: var(--success); margin: 20px 0 15px 0; font-family: 'Orbitron', monospace;">
                            <i class="fas fa-file-code" style="margin-right: 10px;"></i>
                            Processing Results
                        </h4>
                        <div id="processing-results-container">
                            <!-- 处理结果将通过JavaScript动态加载 -->
                        </div>
                    </div>

                    <button class="neon-button" onclick="refreshDatabase()">
                        <i class="fas fa-sync-alt" style="margin-right: 10px;"></i>
                        Refresh Database
                    </button>
                </div>
            </div>





            <!-- 分子处理 -->
            <div id="molecule-tab" class="tab-content hidden">
                <div class="glass-card">
                    <div class="section-header">
                        <i class="fas fa-atom" style="color: var(--secondary);"></i>
                        <h3>Molecular Structure Processing</h3>
                    </div>

                    <!-- SMILES提取部分 -->
                    <div style="margin-bottom: 30px; padding: 20px; background: var(--glass); border-radius: 10px; border: 1px solid var(--secondary);">
                        <h4 style="color: var(--secondary); font-family: 'Orbitron', monospace; margin-bottom: 15px;">
                            <i class="fas fa-atom" style="margin-right: 10px;"></i>
                            SMILES Extraction from Chemical Structures
                        </h4>

                        <div style="display: flex; gap: 15px; margin-bottom: 15px; flex-wrap: wrap;">
                            <button id="extract-db-btn" class="neon-button" onclick="selectMoleculeButton('extract-db-btn', 'extractSmilesFromDatabase')" style="background: linear-gradient(45deg, var(--secondary), var(--accent)); transition: all 0.3s ease;">
                                <i class="fas fa-database" style="margin-right: 10px;"></i>
                                Extract from Database Images
                            </button>
                            <button id="view-smiles-btn" class="neon-button" onclick="selectMoleculeButton('view-smiles-btn', 'loadSmilesResults')" style="background: linear-gradient(45deg, var(--secondary), var(--accent)); transition: all 0.3s ease;">
                                <i class="fas fa-list" style="margin-right: 10px;"></i>
                                View SMILES Results
                            </button>
                            <button id="view-segmented-btn" class="neon-button" onclick="selectMoleculeButton('view-segmented-btn', 'loadSegmentedImages')" style="background: linear-gradient(45deg, var(--secondary), var(--accent)); transition: all 0.3s ease;">
                                <i class="fas fa-images" style="margin-right: 10px;"></i>
                                View Segmented Images
                            </button>
                        </div>

                        <div style="margin-bottom: 15px;">
                            <label style="color: var(--text-secondary); margin-bottom: 5px; display: block;">
                                <i class="fas fa-filter" style="margin-right: 5px;"></i>
                                Filter by Filename (Optional):
                            </label>
                            <select id="smiles-filename-filter" class="white-select-field" style="margin-bottom: 10px;">
                                <option value="">Select a filename to filter results...</option>
                                <!-- 文件名选项将通过JavaScript动态加载 -->
                            </select>
                            <div style="display: flex; gap: 10px;">
                                <button class="neon-button" onclick="loadImageFilenames()" style="font-size: 0.8em; padding: 5px 10px; background: linear-gradient(45deg, var(--primary), var(--secondary));">
                                    <i class="fas fa-refresh" style="margin-right: 5px;"></i>
                                    Load Filenames
                                </button>
                                <button class="neon-button" onclick="clearFilenameFilter()" style="font-size: 0.8em; padding: 5px 10px; background: linear-gradient(45deg, var(--primary), var(--secondary));">
                                    <i class="fas fa-times" style="margin-right: 5px;"></i>
                                    Clear Filter
                                </button>
                            </div>
                        </div>

                        <div id="smiles-result" class="result-box hidden"></div>

                        <!-- AI Model Selection - 移到这里 -->
                        <div style="margin-top: 20px;">
                            <h4 style="color: var(--accent); font-family: 'Orbitron', monospace; margin-bottom: 15px;">
                                <i class="fas fa-brain" style="margin-right: 10px;"></i>
                                AI Model Selection
                            </h4>
                            <div style="display: flex; gap: 30px; flex-wrap: wrap; align-items: flex-start;">
                                <!-- Model A Button with Checkmark -->
                                <div style="display: flex; flex-direction: column; align-items: center; gap: 8px;">
                                    <button id="model-a-btn" class="neon-button" onclick="selectModelButton('model-a-btn', 'modelA')" style="background: linear-gradient(45deg, var(--secondary), var(--accent)); transition: all 0.3s ease; border: 2px solid transparent;">
                                        <i class="fas fa-robot" style="margin-right: 10px;"></i>
                                        Model A (DECIMER)
                                    </button>
                                    <div id="model-a-checkmark" style="font-size: 24px; color: var(--success); display: none; animation: fadeIn 0.3s ease;">
                                        ✅
                                    </div>
                                </div>

                                <!-- Model B Button with Checkmark -->
                                <div style="display: flex; flex-direction: column; align-items: center; gap: 8px;">
                                    <button id="model-b-btn" class="neon-button" onclick="selectModelButton('model-b-btn', 'modelB')" style="background: linear-gradient(45deg, var(--secondary), var(--accent)); transition: all 0.3s ease; border: 2px solid transparent;">
                                        <i class="fas fa-cog" style="margin-right: 10px;"></i>
                                        Model B (AIChemist)
                                    </button>
                                    <div id="model-b-checkmark" style="font-size: 24px; color: var(--success); display: none; animation: fadeIn 0.3s ease;">
                                        ✅
                                    </div>
                                </div>

                                <!-- Model C Button with Checkmark -->
                                <div style="display: flex; flex-direction: column; align-items: center; gap: 8px;">
                                    <button id="model-c-btn" class="neon-button" onclick="selectModelButton('model-c-btn', 'modelC')" style="background: linear-gradient(45deg, var(--secondary), var(--accent)); transition: all 0.3s ease; border: 2px solid transparent;">
                                        <i class="fas fa-eye" style="margin-right: 10px;"></i>
                                        Model C (YOLO)
                                    </button>
                                    <div id="model-c-checkmark" style="font-size: 24px; color: var(--success); display: none; animation: fadeIn 0.3s ease;">
                                        ✅
                                    </div>
                                </div>
                            </div>
                            <div id="selected-model-display" style="margin-top: 15px; padding: 10px; background: rgba(0, 212, 255, 0.1); border: 1px solid var(--primary); border-radius: 8px; display: none; color: var(--primary); font-weight: bold;">
                                <i class="fas fa-check-circle" style="margin-right: 8px;"></i>
                                <span id="selected-model-text">No model selected</span>
                            </div>
                        </div>

                        <!-- 分割图片展示区域 - 在 AI Model Selection 下面 -->
                        <div id="segmented-images-result" class="result-box hidden" style="margin-top: 20px;"></div>

                        <!-- 处理结果展示 -->
                        <div id="molecule-result" class="result-box hidden"></div>
                    </div>
                </div>
            </div>






        </div>
    '''

    # JavaScript部分
    JAVASCRIPT_PART = '''
        <script>
            // 全局变量
            let selectedTemplate = null;
            let selectedWorkflow = null;
            let workflowRunning = false;
            let currentFileFilter = ''; // 当前选择的文件筛选

            // 页面加载时初始化
            document.addEventListener('DOMContentLoaded', function() {
                loadUserInfo();
                loadDatabaseStats();
                loadDocuments();
                loadDocumentOptions();
                loadFileFilterOptions(); // 加载文件筛选选项

                // 延迟加载模板，确保DOM完全加载
                setTimeout(() => {
                    loadPromptTemplates();
                    loadWorkflows();
                }, 500);
            });

            function showTab(tabName) {
                // 隐藏所有标签页内容
                const contents = document.querySelectorAll('.tab-content');
                contents.forEach(content => content.classList.add('hidden'));

                // 移除所有标签页的active类
                const tabs = document.querySelectorAll('.tab');
                tabs.forEach(tab => tab.classList.remove('active'));

                // 显示选中的标签页内容
                const targetTab = document.getElementById(tabName + '-tab');
                if (targetTab) {
                    targetTab.classList.remove('hidden');
                }

                // 添加active类到选中的标签页
                const activeTab = document.querySelector(`a[href="#${tabName}"]`);
                if (activeTab) {
                    activeTab.classList.add('active');
                }

                // 根据标签页刷新相应数据
                if (tabName === 'database') {
                    loadDatabaseStats();
                    loadDocuments();
                    loadFileFilterOptions(); // 刷新文件筛选选项
                } else if (tabName === 'extraction') {
                    loadDocumentOptions();
                    loadPromptTemplates();
                } else if (tabName === 'workflow') {
                    loadDocumentOptions();
                    loadWorkflows();
                } else if (tabName === 'molecule') {
                    loadMoleculeDocumentOptions();
                }
            }



            // 数据库相关函数

            // 文件筛选相关函数
            async function loadFileFilterOptions() {
                try {
                    const response = await fetch('/api/documents');
                    const documents = await response.json();

                    const select = document.getElementById('file-filter-select');
                    if (!select) return;

                    // 保存当前选择
                    const currentValue = select.value;

                    // 清空现有选项（保留第一个"所有文件"选项）
                    select.innerHTML = '<option value="">🗂️ 所有文件 (总体统计)</option>';

                    // 添加文档选项（只显示原始文档，不包括segment files）
                    const originalDocs = documents.filter(doc => !doc.is_segment_file);
                    originalDocs.forEach(doc => {
                        const option = document.createElement('option');
                        option.value = doc.filename;
                        option.textContent = `📄 ${doc.filename}`;
                        select.appendChild(option);
                    });

                    // 恢复之前的选择
                    if (currentValue && [...select.options].some(opt => opt.value === currentValue)) {
                        select.value = currentValue;
                        currentFileFilter = currentValue;
                    }

                } catch (error) {
                    console.error('Error loading file filter options:', error);
                }
            }

            function onFileFilterChange() {
                const select = document.getElementById('file-filter-select');
                const statusDiv = document.getElementById('filter-status');

                currentFileFilter = select.value;

                if (currentFileFilter) {
                    statusDiv.textContent = `当前显示: "${currentFileFilter}" 的统计数据和文档`;
                    statusDiv.style.color = 'var(--accent)';
                } else {
                    statusDiv.textContent = '当前显示: 所有文件的统计数据和文档';
                    statusDiv.style.color = 'var(--text-secondary)';
                }

                // 重新加载统计数据和文档列表
                loadDatabaseStats();
                loadDocuments();
            }

            function resetFileFilter() {
                const select = document.getElementById('file-filter-select');
                const statusDiv = document.getElementById('filter-status');

                select.value = '';
                currentFileFilter = '';
                statusDiv.textContent = '当前显示: 所有文件的统计数据和文档';
                statusDiv.style.color = 'var(--text-secondary)';

                // 重新加载统计数据和文档列表
                loadDatabaseStats();
                loadDocuments();
            }

            async function loadDatabaseStats() {
                try {
                    // 加载数据库连接信息
                    const infoResponse = await fetch('/api/database/info');
                    const infoData = await infoResponse.json();

                    const connectionInfo = document.getElementById('database-connection-info');
                    if (infoData.success) {
                        connectionInfo.innerHTML = `
                            <div style="display: flex; align-items: center; gap: 15px; padding: 15px; background: rgba(0, 99, 207, 0.1); border: 1px solid var(--success); border-radius: 10px;">
                                <i class="fas fa-database" style="color: var(--success); font-size: 1.5em;"></i>
                                <div>
                                    <div style="color: var(--success); font-weight: bold;">${infoData.database_type}</div>
                                    <div style="color: var(--text-secondary); font-size: 0.9em;">${infoData.connection_status}</div>
                                </div>
                            </div>
                        `;
                    }

                    // 加载统计信息（支持文件筛选）
                    let statsUrl = '/api/database/stats';
                    if (currentFileFilter) {
                        statsUrl += `?filename=${encodeURIComponent(currentFileFilter)}`;
                    }
                    const response = await fetch(statsUrl);
                    const stats = await response.json();

                    const statsContainer = document.getElementById('database-stats');
                    const segmentsCard = stats.total_segments !== undefined ? `
                        <div class="stat-card">
                            <i class="fas fa-puzzle-piece" style="color: var(--accent);"></i>
                            <div class="stat-value">${stats.total_segments}</div>
                            <div class="stat-label">Document Segments</div>
                        </div>
                    ` : '';

                    // 筛选状态显示
                    const filterIndicator = stats.filtered_file ? `
                        <div class="stat-card" style="background: linear-gradient(45deg, rgba(255, 193, 7, 0.1), rgba(255, 193, 7, 0.05)); border: 1px solid var(--warning);">
                            <i class="fas fa-filter" style="color: var(--warning);"></i>
                            <div class="stat-value" style="font-size: 1em; color: var(--warning);">筛选中</div>
                            <div class="stat-label" style="color: var(--warning);">文件筛选</div>
                        </div>
                    ` : '';

                    statsContainer.innerHTML = `
                        ${filterIndicator}
                        <div class="stat-card">
                            <i class="fas fa-file-alt" style="color: var(--primary);"></i>
                            <div class="stat-value">${stats.total_documents}</div>
                            <div class="stat-label">${stats.filtered_file ? 'Filtered Documents' : 'Total Documents'}</div>
                        </div>
                        <div class="stat-card">
                            <i class="fas fa-image" style="color: var(--secondary);"></i>
                            <div class="stat-value">${stats.total_chemical_images || 0}</div>
                            <div class="stat-label">Chemical Images</div>
                        </div>
                        ${segmentsCard}
                        <div class="stat-card">
                            <i class="fas fa-database" style="color: var(--success);"></i>
                            <div class="stat-value">${(stats.database_size / 1024).toFixed(1)}KB</div>
                            <div class="stat-label">${stats.filtered_file ? 'File Size' : 'Database Size'}</div>
                        </div>
                        <div class="stat-card">
                            <i class="fas fa-clock" style="color: var(--warning);"></i>
                            <div class="stat-value" style="font-size: 1.2em;">${stats.latest_upload}</div>
                            <div class="stat-label">${stats.filtered_file ? 'Selected File' : 'Latest Upload'}</div>
                        </div>
                    `;
                } catch (error) {
                    console.error('Error loading database stats:', error);
                }
            }

            async function loadDocuments() {
                try {
                    // 添加时间戳防止缓存
                    const timestamp = new Date().getTime();

                    // 构建API URL，支持文件筛选
                    let apiUrl = `/api/documents?_t=${timestamp}`;
                    if (currentFileFilter) {
                        apiUrl += `&filename=${encodeURIComponent(currentFileFilter)}`;
                    }

                    const response = await fetch(apiUrl);
                    const documents = await response.json();
                    console.log('Loaded documents:', documents);

                    const tbody = document.getElementById('documents-tbody');
                    tbody.innerHTML = documents.map(doc => {
                        // 显示所有文档，包括原始文档和segment files
                        if (!doc.is_segment_file) {
                            // 原始文档
                            return `
                                <tr>
                                    <td>${doc.id}</td>
                                    <td>
                                        <i class="fas fa-file-pdf" style="color: var(--accent); margin-right: 8px;"></i>
                                        ${doc.filename}
                                    </td>
                                    <td><span style="color: var(--accent);">${doc.file_type}</span></td>
                                    <td>${new Date(doc.upload_time).toLocaleString()}</td>
                                    <td><span style="color: var(--success);">${doc.processed_status}</span></td>
                                    <td>${(doc.file_size / 1024).toFixed(1)}KB</td>
                                    <td>${doc.page_count}</td>
                                    <td>
                                        <div class="action-buttons-container" style="display: flex; align-items: center; gap: 5px;">
                                            <button class="small-button" onclick="viewDocumentSegments('${doc.id}')"
                                                    style="background: linear-gradient(45deg, var(--accent), var(--primary)); padding: 5px 10px; font-size: 0.8em;">
                                                <i class="fas fa-puzzle-piece" style="margin-right: 5px;"></i>
                                                Segments
                                            </button>
                                            <button class="small-button" onclick="viewProcessingResults('${doc.id}')"
                                                    style="background: linear-gradient(45deg, var(--success), var(--secondary)); padding: 5px 10px; font-size: 0.8em;">
                                                <i class="fas fa-file-code" style="margin-right: 5px;"></i>
                                                Results
                                            </button>

                                            <!-- 删除按钮 -->
                                            <button class="small-button delete-btn" onclick="deleteDocument('${doc.id}', '${doc.filename}', false)"
                                                    title="删除文档">
                                                <i class="fas fa-trash-alt"></i>
                                            </button>
                                        </div>
                                    </td>
                                </tr>
                            `;
                        } else {
                            // Segment files
                            const getFileIcon = (fileType) => {
                                if (fileType.includes('Markdown')) return 'fas fa-file-alt';
                                if (fileType.includes('JSON')) return 'fas fa-file-code';
                                if (fileType.includes('PDF')) return 'fas fa-file-pdf';
                                return 'fas fa-file';
                            };

                            return `
                                <tr style="background: rgba(123, 104, 238, 0.05);">
                                    <td>${doc.id}</td>
                                    <td>
                                        <i class="${getFileIcon(doc.file_type)}" style="color: var(--secondary); margin-right: 8px;"></i>
                                        ${doc.filename}
                                    </td>
                                    <td><span style="color: var(--secondary);">${doc.file_type}</span></td>
                                    <td>${new Date(doc.upload_time).toLocaleString()}</td>
                                    <td><span style="color: var(--info);">Generated</span></td>
                                    <td>Generated</td>
                                    <td>N/A</td>
                                    <td>
                                        <div class="action-buttons-container" style="display: flex; align-items: center; gap: 5px;">
                                            <button class="small-button" onclick="downloadSegmentFile('${doc.parent_document_id}', '${doc.segment_file_type}')"
                                                    style="background: linear-gradient(45deg, var(--info), var(--secondary)); padding: 5px 10px; font-size: 0.8em;">
                                                <i class="fas fa-download" style="margin-right: 5px;"></i>
                                                Download
                                            </button>

                                            <!-- 删除按钮 -->
                                            <button class="small-button delete-btn" onclick="deleteDocument('${doc.id}', '${doc.filename}', true)"
                                                    title="删除文档">
                                                <i class="fas fa-trash-alt"></i>
                                            </button>
                                        </div>
                                    </td>
                                </tr>
                            `;
                        }
                    }).join('');
                } catch (error) {
                    console.error('Error loading documents:', error);
                }
            }

            async function viewDocumentSegments(documentId) {
                try {
                    const response = await fetch(`/api/documents/${documentId}/segments`);
                    const data = await response.json();

                    const segmentsSection = document.getElementById('document-segments-section');
                    const segmentsContainer = document.getElementById('document-segments-container');

                    if (data.success && data.segments.length > 0) {
                        segmentsContainer.innerHTML = `
                            <div style="margin-bottom: 15px; padding: 10px; background: rgba(123, 104, 238, 0.1); border-radius: 8px;">
                                <strong>Document ID:</strong> ${data.document_id} |
                                <strong>Total Segments:</strong> ${data.total_segments}
                            </div>
                            ${data.segments.map((segment, index) => `
                                <div class="segment-card" style="margin-bottom: 15px; padding: 15px; background: rgba(255, 255, 255, 0.05); border-radius: 10px; border-left: 4px solid var(--accent);">
                                    <div style="display: flex; justify-content: between; align-items: center; margin-bottom: 10px;">
                                        <h5 style="color: var(--accent); margin: 0;">
                                            Segment ${segment.segment_index + 1} - ${segment.segment_type.toUpperCase()}
                                        </h5>
                                        <span style="color: var(--text-secondary); font-size: 0.9em;">
                                            Page ${segment.page_number} | Confidence: ${(segment.confidence * 100).toFixed(1)}%
                                        </span>
                                    </div>
                                    <div style="background: rgba(0, 0, 0, 0.3); padding: 10px; border-radius: 5px; font-family: monospace; font-size: 0.9em; max-height: 200px; overflow-y: auto;">
                                        ${segment.content || 'No content available'}
                                    </div>
                                    ${segment.bbox && segment.bbox.length > 0 ? `
                                        <div style="margin-top: 10px; font-size: 0.8em; color: var(--text-secondary);">
                                            BBox: [${segment.bbox.join(', ')}]
                                        </div>
                                    ` : ''}
                                </div>
                            `).join('')}
                        `;
                        segmentsSection.classList.remove('hidden');
                    } else {
                        segmentsContainer.innerHTML = `
                            <div style="text-align: center; padding: 20px; color: var(--text-secondary);">
                                <i class="fas fa-info-circle" style="font-size: 2em; margin-bottom: 10px;"></i>
                                <p>No segments found for this document.</p>
                            </div>
                        `;
                        segmentsSection.classList.remove('hidden');
                    }

                    // 滚动到片段部分
                    segmentsSection.scrollIntoView({ behavior: 'smooth' });

                } catch (error) {
                    console.error('Error loading document segments:', error);
                    alert('Failed to load document segments');
                }
            }

            async function viewProcessingResults(documentId) {
                try {
                    const response = await fetch(`/api/documents/${documentId}/processing-results`);
                    const data = await response.json();

                    const resultsSection = document.getElementById('processing-results-section');
                    const resultsContainer = document.getElementById('processing-results-container');

                    if (data.success && data.results) {
                        const results = data.results;

                        resultsContainer.innerHTML = `
                            <div style="margin-bottom: 15px; padding: 10px; background: rgba(0, 99, 207, 0.1); border-radius: 8px;">
                                <strong>Document:</strong> ${results.filename} |
                                <strong>Created:</strong> ${new Date(results.created_time).toLocaleString()}
                            </div>

                            <!-- 处理结果文件 -->
                            <div class="results-grid" style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin-bottom: 20px;">

                                <!-- Markdown 结果 -->
                                <div class="result-card" style="padding: 15px; background: rgba(255, 255, 255, 0.05); border-radius: 10px; border-left: 4px solid var(--success);">
                                    <div style="display: flex; justify-content: between; align-items: center; margin-bottom: 10px;">
                                        <h5 style="color: var(--success); margin: 0;">
                                            <i class="fab fa-markdown" style="margin-right: 8px;"></i>
                                            Markdown Content
                                        </h5>
                                        <button class="small-button" onclick="downloadResult('${documentId}', 'markdown')"
                                                style="background: var(--success); padding: 3px 8px; font-size: 0.7em;">
                                            <i class="fas fa-download"></i>
                                        </button>
                                    </div>
                                    <div style="background: rgba(0, 0, 0, 0.3); padding: 10px; border-radius: 5px; font-family: monospace; font-size: 0.8em; max-height: 200px; overflow-y: auto;">
                                        ${results.markdown_content ? results.markdown_content.substring(0, 500) + (results.markdown_content.length > 500 ? '...' : '') : 'No markdown content available'}
                                    </div>
                                </div>

                                <!-- Content List JSON -->
                                <div class="result-card" style="padding: 15px; background: rgba(255, 255, 255, 0.05); border-radius: 10px; border-left: 4px solid var(--primary);">
                                    <div style="display: flex; justify-content: between; align-items: center; margin-bottom: 10px;">
                                        <h5 style="color: var(--primary); margin: 0;">
                                            <i class="fas fa-list" style="margin-right: 8px;"></i>
                                            Content List JSON
                                        </h5>
                                        <button class="small-button" onclick="downloadResult('${documentId}', 'content_list')"
                                                style="background: var(--primary); padding: 3px 8px; font-size: 0.7em;">
                                            <i class="fas fa-download"></i>
                                        </button>
                                    </div>
                                    <div style="background: rgba(0, 0, 0, 0.3); padding: 10px; border-radius: 5px; font-family: monospace; font-size: 0.8em; max-height: 200px; overflow-y: auto;">
                                        ${results.content_list ? JSON.stringify(results.content_list, null, 2).substring(0, 500) + '...' : 'No content list available'}
                                    </div>
                                </div>

                                <!-- Middle JSON -->
                                <div class="result-card" style="padding: 15px; background: rgba(255, 255, 255, 0.05); border-radius: 10px; border-left: 4px solid var(--accent);">
                                    <div style="display: flex; justify-content: between; align-items: center; margin-bottom: 10px;">
                                        <h5 style="color: var(--accent); margin: 0;">
                                            <i class="fas fa-cogs" style="margin-right: 8px;"></i>
                                            Middle JSON
                                        </h5>
                                        <button class="small-button" onclick="downloadResult('${documentId}', 'middle_json')"
                                                style="background: var(--accent); padding: 3px 8px; font-size: 0.7em;">
                                            <i class="fas fa-download"></i>
                                        </button>
                                    </div>
                                    <div style="background: rgba(0, 0, 0, 0.3); padding: 10px; border-radius: 5px; font-family: monospace; font-size: 0.8em; max-height: 200px; overflow-y: auto;">
                                        ${results.middle_json ? JSON.stringify(results.middle_json, null, 2).substring(0, 500) + '...' : 'No middle JSON available'}
                                    </div>
                                </div>
                            </div>

                            <!-- 统计信息 -->
                            ${results.middle_json && results.middle_json.statistics ? `
                                <div style="margin-top: 20px; padding: 15px; background: rgba(123, 104, 238, 0.1); border-radius: 8px;">
                                    <h6 style="color: var(--accent); margin-bottom: 10px;">Processing Statistics</h6>
                                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 10px;">
                                        <div><strong>Total Characters:</strong> ${results.middle_json.statistics.total_characters}</div>
                                        <div><strong>Avg Confidence:</strong> ${(results.middle_json.statistics.avg_confidence * 100).toFixed(1)}%</div>
                                        <div><strong>Pages:</strong> ${Object.keys(results.middle_json.statistics.page_distribution || {}).length}</div>
                                    </div>
                                </div>
                            ` : ''}

                            <!-- 片段文件 -->
                            <div id="segment-files-${documentId}" style="margin-top: 20px;">
                                <h6 style="color: var(--secondary); margin-bottom: 10px;">
                                    <i class="fas fa-file-alt" style="margin-right: 8px;"></i>
                                    Generated Segment Files
                                </h6>
                                <div class="loading-segment-files">Loading segment files...</div>
                            </div>
                        `;
                        resultsSection.classList.remove('hidden');
                    } else {
                        resultsContainer.innerHTML = `
                            <div style="text-align: center; padding: 20px; color: var(--text-secondary);">
                                <i class="fas fa-info-circle" style="font-size: 2em; margin-bottom: 10px;"></i>
                                <p>No processing results found for this document.</p>
                            </div>
                        `;
                        resultsSection.classList.remove('hidden');
                    }

                    // 滚动到结果部分
                    resultsSection.scrollIntoView({ behavior: 'smooth' });

                    // 加载片段文件
                    loadSegmentFiles(documentId);

                } catch (error) {
                    console.error('Error loading processing results:', error);
                    alert('Failed to load processing results');
                }
            }

            async function loadSegmentFiles(documentId) {
                try {
                    const response = await fetch(`/api/documents/${documentId}/segment-files`);
                    const data = await response.json();

                    const segmentFilesContainer = document.getElementById(`segment-files-${documentId}`);
                    const loadingElement = segmentFilesContainer.querySelector('.loading-segment-files');

                    if (data.success && data.segment_files && data.segment_files.length > 0) {
                        const segmentFiles = data.segment_files;

                        let filesHtml = `
                            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 15px;">
                        `;

                        segmentFiles.forEach(file => {
                            const iconClass = getSegmentFileIcon(file.file_type);
                            const colorClass = getSegmentFileColor(file.file_type);

                            filesHtml += `
                                <div class="segment-file-card" style="padding: 12px; background: rgba(255, 255, 255, 0.05); border-radius: 8px; border-left: 4px solid ${colorClass};">
                                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">
                                        <div style="display: flex; align-items: center;">
                                            <i class="${iconClass}" style="margin-right: 8px; color: ${colorClass};"></i>
                                            <span style="font-size: 0.9em; font-weight: bold;">${file.virtual_filename}</span>
                                        </div>
                                        <button class="small-button" onclick="downloadSegmentFile('${documentId}', '${file.file_type}')"
                                                style="background: ${colorClass}; padding: 2px 6px; font-size: 0.7em;">
                                            <i class="fas fa-download"></i>
                                        </button>
                                    </div>
                                    <div style="font-size: 0.8em; color: var(--text-secondary);">
                                        Type: ${file.file_type.replace('_', ' ').toUpperCase()}
                                    </div>
                                    <div style="font-size: 0.7em; color: var(--text-secondary); margin-top: 4px;">
                                        Created: ${new Date(file.created_time).toLocaleString()}
                                    </div>
                                </div>
                            `;
                        });

                        filesHtml += '</div>';
                        loadingElement.innerHTML = filesHtml;

                    } else {
                        loadingElement.innerHTML = `
                            <div style="text-align: center; padding: 15px; color: var(--text-secondary); font-style: italic;">
                                No segment files generated for this document.
                            </div>
                        `;
                    }

                } catch (error) {
                    console.error('Error loading segment files:', error);
                    const segmentFilesContainer = document.getElementById(`segment-files-${documentId}`);
                    const loadingElement = segmentFilesContainer.querySelector('.loading-segment-files');
                    loadingElement.innerHTML = `
                        <div style="text-align: center; padding: 15px; color: var(--error);">
                            Failed to load segment files.
                        </div>
                    `;
                }
            }

            function getFileIcon(fileType) {
                switch(fileType) {
                    case 'layout_pdf': return 'fas fa-th-large';
                    case 'model_pdf': return 'fas fa-sitemap';
                    case 'spans_pdf': return 'fas fa-text-width';
                    default: return 'fas fa-file';
                }
            }

            function getFileColor(fileType) {
                switch(fileType) {
                    case 'layout_pdf': return 'var(--primary)';
                    case 'model_pdf': return 'var(--accent)';
                    case 'spans_pdf': return 'var(--success)';
                    default: return 'var(--secondary)';
                }
            }

            function getSegmentFileIcon(fileType) {
                switch(fileType) {
                    case 'layout_pdf': return 'fas fa-file-pdf';
                    case 'model_pdf': return 'fas fa-file-pdf';
                    case 'spans_pdf': return 'fas fa-file-pdf';
                    case 'markdown': return 'fab fa-markdown';
                    case 'content_list': return 'fas fa-list';
                    case 'middle_json': return 'fas fa-cogs';
                    default: return 'fas fa-file';
                }
            }

            function getSegmentFileColor(fileType) {
                switch(fileType) {
                    case 'layout_pdf': return '#dc3545';
                    case 'model_pdf': return '#dc3545';
                    case 'spans_pdf': return '#dc3545';
                    case 'markdown': return '#0366d6';
                    case 'content_list_json': return '#f39c12';
                    case 'middle_json': return '#9b59b6';
                    case 'content_list': return '#f39c12';
                    case 'middle': return '#9b59b6';
                    default: return 'var(--secondary)';
                }
            }

            function getSegmentFileIcon(fileType) {
                switch(fileType) {
                    case 'layout_pdf': return 'fas fa-file-pdf';
                    case 'model_pdf': return 'fas fa-file-pdf';
                    case 'spans_pdf': return 'fas fa-file-pdf';
                    case 'markdown': return 'fab fa-markdown';
                    case 'content_list_json': return 'fas fa-file-code';
                    case 'middle_json': return 'fas fa-file-code';
                    case 'content_list': return 'fas fa-file-code';
                    case 'middle': return 'fas fa-file-code';
                    default: return 'fas fa-file';
                }
            }

            function downloadSegmentFile(documentId, fileType) {
                const url = `/api/documents/${documentId}/segment-files/${fileType}/download`;
                const link = document.createElement('a');
                link.href = url;
                link.download = '';
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
            }

            function downloadGeneratedFile(documentId, fileType) {
                // 根据文件类型选择正确的下载URL
                let url;
                if (['markdown', 'content_list', 'middle_json'].includes(fileType)) {
                    // 处理结果文件
                    url = `/api/documents/${documentId}/download/${fileType}`;
                } else {
                    // 片段文件（包括新的JSON格式和旧的PDF格式）
                    url = `/api/documents/${documentId}/segment-files/${fileType}/download`;
                }

                const link = document.createElement('a');
                link.href = url;
                link.download = '';
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
            }

            function downloadSegmentFile(documentId, fileType) {
                // 统一使用segment-files下载接口
                const url = `/api/documents/${documentId}/segment-files/${fileType}/download`;
                const link = document.createElement('a');
                link.href = url;
                link.download = '';
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
            }

            async function useMagicPdfContent() {
                if (!confirm('确定要批量更新所有文档使用Magic PDF生成的完整内容吗？这将替换现有的简化内容。')) {
                    return;
                }

                try {
                    // 显示进度提示
                    const progressDiv = document.createElement('div');
                    progressDiv.innerHTML = `
                        <div style="position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%);
                                    background: rgba(255,255,255,0.95); color: black; padding: 20px; border-radius: 10px; z-index: 10000; border: 1px solid #ccc;">
                            <i class="fas fa-spinner fa-spin" style="margin-right: 10px;"></i>
                            正在更新文档内容...
                        </div>
                    `;
                    document.body.appendChild(progressDiv);

                    const response = await fetch('/api/documents/batch-use-magic-pdf', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json'
                        }
                    });

                    const result = await response.json();
                    document.body.removeChild(progressDiv);

                    if (result.success) {
                        let message = `更新完成！\\n`;
                        message += `- 成功更新: ${result.updated_count} 个文档\\n`;
                        message += `- 失败: ${result.failed_count} 个文档\\n`;
                        message += `- 总文档数: ${result.total_documents} 个\\n\\n`;

                        if (result.results && result.results.length > 0) {
                            message += `详细结果:\\n`;
                            result.results.forEach(r => {
                                if (r.status === 'updated') {
                                    message += `✅ ${r.filename}: 已更新 (${r.content_length} 字符)\\n`;
                                } else if (r.status === 'failed') {
                                    message += `❌ ${r.filename}: 失败 - ${r.reason}\\n`;
                                } else {
                                    message += `⏭️ ${r.filename}: 跳过 - ${r.reason}\\n`;
                                }
                            });
                        }

                        alert(message);
                        loadDatabaseStats();
                        loadDocuments();
                    } else {
                        alert('更新失败: ' + result.message);
                    }
                } catch (error) {
                    console.error('Error updating with Magic PDF content:', error);
                    alert('更新Magic PDF内容时发生错误: ' + error.message);
                }
            }

            async function clearDatabase() {
                if (!confirm('确定要清空所有数据库数据吗？此操作不可恢复！')) {
                    return;
                }

                try {
                    const response = await fetch('/api/database/clear', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json'
                        }
                    });

                    const result = await response.json();

                    if (result.success) {
                        alert('数据库已成功清空！');
                        // 刷新页面数据
                        loadDocuments();
                        loadDatabaseStats();
                        loadDatabaseInfo();
                    } else {
                        alert('清空数据库失败: ' + result.error);
                    }
                } catch (error) {
                    console.error('Error clearing database:', error);
                    alert('清空数据库时发生错误: ' + error.message);
                }
            }

            function downloadResult(documentId, fileType) {
                const url = `/api/documents/${documentId}/download/${fileType}`;
                const link = document.createElement('a');
                link.href = url;
                link.download = '';
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
            }

            async function loadDocumentOptions() {
                try {
                    const response = await fetch('/api/documents');
                    const documents = await response.json();

                    // 更新信息提取页面的多选框
                    const multiSelect = document.getElementById('document-multi-select');
                    if (multiSelect) {
                        multiSelect.innerHTML = documents.map(doc => `
                            <option value="${doc.id}">${doc.filename} (${doc.file_type}) - ${new Date(doc.upload_time).toLocaleDateString()}</option>
                        `).join('');
                    }

                    // 更新工作流页面的多选框
                    const workflowSelect = document.getElementById('workflow-document-select');
                    if (workflowSelect) {
                        workflowSelect.innerHTML = documents.map(doc => `
                            <option value="${doc.id}">${doc.filename} (${doc.file_type}) - ${new Date(doc.upload_time).toLocaleDateString()}</option>
                        `).join('');
                    }
                } catch (error) {
                    console.error('Error loading document options:', error);
                }
            }

            async function loadPromptTemplates() {
                try {
                    const response = await fetch('/api/templates');
                    const templates = await response.json();

                    // 检查元素是否存在
                    const templateSelect = document.getElementById('template-select');
                    if (!templateSelect) {
                        console.warn('⚠️ template-select element not found');
                        return;
                    }

                    templateSelect.innerHTML = '<option value="">-- 请选择提示模板 --</option>' +
                        Object.entries(templates).map(([id, template]) => `
                            <option value="${id}">${template.name} - ${template.description}</option>
                        `).join('');
                    
                    console.log('✅ Templates loaded successfully');
                } catch (error) {
                    console.error('Error loading templates:', error);
                }
            }

            async function loadWorkflows() {
                try {
                    const response = await fetch('/api/workflows');
                    const workflows = await response.json();

                    const workflowGrid = document.getElementById('workflow-grid');
                    workflowGrid.innerHTML = Object.entries(workflows).map(([id, workflow]) => `
                        <div class="workflow-card" onclick="selectWorkflow('${id}')">
                            <h4>${workflow.name}</h4>
                            <p>${workflow.description}</p>
                            <div class="workflow-info">
                                <span class="workflow-time">
                                    <i class="fas fa-clock" style="margin-right: 5px;"></i>
                                    ${workflow.estimated_time}
                                </span>
                                <span class="workflow-steps">
                                    <i class="fas fa-list" style="margin-right: 5px;"></i>
                                    ${workflow.steps.length} steps
                                </span>
                            </div>
                        </div>
                    `).join('');
                } catch (error) {
                    console.error('Error loading workflows:', error);
                }
            }

            function selectTemplate(templateId) {
                // 移除之前的选择
                document.querySelectorAll('.template-card').forEach(card => {
                    card.classList.remove('selected');
                });

                // 选择当前模板
                event.target.closest('.template-card').classList.add('selected');
                selectedTemplate = templateId;

                // 更新查询输入框
                fetch(`/api/templates/${templateId}`)
                    .then(response => response.json())
                    .then(template => {
                        const queryInput = document.getElementById('query-input');
                        queryInput.placeholder = `Template: ${template.name} - ${template.description}`;
                    });
            }



            // 启动自动工作流




            function clearSelection() {
                // 清除模板选择
                document.querySelectorAll('.template-card').forEach(card => {
                    card.classList.remove('selected');
                });
                selectedTemplate = null;

                // 清除文档选择
                const multiSelect = document.getElementById('document-multi-select');
                multiSelect.selectedIndex = -1;

                // 清除查询输入
                document.getElementById('query-input').value = '';
                document.getElementById('query-input').placeholder = 'Enter your custom query or modify the selected template...';
            }

            function refreshDatabase() {
                console.log('🔄 开始刷新数据库...');

                // 显示刷新状态
                const refreshBtn = document.querySelector('button[onclick="refreshDatabase()"]');
                if (refreshBtn) {
                    const originalText = refreshBtn.innerHTML;
                    refreshBtn.innerHTML = '<i class="fas fa-spinner fa-spin" style="margin-right: 10px;"></i>Refreshing...';
                    refreshBtn.disabled = true;

                    // 执行刷新操作
                    Promise.all([
                        loadDatabaseStats(),
                        loadDocuments(),
                        loadDocumentOptions()
                    ]).then(() => {
                        console.log('✅ 数据库刷新完成');

                        // 显示成功提示
                        showNotification('数据库已刷新', 'success');

                        // 恢复按钮状态
                        refreshBtn.innerHTML = originalText;
                        refreshBtn.disabled = false;
                    }).catch(error => {
                        console.error('❌ 数据库刷新失败:', error);

                        // 显示错误提示
                        showNotification('数据库刷新失败: ' + error.message, 'error');

                        // 恢复按钮状态
                        refreshBtn.innerHTML = originalText;
                        refreshBtn.disabled = false;
                    });
                }
            }

            function showNotification(message, type = 'info') {
                // 创建通知元素
                const notification = document.createElement('div');
                notification.className = `notification notification-${type}`;
                notification.style.cssText = `
                    position: fixed;
                    top: 20px;
                    right: 20px;
                    padding: 15px 20px;
                    border-radius: 10px;
                    color: white;
                    font-weight: bold;
                    z-index: 10000;
                    max-width: 300px;
                    box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
                    transform: translateX(100%);
                    transition: transform 0.3s ease;
                `;

                // 根据类型设置颜色
                if (type === 'success') {
                    notification.style.background = 'linear-gradient(45deg, var(--success), #004A9F)';
                } else if (type === 'error') {
                    notification.style.background = 'linear-gradient(45deg, var(--danger), #dc2626)';
                } else {
                    notification.style.background = 'linear-gradient(45deg, var(--primary), var(--accent))';
                }

                notification.innerHTML = `
                    <i class="fas fa-${type === 'success' ? 'check-circle' : type === 'error' ? 'exclamation-circle' : 'info-circle'}" style="margin-right: 10px;"></i>
                    ${message}
                `;

                document.body.appendChild(notification);

                // 显示动画
                setTimeout(() => {
                    notification.style.transform = 'translateX(0)';
                }, 100);

                // 自动隐藏
                setTimeout(() => {
                    notification.style.transform = 'translateX(100%)';
                    setTimeout(() => {
                        if (notification.parentNode) {
                            document.body.removeChild(notification);
                        }
                    }, 300);
                }, 3000);
            }

            async function deleteDocument(documentId, filename, isSegmentFile = false) {
                // 简化确认消息
                const confirmMessage = `确定要删除文档 "${filename}" 吗？\n\n此操作不可撤销！`;

                if (!confirm(confirmMessage)) {
                    return;
                }

                try {
                    // 显示删除进度
                    const progressDiv = document.createElement('div');
                    progressDiv.innerHTML = `
                        <div style="position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%);
                                    background: rgba(255, 255, 255, 0.95); padding: 20px; border-radius: 10px; border: 1px solid #ccc;
                                    color: black; z-index: 10000; text-align: center;">
                            <i class="fas fa-spinner fa-spin" style="font-size: 2em; margin-bottom: 10px; color: var(--danger);"></i>
                            <p>正在删除文档: ${filename}</p>
                            <p style="font-size: 0.9em; color: var(--text-secondary);">请稍候...</p>
                        </div>
                    `;
                    document.body.appendChild(progressDiv);

                    const response = await fetch(`/api/documents/${documentId}`, {
                        method: 'DELETE',
                        headers: {
                            'Content-Type': 'application/json'
                        }
                    });

                    const result = await response.json();

                    // 移除进度提示
                    document.body.removeChild(progressDiv);

                    if (result.success) {
                        // 显示成功消息
                        alert(`✅ 文档 "${filename}" 删除成功！`);

                        // 刷新页面数据
                        loadDatabaseStats();
                        loadDocuments();
                        loadDocumentOptions();
                        loadFileFilterOptions(); // 更新筛选下拉框
                    } else {
                        alert(`❌ 删除文档失败: ${result.error}`);
                    }

                } catch (error) {
                    // 移除进度提示（如果还存在）
                    const progressDiv = document.querySelector('div[style*="position: fixed"]');
                    if (progressDiv) {
                        document.body.removeChild(progressDiv);
                    }

                    console.error('Error deleting document:', error);
                    alert(`❌ 删除文档时发生错误: ${error.message}`);
                }
            }

            // 文档处理相关函数
            async function processDocument() {
                const fileInput = document.getElementById('file-upload');
                const files = fileInput.files;

                if (files.length === 0) {
                    alert('Please select at least one file');
                    return;
                }

                // 检查文件大小
                let totalSize = 0;
                let largeFiles = [];
                for (let file of files) {
                    totalSize += file.size;
                    if (file.size > 500 * 1024 * 1024) { // 500MB
                        largeFiles.push(file.name);
                    }
                }

                // 显示大文件警告
                if (largeFiles.length > 0) {
                    const proceed = confirm(`检测到大文件 (${largeFiles.join(', ')})。\n处理可能需要较长时间，是否继续？`);
                    if (!proceed) return;
                }

                // 检查是否有PDF文件
                const hasPDF = Array.from(files).some(file => file.name.toLowerCase().endsWith('.pdf'));

                // 显示处理进度
                const progressDiv = document.getElementById('processing-progress');
                const progressFill = document.getElementById('progress-fill');
                const progressText = document.getElementById('progress-text');
                const resultDiv = document.getElementById('document-result');

                progressDiv.classList.remove('hidden');
                resultDiv.classList.add('hidden');

                // 显示轻量级模式状态和文件信息
                progressText.innerHTML = `
                    <div style="margin-bottom: 10px;">
                        <i class="fas fa-bolt" style="color: var(--accent); margin-right: 5px;"></i>
                        轻量级模式: 启用 | 文件总数: ${files.length} | 总大小: ${(totalSize / 1024 / 1024).toFixed(1)}MB
                    </div>
                    <div>正在上传文件...</div>
                `;

                try {
                    const formData = new FormData();

                    // 添加文件
                    for (let i = 0; i < files.length; i++) {
                        formData.append('files', files[i]);
                    }

                    // 如果有PDF文件，添加PDF处理选项
                    if (hasPDF) {
                        const enableOCR = document.getElementById('enable-ocr').checked;
                        const enableFormula = document.getElementById('enable-formula').checked;
                        const enableTable = document.getElementById('enable-table').checked;

                        formData.append('enable_ocr', enableOCR);
                        formData.append('enable_formula', enableFormula);
                        formData.append('enable_table', enableTable);
                    }

                    progressFill.style.width = '20%';

                    // 启动处理请求
                    const response = await fetch('/api/documents/process', {
                        method: 'POST',
                        body: formData
                    });

                    progressText.innerHTML = `
                        <div style="margin-bottom: 10px;">
                            <i class="fas fa-cogs" style="color: var(--warning); margin-right: 5px;"></i>
                            轻量级处理中... | 大文件优化: ${largeFiles.length > 0 ? '启用' : '标准'}
                        </div>
                        <div>正在快速分析文档结构...</div>
                    `;
                    progressFill.style.width = '60%';

                    const data = await response.json();

                    progressText.innerHTML = `
                        <div style="margin-bottom: 10px;">
                            <i class="fas fa-check" style="color: var(--success); margin-right: 5px;"></i>
                            轻量级处理完成 | 速度优化: 已启用
                        </div>
                        <div>正在保存结果...</div>
                    `;
                    progressFill.style.width = '100%';

                    setTimeout(() => {
                        progressDiv.classList.add('hidden');

                        if (data.success) {
                            resultDiv.innerHTML = `
                                <h4 style="color: var(--success); margin-bottom: 15px;">
                                    <i class="fas fa-check-circle" style="margin-right: 10px;"></i>
                                    Documents Processed Successfully
                                </h4>
                                <div style="margin-bottom: 15px;">
                                    <strong>Processed Files:</strong> ${data.processed_count || files.length}
                                </div>
                                ${data.total_chemical_structures > 0 ? `
                                    <div style="margin-bottom: 15px; padding: 10px; background: rgba(123, 104, 238, 0.1); border-radius: 8px;">
                                        <strong style="color: var(--accent);">
                                            <i class="fas fa-image" style="margin-right: 8px;"></i>
                                            Chemical Structures Extracted: ${data.total_chemical_structures}
                                        </strong>
                                    </div>
                                ` : ''}
                                ${data.pdf_results ? `
                                    <div style="margin-bottom: 15px;">
                                        <strong>PDF Processing Results:</strong>
                                        <ul style="margin-top: 10px; padding-left: 20px;">
                                            ${data.pdf_results.map(result => `
                                                <li style="margin-bottom: 5px;">
                                                    ${result.filename}: ${result.segments || 0} segments extracted
                                                    ${result.chemical_structures > 0 ? `, <span style="color: var(--accent);">${result.chemical_structures} chemical structures</span>` : ''}
                                                </li>
                                            `).join('')}
                                        </ul>
                                    </div>
                                ` : ''}
                                <div style="margin-top: 15px;">
                                    <button class="neon-button" onclick="refreshDatabase()" style="background: linear-gradient(45deg, var(--success), var(--primary));">
                                        <i class="fas fa-sync" style="margin-right: 10px;"></i>
                                        Refresh Database View
                                    </button>
                                </div>
                            `;
                        } else {
                            resultDiv.innerHTML = `
                                <h4 style="color: var(--danger); margin-bottom: 15px;">
                                    <i class="fas fa-exclamation-triangle" style="margin-right: 10px;"></i>
                                    Processing Failed
                                </h4>
                                <p style="color: var(--text-secondary);">${data.message || 'Unknown error occurred'}</p>
                            `;
                        }

                        resultDiv.classList.remove('hidden');

                        // 清空文件输入
                        fileInput.value = '';
                        document.getElementById('pdf-options').classList.add('hidden');

                    }, 1000);

                } catch (error) {
                    progressDiv.classList.add('hidden');

                    let errorTitle = "Processing Error";
                    let errorMessage = error.message;
                    let errorIcon = "fas fa-exclamation-triangle";

                    // 检查是否是文件大小错误
                    if (error.message.includes('413') || error.message.includes('too large')) {
                        errorTitle = "File Too Large";
                        errorMessage = "文件太大，请选择小于500MB的文件，或者尝试压缩PDF文件。";
                        errorIcon = "fas fa-file-alt";
                    } else if (error.message.includes('timeout')) {
                        errorTitle = "Processing Timeout";
                        errorMessage = "处理超时，请尝试处理较小的文件或稍后重试。";
                        errorIcon = "fas fa-clock";
                    }

                    resultDiv.innerHTML = `
                        <h4 style="color: var(--danger); margin-bottom: 15px;">
                            <i class="${errorIcon}" style="margin-right: 10px;"></i>
                            ${errorTitle}
                        </h4>
                        <p style="color: var(--text-secondary); margin-bottom: 15px;">${errorMessage}</p>
                        <div style="padding: 10px; background: rgba(255, 107, 107, 0.1); border-radius: 8px; border-left: 4px solid var(--danger);">
                            <strong>建议:</strong>
                            <ul style="margin: 10px 0 0 20px; color: var(--text-secondary);">
                                <li>确保PDF文件小于500MB</li>
                                <li>尝试压缩PDF文件</li>
                                <li>分批处理多个文件</li>
                                <li>检查网络连接稳定性</li>
                            </ul>
                        </div>
                    `;
                    resultDiv.classList.remove('hidden');
                }
            }

            // 文件选择变化时显示/隐藏PDF选项
            document.addEventListener('DOMContentLoaded', function() {
                const fileInput = document.getElementById('file-upload');
                const pdfOptions = document.getElementById('pdf-options');

                if (fileInput && pdfOptions) {
                    fileInput.addEventListener('change', function() {
                        const files = this.files;
                        const hasPDF = Array.from(files).some(file => file.name.toLowerCase().endsWith('.pdf'));

                        if (hasPDF) {
                            pdfOptions.classList.remove('hidden');
                        } else {
                            pdfOptions.classList.add('hidden');
                        }
                    });
                }
            });

            // 分子处理相关函数
            async function loadMoleculeDocumentOptions() {
                try {
                    const response = await fetch('/api/documents');
                    const documents = await response.json();

                    const select = document.getElementById('molecule-document-select');
                    if (select) {
                        select.innerHTML = '<option value="">Choose a document...</option>' +
                            documents.map(doc => `
                                <option value="${doc.id}">${doc.filename} (${doc.file_type}) - ${new Date(doc.upload_time).toLocaleDateString()}</option>
                            `).join('');
                    }
                } catch (error) {
                    console.error('Error loading molecule document options:', error);
                }
            }

            // 当选择文档时提取图片
            document.addEventListener('DOMContentLoaded', function() {
                const moleculeSelect = document.getElementById('molecule-document-select');
                if (moleculeSelect) {
                    moleculeSelect.addEventListener('change', function() {
                        if (this.value) {
                            extractImagesFromDocument(this.value);
                        } else {
                            document.getElementById('extracted-images').classList.add('hidden');
                        }
                    });
                }
            });

            async function extractImagesFromDocument(documentId) {
                try {
                    const response = await fetch(`/api/documents/${documentId}/images`);
                    const data = await response.json();

                    const imagesContainer = document.getElementById('extracted-images');
                    const imagesGrid = document.getElementById('images-grid');

                    if (data.images && data.images.length > 0) {
                        imagesGrid.innerHTML = data.images.map((img, index) => `
                            <div class="image-card" style="border: 1px solid var(--border); border-radius: 10px; padding: 15px; background: var(--glass); text-align: center;">
                                <img src="${img.url}" alt="Molecular Structure ${index + 1}"
                                     style="max-width: 100%; height: 150px; object-fit: contain; border-radius: 5px; background: white;">
                                <p style="margin-top: 10px; color: var(--text-secondary); font-size: 0.9em;">
                                    Image ${index + 1} - ${img.type || 'Molecular Structure'}
                                </p>
                            </div>
                        `).join('');
                        imagesContainer.classList.remove('hidden');
                    } else {
                        imagesGrid.innerHTML = '<p style="color: var(--text-secondary); text-align: center;">No molecular images found in this document.</p>';
                        imagesContainer.classList.remove('hidden');
                    }
                } catch (error) {
                    console.error('Error extracting images:', error);
                    const imagesGrid = document.getElementById('images-grid');
                    imagesGrid.innerHTML = '<p style="color: var(--danger); text-align: center;">Error extracting images from document.</p>';
                    document.getElementById('extracted-images').classList.remove('hidden');
                }
            }

            // 全局变量存储当前处理的结果
            let currentModelResults = {
                modelA: null,
                modelB: null,
                selectedImages: []
            };

            // AIChemist API 集成函数
            async function processWithAIChemist(documentId) {
                try {
                    // 调用后端 API 获取化学结构图像并使用 AIChemist 处理
                    const response = await fetch('/api/molecules/process-aichemist', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            document_id: documentId
                        })
                    });

                    if (!response.ok) {
                        throw new Error(`HTTP error! status: ${response.status}`);
                    }

                    const data = await response.json();

                    if (!data.status || data.status !== 'success') {
                        throw new Error(data.error || 'AIChemist processing failed');
                    }

                    return data;

                } catch (error) {
                    console.error('AIChemist API error:', error);
                    throw new Error(`AIChemist processing error: ${error.message}`);
                }
            }

            // 存储当前选中的模型
            let currentSelectedModel = null;

            // 存储当前选中的Molecule Processing按钮
            let currentSelectedMoleculeButton = null;

            // 处理Molecule Processing模块的数据按钮点击
            function selectMoleculeButton(buttonId, functionName) {
                // 获取所有Molecule Processing数据按钮
                const dataButtons = document.querySelectorAll('#extract-db-btn, #view-smiles-btn, #view-segmented-btn');

                // 移除所有数据按钮的selected类
                dataButtons.forEach(btn => {
                    btn.classList.remove('selected');
                });

                // 为当前点击的按钮添加selected类
                const currentButton = document.getElementById(buttonId);
                if (currentButton) {
                    currentButton.classList.add('selected');
                    currentSelectedMoleculeButton = buttonId;
                }

                // 执行对应的函数
                if (typeof window[functionName] === 'function') {
                    window[functionName]();
                }
            }

            // 处理Model按钮点击
            function selectModelButton(buttonId, modelType) {
                // 获取所有Model按钮
                const modelButtons = document.querySelectorAll('#model-a-btn, #model-b-btn, #model-c-btn');

                // 移除所有Model按钮的selected类
                modelButtons.forEach(btn => {
                    btn.classList.remove('selected');
                });

                // 为当前点击的按钮添加selected类
                const currentButton = document.getElementById(buttonId);
                if (currentButton) {
                    currentButton.classList.add('selected');
                }

                // 调用processWithModel
                processWithModel(modelType);
            }

            async function processWithModel(modelType) {
                // 更新按钮高亮状态 - 这个应该总是执行，不管是否选择了文档
                updateModelButtonHighlight(modelType);
                currentSelectedModel = modelType;
                console.log('✅ Model selected:', modelType);

                const documentId = document.getElementById('molecule-document-select').value;
                const result = document.getElementById('molecule-result');

                if (!documentId) {
                    // 如果没有选择文档，只更新模型选择，不显示错误
                    console.log('ℹ️ No document selected, but model selection updated');
                    return;
                }

                // 显示加载状态
                const modelName = modelType === 'modelA' ? 'DECIMER' :
                                 modelType === 'modelB' ? 'AIChemist' : 'YOLO';
                result.innerHTML = `
                    <div style="padding: 15px; background: linear-gradient(135deg, rgba(0, 212, 255, 0.1), rgba(0, 212, 255, 0.05)); border: 1px solid var(--primary); border-radius: 10px;">
                        <i class="fas fa-spinner fa-spin" style="color: var(--primary); margin-right: 10px;"></i>
                        Processing with ${modelName}...
                    </div>
                `;
                result.classList.remove('hidden');

                try {
                    // Model A 使用 DECIMER，Model B 使用 AIChemist API，Model C 使用 YOLO
                    if (modelType === 'modelB') {
                        // Model B: 使用 AIChemist API
                        const data = await processWithAIChemist(documentId);
                        await handleSingleModelProcessing(modelType, data);
                    } else if (modelType === 'modelA') {
                        // Model A: 使用 DECIMER
                        const response = await fetch('/api/molecules/process', {
                            method: 'POST',
                            headers: {
                                'Content-Type': 'application/json',
                            },
                            body: JSON.stringify({
                                document_id: documentId,
                                model_type: 'modelA'
                            })
                        });

                        const data = await response.json();
                        await handleSingleModelProcessing(modelType, data);
                    } else if (modelType === 'modelC') {
                        // Model C: 使用 YOLO 预处理
                        console.log('ℹ️ Model C (YOLO) selected - YOLO preprocessing will be used when extracting SMILES');
                        result.innerHTML = `
                            <div style="padding: 15px; background: linear-gradient(135deg, rgba(155, 89, 182, 0.1), rgba(155, 89, 182, 0.05)); border: 1px solid #9b59b6; border-radius: 10px;">
                                <i class="fas fa-check-circle" style="color: #9b59b6; margin-right: 10px;"></i>
                                Model C (YOLO) selected. YOLO preprocessing will be applied when you click "Extract SMILES".
                                <br><br>
                                <strong>Next steps:</strong>
                                <ol style="margin-top: 10px; padding-left: 20px;">
                                    <li>Load a document using "Load Filenames"</li>
                                    <li>View segmented images</li>
                                    <li>Click "Extract SMILES" - images will be preprocessed with YOLO first</li>
                                    <li>Then select Model A or B to extract SMILES from YOLO-processed images</li>
                                </ol>
                            </div>
                        `;
                        result.classList.remove('hidden');
                    }

                } catch (error) {
                    result.innerHTML = `
                        <div style="padding: 15px; background: linear-gradient(135deg, rgba(255, 23, 68, 0.1), rgba(255, 23, 68, 0.05)); border: 1px solid var(--danger); border-radius: 10px;">
                            <i class="fas fa-exclamation-triangle" style="color: var(--danger); margin-right: 10px;"></i>
                            Error processing with ${modelName}: ${error.message}
                        </div>
                    `;
                    result.classList.remove('hidden');
                }
            }

            // 更新模型按钮高亮状态
            function updateModelButtonHighlight(modelType) {
                const modelABtn = document.getElementById('model-a-btn');
                const modelBBtn = document.getElementById('model-b-btn');
                const modelCBtn = document.getElementById('model-c-btn');
                const modelACheckmark = document.getElementById('model-a-checkmark');
                const modelBCheckmark = document.getElementById('model-b-checkmark');
                const modelCCheckmark = document.getElementById('model-c-checkmark');
                const selectedDisplay = document.getElementById('selected-model-display');
                const selectedText = document.getElementById('selected-model-text');

                console.log('🔍 updateModelButtonHighlight called with modelType:', modelType);
                console.log('modelABtn:', modelABtn);
                console.log('modelBBtn:', modelBBtn);
                console.log('modelCBtn:', modelCBtn);

                // 重置所有按钮和对勾 - 移除selected类
                if (modelABtn) {
                    modelABtn.classList.remove('selected');
                }
                if (modelBBtn) {
                    modelBBtn.classList.remove('selected');
                }
                if (modelCBtn) {
                    modelCBtn.classList.remove('selected');
                }
                // 隐藏所有对勾
                if (modelACheckmark) {
                    modelACheckmark.style.display = 'none';
                }
                if (modelBCheckmark) {
                    modelBCheckmark.style.display = 'none';
                }
                if (modelCCheckmark) {
                    modelCCheckmark.style.display = 'none';
                }

                // 高亮选中的按钮 - 添加selected类（灰色填充，无荧光效果）
                if (modelType === 'modelA' && modelABtn) {
                    console.log('✅ Highlighting Model A button');
                    modelABtn.classList.add('selected');
                    selectedText.textContent = '✓ Model A (DECIMER) Selected';
                } else if (modelType === 'modelB' && modelBBtn) {
                    console.log('✅ Highlighting Model B button');
                    modelBBtn.classList.add('selected');
                    selectedText.textContent = '✓ Model B (AIChemist) Selected';
                } else if (modelType === 'modelC' && modelCBtn) {
                    console.log('✅ Highlighting Model C button');
                    modelCBtn.classList.add('selected');
                    selectedText.textContent = '✓ Model C (YOLO) Selected';
                }

                // 显示选中状态
                if (selectedDisplay) {
                    selectedDisplay.style.display = 'block';
                }
            }

            async function handleSingleModelProcessing(modelType, data) {
                const result = document.getElementById('molecule-result');

                // 存储结果
                currentModelResults[modelType] = data;

                // 获取模型名称
                const modelName = modelType === 'modelA' ? 'DECIMER' : 'AIChemist';
                const methodName = modelType === 'modelA' ? 'DECIMER Model' : 'AIChemist API';

                let resultHtml = `
                    <div class="success-box" style="padding: 15px; border-radius: 10px; margin-bottom: 15px;">
                        <i class="fas fa-check-circle" style="color: var(--success); margin-right: 10px;"></i>
                        ${modelName} processing completed successfully!
                    </div>
                    <p><strong>Method:</strong> <span style="color: var(--accent); font-weight: bold;">${methodName}</span></p>
                    <p><strong>Model:</strong> ${modelName}</p>
                    <p><strong>Processed Images:</strong> ${data.processed_count}</p>
                    <p><strong>SMILES Generated:</strong> ${data.smiles_count}</p>
                    <p><strong>Confidence:</strong> <span style="color: var(--success); font-weight: bold;">${data.confidence}%</span></p>
                `;

                // 显示生成的分子结果
                if (data.results && data.results.length > 0) {
                    resultHtml += `
                        <div style="margin-top: 20px;">
                            <h4 style="color: var(--primary); margin-bottom: 15px;">Generated Molecular Structures:</h4>
                    `;

                    if (data.results.length > 1) {
                        resultHtml += `
                            <label style="color: var(--text-secondary); margin-bottom: 10px; display: block;">
                                Select result to view:
                            </label>
                            <select id="result-selector" class="select-field" onchange="showSelectedResult(this.value, '${modelType}')">
                                ${data.results.map((result, index) => `
                                    <option value="${index}">Result ${index + 1} - ${result.smiles}</option>
                                `).join('')}
                            </select>
                        `;
                    }

                    resultHtml += `
                            <div id="selected-result-display">
                                ${generateResultDisplay(data.results[0], 0)}
                            </div>
                        </div>
                    `;
                }

                result.innerHTML = resultHtml;
                result.classList.remove('hidden');
            }

            async function handleEnsembleProcessing(data) {
                const result = document.getElementById('molecule-result');
                const ensembleSelection = document.getElementById('ensemble-selection');

                // 存储两个模型的结果
                currentModelResults.modelA = data.modelA_results;
                currentModelResults.modelB = data.modelB_results;

                // 检查SMILES是否一致
                const disagreements = data.disagreements || [];

                if (disagreements.length === 0) {
                    // 所有结果一致，显示ensemble结果
                    result.innerHTML = `
                        <div class="success-box" style="padding: 15px; border-radius: 10px; margin-bottom: 15px;">
                            <i class="fas fa-check-circle" style="color: var(--success); margin-right: 10px;"></i>
                            Ensemble processing completed! All models agree.
                        </div>
                        <p><strong>Model Agreement:</strong> <span style="color: var(--success); font-weight: bold;">100%</span></p>
                        <p><strong>Processed Images:</strong> ${data.processed_count}</p>
                        <p><strong>SMILES Generated:</strong> ${data.smiles_count}</p>
                        <p><strong>Ensemble Confidence:</strong> <span style="color: var(--success); font-weight: bold;">${data.confidence}%</span></p>

                        <div style="margin-top: 20px;">
                            <h4 style="color: var(--primary); margin-bottom: 15px;">Consensus Results:</h4>
                            <div id="consensus-results">
                                ${data.consensus_results.map((result, index) => generateResultDisplay(result, index)).join('')}
                            </div>
                        </div>
                    `;
                    result.classList.remove('hidden');
                } else {
                    // 有不一致的结果，需要人工选择
                    result.innerHTML = `
                        <div style="padding: 15px; background: linear-gradient(135deg, rgba(255, 215, 0, 0.1), rgba(255, 215, 0, 0.05)); border: 1px solid var(--warning); border-radius: 10px; margin-bottom: 15px;">
                            <i class="fas fa-exclamation-triangle" style="color: var(--warning); margin-right: 10px;"></i>
                            Ensemble processing completed with disagreements requiring manual review.
                        </div>
                        <p><strong>Model Agreement:</strong> <span style="color: var(--warning); font-weight: bold;">${((data.processed_count - disagreements.length) / data.processed_count * 100).toFixed(1)}%</span></p>
                        <p><strong>Processed Images:</strong> ${data.processed_count}</p>
                        <p><strong>Disagreements:</strong> <span style="color: var(--warning); font-weight: bold;">${disagreements.length}</span></p>
                        <p><strong>Consensus Results:</strong> ${data.processed_count - disagreements.length}</p>
                    `;
                    result.classList.remove('hidden');

                    // 显示不一致的结果供人工选择
                    showDisagreements(disagreements);
                }
            }

            function showDisagreements(disagreements) {
                const ensembleSelection = document.getElementById('ensemble-selection');
                const disagreementImages = document.getElementById('disagreement-images');

                disagreementImages.innerHTML = disagreements.map((disagreement, index) => `
                    <div class="disagreement-card" style="border: 1px solid var(--warning); border-radius: 15px; padding: 20px; background: var(--glass);">
                        <h5 style="color: var(--warning); margin-bottom: 15px; text-align: center;">
                            Disagreement ${index + 1}
                        </h5>

                        <!-- 原始图片 -->
                        <div style="text-align: center; margin-bottom: 20px;">
                            <h6 style="color: var(--text-secondary); margin-bottom: 10px;">Original Image:</h6>
                            <img src="${disagreement.original_image}" alt="Original"
                                 style="max-width: 100%; height: 150px; object-fit: contain; border-radius: 5px; background: white;">
                        </div>

                        <!-- Model A结果 -->
                        <div style="border: 1px solid var(--primary); border-radius: 10px; padding: 15px; margin-bottom: 15px; background: rgba(0, 212, 255, 0.05);">
                            <h6 style="color: var(--primary); margin-bottom: 10px;">
                                <i class="fas fa-robot" style="margin-right: 5px;"></i>
                                Model A Result
                            </h6>
                            <p><strong>SMILES:</strong> <code style="background: var(--surface); padding: 3px; border-radius: 3px;">${disagreement.modelA_smiles}</code></p>
                            <p><strong>Confidence:</strong> <span style="color: var(--primary);">${disagreement.modelA_confidence}%</span></p>
                            <img src="${disagreement.modelA_image}" alt="Model A Result"
                                 style="max-width: 100%; height: 120px; object-fit: contain; border-radius: 5px; background: white; margin-top: 10px;">
                        </div>

                        <!-- Model B结果 -->
                        <div style="border: 1px solid var(--secondary); border-radius: 10px; padding: 15px; background: rgba(0, 99, 207, 0.05);">
                            <h6 style="color: var(--secondary); margin-bottom: 10px;">
                                <i class="fas fa-cog" style="margin-right: 5px;"></i>
                                Model B Result
                            </h6>
                            <p><strong>SMILES:</strong> <code style="background: var(--surface); padding: 3px; border-radius: 3px;">${disagreement.modelB_smiles}</code></p>
                            <p><strong>Confidence:</strong> <span style="color: var(--secondary);">${disagreement.modelB_confidence}%</span></p>
                            <img src="${disagreement.modelB_image}" alt="Model B Result"
                                 style="max-width: 100%; height: 120px; object-fit: contain; border-radius: 5px; background: white; margin-top: 10px;">
                        </div>
                    </div>
                `).join('');

                ensembleSelection.classList.remove('hidden');
            }

            function generateResultDisplay(result, index) {
                return `
                    <div style="border: 1px solid var(--border); border-radius: 10px; padding: 15px; margin: 10px 0; background: var(--glass);">
                        <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 15px;">
                            <h5 style="color: var(--primary); margin: 0;">Molecule ${index + 1}</h5>
                            <span style="color: var(--success); font-weight: bold;">Confidence: ${result.confidence}%</span>
                        </div>

                        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 20px; align-items: center;">
                            <div>
                                <p><strong>SMILES:</strong></p>
                                <code style="background: var(--surface); padding: 8px; border-radius: 5px; display: block; margin: 5px 0; word-break: break-all;">
                                    ${result.smiles}
                                </code>
                                <p><strong>Molecular Weight:</strong> ${result.molecular_weight} g/mol</p>
                                <p><strong>Formula:</strong> ${result.molecular_formula}</p>
                            </div>
                            <div style="text-align: center;">
                                <img src="${result.structure_image}" alt="Generated Structure"
                                     style="max-width: 100%; height: 150px; object-fit: contain; border-radius: 5px; background: white;">
                            </div>
                        </div>
                    </div>
                `;
            }

            function showSelectedResult(index, modelType) {
                const results = currentModelResults[modelType]?.results || [];
                if (results[index]) {
                    document.getElementById('selected-result-display').innerHTML = generateResultDisplay(results[index], index);
                }
            }

            function selectManualResult(selectedModel) {
                const result = document.getElementById('molecule-result');
                const ensembleSelection = document.getElementById('ensemble-selection');

                // 隐藏人工选择界面
                ensembleSelection.classList.add('hidden');

                // 显示最终选择的结果
                const selectedResults = currentModelResults[selectedModel];

                result.innerHTML = `
                    <div class="success-box" style="padding: 15px; border-radius: 10px; margin-bottom: 15px;">
                        <i class="fas fa-check-circle" style="color: var(--success); margin-right: 10px;"></i>
                        Manual selection completed! ${selectedModel.toUpperCase()} results selected.
                    </div>
                    <p><strong>Selected Model:</strong> ${selectedModel.toUpperCase()}</p>
                    <p><strong>Selection Method:</strong> Manual Review</p>
                    <p><strong>Final Results:</strong> ${selectedResults.results?.length || 0}</p>

                    <div style="margin-top: 20px;">
                        <h4 style="color: var(--primary); margin-bottom: 15px;">Final Selected Results:</h4>
                        <div id="final-results">
                            ${selectedResults.results?.map((result, index) => generateResultDisplay(result, index)).join('') || 'No results available'}
                        </div>
                    </div>
                `;

                result.classList.remove('hidden');
            }

            // SMILES提取相关函数
            async function extractSmilesFromDatabase() {
                const filenameFilter = document.getElementById('smiles-filename-filter').value.trim();
                const resultDiv = document.getElementById('smiles-result');

                // 确定显示的文件名信息
                const displayFilename = filenameFilter ? filenameFilter : 'All Images';
                const isAllImages = !filenameFilter;

                resultDiv.innerHTML = `
                    <div style="text-align: center; padding: 20px;">
                        <i class="fas fa-spinner fa-spin" style="font-size: 2em; color: var(--secondary);"></i>
                        <br><br>Extracting SMILES from ${isAllImages ? 'all images' : `images in file: <strong>${filenameFilter}</strong>`}...
                    </div>
                `;
                resultDiv.classList.remove('hidden');

                try {
                    const formData = new FormData();
                    formData.append('filename', filenameFilter);

                    const response = await fetch('/api/smiles/extract', {
                        method: 'POST',
                        body: formData
                    });

                    const data = await response.json();

                    if (data.success && data.results.length > 0) {
                        let resultHtml = `
                            <div style="color: var(--success); padding: 15px; text-align: center; margin-bottom: 20px;">
                                <i class="fas fa-check-circle" style="margin-right: 10px;"></i>
                                Successfully processed ${data.total_images} images! (${data.success_count} SMILES extracted and stored)
                            </div>
                            <div style="max-height: 400px; overflow-y: auto;">
                        `;

                        data.results.forEach((result, index) => {
                            const smilesResult = result.smiles_result || {};
                            const isSuccess = smilesResult.success;
                            const smiles = smilesResult.smiles;

                            // 构建数据库存储状态
                            let dbStatusHtml = '';
                            if (result.smiles_db_id) {
                                dbStatusHtml = `
                                    <p style="color: var(--success); font-size: 12px;">
                                        <i class="fas fa-database" style="margin-right: 5px;"></i>
                                        <strong>Stored to DB:</strong> ${result.smiles_db_id}
                                    </p>
                                `;
                            }

                            resultHtml += `
                                <div style="border: 1px solid var(--border); border-radius: 8px; padding: 15px; margin-bottom: 15px; background: var(--card-bg);">
                                    <h4 style="color: var(--primary); margin-bottom: 10px;">
                                        Result ${index + 1}
                                        ${isSuccess ? '<i class="fas fa-check-circle" style="color: var(--success); margin-left: 10px;"></i>' : '<i class="fas fa-times-circle" style="color: var(--error); margin-left: 10px;"></i>'}
                                    </h4>
                                    <p><strong>File ID:</strong> ${result.file_id}</p>
                                    <p><strong>Filename:</strong> ${result.filename}</p>
                                    <p><strong>Image Type:</strong> ${result.metadata?.image_type || 'N/A'}</p>
                                    ${isSuccess ? `
                                        <p><strong>SMILES:</strong> <code style="background: var(--bg-secondary); padding: 5px; border-radius: 4px; word-break: break-all;">${smiles}</code></p>
                                        <p><strong>Confidence:</strong> ${smilesResult.confidence || 'N/A'}</p>
                                        <p><strong>Method:</strong> ${smilesResult.method || 'DECIMER'}</p>
                                        ${dbStatusHtml}
                                    ` : `
                                        <p style="color: var(--error);"><strong>Error:</strong> ${smilesResult.error || result.error || 'SMILES extraction failed'}</p>
                                    `}
                                </div>
                            `;
                        });

                        resultHtml += '</div>';
                        resultDiv.innerHTML = resultHtml;
                    } else {
                        resultDiv.innerHTML = `
                            <div style="color: var(--warning); padding: 15px; text-align: center;">
                                <i class="fas fa-exclamation-triangle" style="margin-right: 10px;"></i>
                                ${data.error || 'No SMILES results found'}
                            </div>
                        `;
                    }
                } catch (error) {
                    resultDiv.innerHTML = `
                        <div style="color: var(--error); padding: 15px; text-align: center;">
                            <i class="fas fa-times-circle" style="margin-right: 10px;"></i>
                            Error extracting SMILES: ${error.message}
                        </div>
                    `;
                }
            }

            // 无限滚动全局变量
            let smilesInfiniteScroll = {
                offset: 0,
                limit: 50,
                total_count: 0,
                has_more: true,
                is_loading: false,
                table_container: null,
                table_body: null
            };

            async function loadSmilesResults() {
                const resultDiv = document.getElementById('smiles-result');

                // 初始化无限滚动
                smilesInfiniteScroll.offset = 0;
                smilesInfiniteScroll.has_more = true;
                smilesInfiniteScroll.is_loading = false;

                resultDiv.innerHTML = '<div style="text-align: center; padding: 20px;"><i class="fas fa-spinner fa-spin" style="font-size: 2em; color: var(--accent);"></i><br><br>Loading SMILES results...</div>';
                resultDiv.classList.remove('hidden');

                try {
                    // 加载第一批数据
                    const data = await loadMoreSmilesResultsWithData();

                    if (data && data.success && data.total_count > 0) {
                        // 更新无限滚动状态
                        smilesInfiniteScroll.total_count = data.total_count;
                        smilesInfiniteScroll.has_more = data.has_more;

                        // 存储数据供下载使用
                        if (!window.smilesResultsData) {
                            window.smilesResultsData = [];
                        }

                        let resultHtml = `
                            <!-- 标题和下载按钮容器 -->
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px; flex-wrap: wrap; gap: 10px;">
                                <div style="color: var(--success); padding: 15px; text-align: center;">
                                    <i class="fas fa-database" style="margin-right: 10px;"></i>
                                    Found <span id="smiles-total-count">${data.total_count}</span> SMILES results in database!
                                    <span id="smiles-loading-status" style="margin-left: 10px; color: var(--text-secondary); font-size: 12px;"></span>
                                </div>
                                <div style="display: flex; gap: 10px;">
                                    <button onclick="downloadSmilesXLSXDirect()"
                                            style="background: linear-gradient(45deg, #4CAF50, #45a049);
                                                   color: white; border: none; padding: 10px 15px; border-radius: 6px;
                                                   cursor: pointer; font-size: 12px; display: flex; align-items: center;
                                                   transition: all 0.3s ease; box-shadow: 0 2px 4px rgba(0,0,0,0.2); font-weight: bold;"
                                            onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 8px rgba(0,0,0,0.3)';"
                                            onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='0 2px 4px rgba(0,0,0,0.2)';">
                                        <i class="fas fa-download" style="margin-right: 8px;"></i>
                                        Download All (Direct)
                                    </button>
                                    <button onclick="downloadSmilesXLSX()"
                                            style="background: linear-gradient(45deg, var(--success), var(--accent));
                                                   color: white; border: none; padding: 10px 15px; border-radius: 6px;
                                                   cursor: pointer; font-size: 12px; display: flex; align-items: center;
                                                   transition: all 0.3s ease; box-shadow: 0 2px 4px rgba(0,0,0,0.2);"
                                            onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 8px rgba(0,0,0,0.3)';"
                                            onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='0 2px 4px rgba(0,0,0,0.2)';">
                                        <i class="fas fa-file-excel" style="margin-right: 8px;"></i>
                                        Download Loaded
                                    </button>
                                </div>
                            </div>

                            <!-- 表格容器 (启用无限滚动) -->
                            <div id="smiles-table-container" style="max-height: 600px; overflow-y: auto; border: 1px solid var(--border); border-radius: 8px;">
                                <table id="smiles-results-table" style="width: 100%; border-collapse: collapse; background: var(--card-bg);">
                                    <thead style="background: var(--bg-secondary); position: sticky; top: 0; z-index: 10;">
                                        <tr>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 60px;">#</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: center; min-width: 120px;">Chemical Image</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 200px;">SMILES</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: center; min-width: 120px;">Structure (RDKit)</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 120px;">Image File ID</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 120px;">Database ID</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: center; min-width: 80px;">Confidence</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 120px;">Filename</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 100px;">Source File</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 80px;">Method</th>
                                            <th style="padding: 12px 8px; border-bottom: 2px solid var(--border); color: var(--primary); font-weight: bold; text-align: left; min-width: 120px;">Extraction Date</th>
                                        </tr>
                                    </thead>
                                    <tbody id="smiles-table-body">
                        `;

                        // 添加行到表格
                        const addRowsToTable = (results, startIndex) => {
                            let rowsHtml = '';
                            results.forEach((result, index) => {
                                const rowNumber = startIndex + index + 1;

                                // 格式化日期
                                const extractionDate = result.extraction_date ?
                                    new Date(result.extraction_date).toLocaleString() : 'N/A';

                                // 截断长ID显示
                                const shortImageId = result.image_file_id ?
                                    result.image_file_id.substring(0, 8) + '...' : 'N/A';
                                const shortDbId = result._id ?
                                    result._id.substring(0, 8) + '...' : 'N/A';

                                // 截断长文件名（切割后的文件名）
                                const shortFilename = result.filename ?
                                    (result.filename.length > 20 ?
                                        result.filename.substring(0, 17) + '...' :
                                        result.filename) : 'N/A';

                                // 截断长源文件名
                                const shortSourceFile = result.source_file ?
                                    (result.source_file.length > 20 ?
                                        result.source_file.substring(0, 17) + '...' :
                                        result.source_file) : 'N/A';

                                // 使用替换后的SMILES(如果有的话),否则使用原始SMILES
                                const finalSmiles = result.smiles_replaced || result.smiles || 'N/A';
                                const hasReplacement = result.has_ce_replacement === true;

                                rowsHtml += `
                                    <tr style="border-bottom: 1px solid var(--border); transition: background-color 0.2s;"
                                        onmouseover="this.style.backgroundColor='var(--bg-secondary)'"
                                        onmouseout="this.style.backgroundColor='transparent'">
                                        <td style="padding: 10px 8px; font-weight: bold; color: var(--accent);">${rowNumber}</td>
                                        <td style="padding: 8px; text-align: center;">
                                            <div id="image-cell-${rowNumber - 1}" style="width: 80px; height: 80px; border: 1px solid var(--border); border-radius: 4px; display: flex; align-items: center; justify-content: center; background: var(--bg-secondary); margin: 0 auto;">
                                                <i class="fas fa-spinner fa-spin" style="color: var(--text-secondary);"></i>
                                            </div>
                                        </td>
                                        <td style="padding: 10px 8px; font-family: monospace; font-size: 12px; word-break: break-all; max-width: 200px; ${hasReplacement ? 'background: rgba(76, 175, 80, 0.1);' : ''}"
                                            title="${finalSmiles}">
                                            ${finalSmiles}
                                            ${hasReplacement ? '<span style="color: var(--success); margin-left: 5px;" title="Marker replaced">✓</span>' : ''}
                                        </td>
                                        <td style="padding: 8px; text-align: center;">
                                            <div id="rdkit-cell-${rowNumber - 1}" style="width: 80px; height: 80px; border: 1px solid var(--border); border-radius: 4px; display: flex; align-items: center; justify-content: center; background: var(--bg-secondary); margin: 0 auto;">
                                                ${hasReplacement ? '<i class="fas fa-spinner fa-spin" style="color: var(--text-secondary);"></i>' : '<span style="color: var(--text-secondary); font-size: 10px;">N/A</span>'}
                                            </div>
                                        </td>
                                        <td style="padding: 10px 8px; font-family: monospace; font-size: 11px; color: var(--text-secondary);"
                                            title="${result.image_file_id}">${shortImageId}</td>
                                        <td style="padding: 10px 8px; font-family: monospace; font-size: 11px; color: var(--text-secondary);"
                                            title="${result._id}">${shortDbId}</td>
                                        <td style="padding: 10px 8px; text-align: center; color: var(--success); font-weight: bold;">
                                            ${result.confidence !== undefined && result.confidence !== null ? result.confidence : 'N/A'}
                                        </td>
                                        <td style="padding: 10px 8px; font-size: 12px;"
                                            title="${result.filename || 'N/A'}">${shortFilename}</td>
                                        <td style="padding: 10px 8px; font-size: 12px;"
                                            title="${result.source_file || 'N/A'}">${shortSourceFile}</td>
                                        <td style="padding: 10px 8px; font-size: 12px;">
                                            <span style="background: var(--accent); color: var(--bg-primary); padding: 2px 6px; border-radius: 10px; font-size: 10px;">
                                                ${result.extraction_method || 'DECIMER'}
                                            </span>
                                        </td>
                                        <td style="padding: 10px 8px; font-size: 11px; color: var(--text-secondary);">
                                            ${extractionDate}
                                        </td>
                                    </tr>
                                `;
                            });
                            return rowsHtml;
                        };

                        // 初始化表格
                        resultHtml += addRowsToTable(data.results, smilesInfiniteScroll.offset);

                        resultHtml += `
                                    </tbody>
                                </table>
                                <!-- 加载更多提示 -->
                                <div id="smiles-loading-indicator" style="text-align: center; padding: 20px; display: none;">
                                    <i class="fas fa-spinner fa-spin" style="color: var(--accent); margin-right: 10px;"></i>
                                    <span style="color: var(--text-secondary);">Loading more results...</span>
                                </div>
                            </div>

                            <!-- 统计信息 -->
                            <div style="margin-top: 15px; padding: 10px; background: var(--bg-secondary); border-radius: 6px; font-size: 12px; color: var(--text-secondary);">
                                <i class="fas fa-info-circle" style="margin-right: 5px;"></i>
                                Total: <span id="smiles-loaded-count">${data.results.length}</span> / ${data.total_count} SMILES results loaded |
                                Scroll down to load more | Hover over cells to see full content
                            </div>
                        `;

                        resultDiv.innerHTML = resultHtml;

                        // 保存表格容器和tbody引用
                        smilesInfiniteScroll.table_container = document.getElementById('smiles-table-container');
                        smilesInfiniteScroll.table_body = document.getElementById('smiles-table-body');

                        // 表格插入DOM后，异步加载所有图像
                        data.results.forEach((result, index) => {
                            const rowIndex = smilesInfiniteScroll.offset + index;
                            if (result.image_file_id) {
                                loadImageForRow(result.image_file_id, rowIndex);
                            }
                            // 如果有RDKit图像ID，从数据库加载RDKit结构图像
                            if (result.rdkit_image_id && result.rdkit_image_id !== 'None') {
                                loadRDKitStructureFromDB(result.rdkit_image_id, rowIndex);
                            }
                        });

                        // 更新加载状态
                        smilesInfiniteScroll.offset += data.results.length;
                        window.smilesResultsData = window.smilesResultsData.concat(data.results);

                        // 添加滚动事件监听
                        if (smilesInfiniteScroll.table_container) {
                            smilesInfiniteScroll.table_container.addEventListener('scroll', handleSmilesScroll);
                        }
                    } else {
                        resultDiv.innerHTML = `
                            <div style="color: var(--warning); padding: 15px; text-align: center;">
                                <i class="fas fa-exclamation-triangle" style="margin-right: 10px;"></i>
                                No SMILES results found in database
                            </div>
                        `;
                    }
                } catch (error) {
                    resultDiv.innerHTML = `
                        <div style="color: var(--error); padding: 15px; text-align: center;">
                            <i class="fas fa-times-circle" style="margin-right: 10px;"></i>
                            Error loading SMILES results: ${error.message}
                        </div>
                    `;
                }
            }

            // 加载更多SMILES结果并返回数据（用于初始加载）
            async function loadMoreSmilesResultsWithData() {
                try {
                    const response = await fetch(`/api/smiles/results?offset=${smilesInfiniteScroll.offset}&limit=${smilesInfiniteScroll.limit}`);
                    const data = await response.json();

                    if (data.success) {
                        smilesInfiniteScroll.total_count = data.total_count;
                        smilesInfiniteScroll.has_more = data.has_more;
                    }

                    return data;
                } catch (error) {
                    console.error('❌ Error loading SMILES results:', error);
                    return { success: false, error: error.message };
                }
            }

            // 加载更多SMILES结果（无限滚动）
            async function loadMoreSmilesResults() {
                if (smilesInfiniteScroll.is_loading || !smilesInfiniteScroll.has_more) {
                    return;
                }

                smilesInfiniteScroll.is_loading = true;

                // 显示加载指示器
                const loadingIndicator = document.getElementById('smiles-loading-indicator');
                if (loadingIndicator) {
                    loadingIndicator.style.display = 'block';
                }

                try {
                    const response = await fetch(`/api/smiles/results?offset=${smilesInfiniteScroll.offset}&limit=${smilesInfiniteScroll.limit}`);
                    const data = await response.json();

                    if (data.success && data.results.length > 0) {
                        // 添加新行到表格
                        const newRowsHtml = (() => {
                            let html = '';
                            data.results.forEach((result, index) => {
                                const rowNumber = smilesInfiniteScroll.offset + index + 1;

                                // 格式化日期
                                const extractionDate = result.extraction_date ?
                                    new Date(result.extraction_date).toLocaleString() : 'N/A';

                                // 截断长ID显示
                                const shortImageId = result.image_file_id ?
                                    result.image_file_id.substring(0, 8) + '...' : 'N/A';
                                const shortDbId = result._id ?
                                    result._id.substring(0, 8) + '...' : 'N/A';

                                // 截断长文件名
                                const shortFilename = result.filename ?
                                    (result.filename.length > 20 ?
                                        result.filename.substring(0, 17) + '...' :
                                        result.filename) : 'N/A';

                                // 截断长源文件名
                                const shortSourceFile = result.source_file ?
                                    (result.source_file.length > 20 ?
                                        result.source_file.substring(0, 17) + '...' :
                                        result.source_file) : 'N/A';

                                // 使用替换后的SMILES
                                const finalSmiles = result.smiles_replaced || result.smiles || 'N/A';
                                const hasReplacement = result.has_ce_replacement === true;

                                html += `
                                    <tr style="border-bottom: 1px solid var(--border); transition: background-color 0.2s;"
                                        onmouseover="this.style.backgroundColor='var(--bg-secondary)'"
                                        onmouseout="this.style.backgroundColor='transparent'">
                                        <td style="padding: 10px 8px; font-weight: bold; color: var(--accent);">${rowNumber}</td>
                                        <td style="padding: 8px; text-align: center;">
                                            <div id="image-cell-${rowNumber - 1}" style="width: 80px; height: 80px; border: 1px solid var(--border); border-radius: 4px; display: flex; align-items: center; justify-content: center; background: var(--bg-secondary); margin: 0 auto;">
                                                <i class="fas fa-spinner fa-spin" style="color: var(--text-secondary);"></i>
                                            </div>
                                        </td>
                                        <td style="padding: 10px 8px; font-family: monospace; font-size: 12px; word-break: break-all; max-width: 200px; ${hasReplacement ? 'background: rgba(76, 175, 80, 0.1);' : ''}"
                                            title="${finalSmiles}">
                                            ${finalSmiles}
                                            ${hasReplacement ? '<span style="color: var(--success); margin-left: 5px;" title="Marker replaced">✓</span>' : ''}
                                        </td>
                                        <td style="padding: 8px; text-align: center;">
                                            <div id="rdkit-cell-${rowNumber - 1}" style="width: 80px; height: 80px; border: 1px solid var(--border); border-radius: 4px; display: flex; align-items: center; justify-content: center; background: var(--bg-secondary); margin: 0 auto;">
                                                ${hasReplacement ? '<i class="fas fa-spinner fa-spin" style="color: var(--text-secondary);"></i>' : '<span style="color: var(--text-secondary); font-size: 10px;">N/A</span>'}
                                            </div>
                                        </td>
                                        <td style="padding: 10px 8px; font-family: monospace; font-size: 11px; color: var(--text-secondary);"
                                            title="${result.image_file_id}">${shortImageId}</td>
                                        <td style="padding: 10px 8px; font-family: monospace; font-size: 11px; color: var(--text-secondary);"
                                            title="${result._id}">${shortDbId}</td>
                                        <td style="padding: 10px 8px; text-align: center; color: var(--success); font-weight: bold;">
                                            ${result.confidence !== undefined && result.confidence !== null ? result.confidence : 'N/A'}
                                        </td>
                                        <td style="padding: 10px 8px; font-size: 12px;"
                                            title="${result.filename || 'N/A'}">${shortFilename}</td>
                                        <td style="padding: 10px 8px; font-size: 12px;"
                                            title="${result.source_file || 'N/A'}">${shortSourceFile}</td>
                                        <td style="padding: 10px 8px; font-size: 12px;">
                                            <span style="background: var(--accent); color: var(--bg-primary); padding: 2px 6px; border-radius: 10px; font-size: 10px;">
                                                ${result.extraction_method || 'DECIMER'}
                                            </span>
                                        </td>
                                        <td style="padding: 10px 8px; font-size: 11px; color: var(--text-secondary);">
                                            ${extractionDate}
                                        </td>
                                    </tr>
                                `;
                            });
                            return html;
                        })();

                        // 将新行添加到表格
                        if (smilesInfiniteScroll.table_body) {
                            smilesInfiniteScroll.table_body.innerHTML += newRowsHtml;
                        }

                        // 异步加载新行的图像
                        data.results.forEach((result, index) => {
                            const rowIndex = smilesInfiniteScroll.offset + index;
                            if (result.image_file_id) {
                                loadImageForRow(result.image_file_id, rowIndex);
                            }
                            if (result.rdkit_image_id && result.rdkit_image_id !== 'None') {
                                loadRDKitStructureFromDB(result.rdkit_image_id, rowIndex);
                            }
                        });

                        // 更新状态
                        smilesInfiniteScroll.offset += data.results.length;
                        smilesInfiniteScroll.has_more = data.has_more;
                        window.smilesResultsData = window.smilesResultsData.concat(data.results);

                        // 更新统计信息
                        const loadedCount = document.getElementById('smiles-loaded-count');
                        if (loadedCount) {
                            loadedCount.textContent = smilesInfiniteScroll.offset;
                        }

                        console.log(`✅ Loaded ${data.results.length} more results. Total loaded: ${smilesInfiniteScroll.offset}/${smilesInfiniteScroll.total_count}`);
                    }
                } catch (error) {
                    console.error('❌ Error loading more SMILES results:', error);
                } finally {
                    smilesInfiniteScroll.is_loading = false;

                    // 隐藏加载指示器
                    const loadingIndicator = document.getElementById('smiles-loading-indicator');
                    if (loadingIndicator) {
                        loadingIndicator.style.display = 'none';
                    }
                }
            }

            // 滚动事件处理
            function handleSmilesScroll(event) {
                const container = event.target;
                const scrollTop = container.scrollTop;
                const scrollHeight = container.scrollHeight;
                const clientHeight = container.clientHeight;

                // 当滚动到底部时，加载更多数据
                if (scrollHeight - scrollTop - clientHeight < 100) {
                    loadMoreSmilesResults();
                }
            }

            // 异步加载表格行中的图像
            async function loadImageForRow(imageFileId, rowIndex) {
                const imageCell = document.getElementById(`image-cell-${rowIndex}`);
                if (!imageCell || !imageFileId) {
                    if (imageCell) {
                        imageCell.innerHTML = '<i class="fas fa-exclamation-triangle" style="color: var(--warning);"></i>';
                    }
                    return;
                }

                try {
                    const response = await fetch(`/api/images/${imageFileId}`);
                    const data = await response.json();

                    if (data.success && data.image_data) {
                        imageCell.innerHTML = `
                            <img src="data:image/png;base64,${data.image_data}"
                                 style="width: 100%; height: 100%; object-fit: contain; border-radius: 4px; cursor: pointer;"
                                 alt="Chemical Structure"
                                 onclick="showImageModal('data:image/png;base64,${data.image_data}', '${imageFileId}')"
                                 title="Click to view full size">
                        `;
                    } else {
                        imageCell.innerHTML = '<i class="fas fa-image" style="color: var(--text-secondary);"></i>';
                    }
                } catch (error) {
                    console.error(`Error loading image for row ${rowIndex}:`, error);
                    imageCell.innerHTML = '<i class="fas fa-exclamation-triangle" style="color: var(--error);"></i>';
                }
            }

            // 异步从数据库加载RDKit化学结构图像
            async function loadRDKitStructureFromDB(rdkitImageId, rowIndex) {
                const rdkitCell = document.getElementById(`rdkit-cell-${rowIndex}`);
                if (!rdkitCell || !rdkitImageId || rdkitImageId === 'None') {
                    if (rdkitCell) {
                        rdkitCell.innerHTML = '<span style="color: var(--text-secondary); font-size: 10px;">N/A</span>';
                    }
                    return;
                }

                try {
                    const response = await fetch(`/api/rdkit-images/${rdkitImageId}`);

                    if (response.ok) {
                        const blob = await response.blob();
                        const imageUrl = URL.createObjectURL(blob);

                        rdkitCell.innerHTML = `
                            <img src="${imageUrl}"
                                 style="width: 100%; height: 100%; object-fit: contain; border-radius: 4px; cursor: pointer;"
                                 alt="RDKit Structure"
                                 onclick="showImageModal('${imageUrl}', 'RDKit Structure')"
                                 title="RDKit generated structure - Click to view full size">
                        `;
                    } else {
                        rdkitCell.innerHTML = '<i class="fas fa-exclamation-triangle" style="color: var(--warning);"></i>';
                    }
                } catch (error) {
                    console.error(`Error loading RDKit structure for row ${rowIndex}:`, error);
                    rdkitCell.innerHTML = '<i class="fas fa-exclamation-triangle" style="color: var(--error);"></i>';
                }
            }

            // 显示图像模态框
            function showImageModal(imageSrc, fileId) {
                // 创建模态框
                const modal = document.createElement('div');
                modal.style.cssText = `
                    position: fixed;
                    top: 0;
                    left: 0;
                    width: 100%;
                    height: 100%;
                    background: rgba(0, 0, 0, 0.8);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    z-index: 10000;
                    cursor: pointer;
                `;

                modal.innerHTML = `
                    <div style="max-width: 90%; max-height: 90%; background: var(--card-bg); border-radius: 8px; padding: 20px; position: relative;">
                        <button onclick="this.parentElement.parentElement.remove()"
                                style="position: absolute; top: 10px; right: 10px; background: var(--error); color: white; border: none; border-radius: 50%; width: 30px; height: 30px; cursor: pointer; font-size: 16px;">
                            ×
                        </button>
                        <div style="text-align: center; margin-bottom: 15px;">
                            <h3 style="color: var(--primary); margin: 0;">Chemical Structure Image</h3>
                            <p style="color: var(--text-secondary); font-size: 12px; margin: 5px 0;">File ID: ${fileId}</p>
                        </div>
                        <img src="${imageSrc}"
                             style="max-width: 100%; max-height: 70vh; border-radius: 4px; border: 1px solid var(--border);"
                             alt="Chemical Structure">
                    </div>
                `;

                // 点击背景关闭模态框
                modal.addEventListener('click', (e) => {
                    if (e.target === modal) {
                        modal.remove();
                    }
                });

                document.body.appendChild(modal);
            }

            // 直接从后端导出所有SMILES结果为XLSX（包含嵌入图片） - 优化版本
            async function downloadSmilesXLSXDirect() {
                try {
                    // 显示加载提示
                    const loadingMessage = showLoadingMessage('🚀 Preparing to download all SMILES data from database (this may take a while for large datasets)...');

                    // 调用后端API直接导出所有数据
                    const response = await fetch('/api/smiles/export/all/excel', {
                        method: 'GET',
                        timeout: 600000  // 10分钟超时
                    });

                    if (!response.ok) {
                        const errorData = await response.json().catch(() => ({}));
                        throw new Error(errorData.error || `HTTP error! status: ${response.status}`);
                    }

                    // 获取blob数据
                    const blob = await response.blob();

                    // 创建下载链接
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;

                    // 生成文件名（包含时间戳）
                    const timestamp = new Date().toISOString().replace(/[:.]/g, '-').slice(0, 19);
                    a.download = `SMILES_All_Results_${timestamp}.xlsx`;

                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                    window.URL.revokeObjectURL(url);

                    // 隐藏加载提示
                    if (loadingMessage && loadingMessage.parentNode) {
                        loadingMessage.parentNode.removeChild(loadingMessage);
                    }

                    // 显示成功消息
                    showDownloadSuccess('Successfully downloaded all SMILES data with images!');

                } catch (error) {
                    console.error('❌ Error downloading SMILES data:', error);
                    alert(`Failed to download SMILES data: ${error.message}\n\nPlease try again or use the "Download Loaded" option.`);

                    // 隐藏加载提示
                    const loadingMessage = document.querySelector('.loading-message');
                    if (loadingMessage && loadingMessage.parentNode) {
                        loadingMessage.parentNode.removeChild(loadingMessage);
                    }
                }
            }

            // 下载SMILES结果为XLSX文件（包含嵌入图片）
            async function downloadSmilesXLSX() {
                if (!window.smilesResultsData || window.smilesResultsData.length === 0) {
                    alert('No SMILES data available for download');
                    return;
                }

                try {
                    // 显示加载提示
                    const loadingMessage = showLoadingMessage('正在准备下载文件，请稍候...');

                    // 创建ExcelJS工作簿
                    const workbook = new ExcelJS.Workbook();
                    const worksheet = workbook.addWorksheet('SMILES Results');

                    // 定义列
                    worksheet.columns = [
                        { header: 'Index', key: 'index', width: 8 },
                        { header: 'Chemical Image', key: 'image', width: 25 },
                        { header: 'SMILES', key: 'smiles', width: 50 },
                        { header: 'Structure (RDKit)', key: 'rdkitStructure', width: 25 },
                        { header: 'Image File ID', key: 'imageFileId', width: 25 },
                        { header: 'Database ID', key: 'databaseId', width: 25 },
                        { header: 'Confidence', key: 'confidence', width: 12 },
                        { header: 'Source File', key: 'sourceFile', width: 30 },
                        { header: 'Extraction Method', key: 'method', width: 15 },
                        { header: 'Extraction Date', key: 'date', width: 20 },
                        { header: 'Image Type', key: 'imageType', width: 15 },
                        { header: 'Filename', key: 'filename', width: 30 }
                    ];

                    // 设置表头样式
                    const headerRow = worksheet.getRow(1);
                    headerRow.font = { bold: true, color: { argb: 'FFFFFFFF' } };
                    headerRow.fill = {
                        type: 'pattern',
                        pattern: 'solid',
                        fgColor: { argb: 'FF4472C4' }
                    };
                    headerRow.alignment = { horizontal: 'center', vertical: 'middle' };

                    // 收集所有图片数据
                    const imagePromises = [];
                    const imageDataMap = new Map();

                    // 加载所有图片
                    for (let i = 0; i < window.smilesResultsData.length; i++) {
                        const result = window.smilesResultsData[i];
                        if (result.image_file_id) {
                            const imagePromise = loadImageAsArrayBuffer(result.image_file_id)
                                .then(arrayBuffer => {
                                    if (arrayBuffer) {
                                        imageDataMap.set(i, arrayBuffer);
                                    }
                                })
                                .catch(error => {
                                    console.error(`Failed to load image ${result.image_file_id}:`, error);
                                });
                            imagePromises.push(imagePromise);
                        }
                    }

                    // 等待所有图片加载完成
                    await Promise.all(imagePromises);

                    // 加载RDKit化学结构图像(从数据库读取)
                    const rdkitImagePromises = [];
                    const rdkitImageMap = new Map();

                    console.log(`📊 Total SMILES records: ${window.smilesResultsData.length}`);

                    for (let i = 0; i < window.smilesResultsData.length; i++) {
                        const result = window.smilesResultsData[i];
                        console.log(`Record ${i}:`, {
                            has_ce_replacement: result.has_ce_replacement,
                            rdkit_image_id: result.rdkit_image_id,
                            smiles: result.smiles
                        });

                        // 只为有rdkit_image_id的记录加载图像
                        const hasRdkitImage = result.rdkit_image_id && result.rdkit_image_id !== 'None' && result.rdkit_image_id !== null;

                        if (hasRdkitImage) {
                            console.log(`🔄 Loading RDKit structure for index ${i} from database: ${result.rdkit_image_id}`);
                            const rdkitPromise = fetch(`/api/rdkit-images/${result.rdkit_image_id}`)
                            .then(response => {
                                if (!response.ok) {
                                    throw new Error(`HTTP error! status: ${response.status}`);
                                }
                                return response.blob();
                            })
                            .then(blob => blob.arrayBuffer())
                            .then(arrayBuffer => {
                                rdkitImageMap.set(i, arrayBuffer);
                                console.log(`✅ RDKit structure loaded for index ${i}, size: ${arrayBuffer.byteLength} bytes`);
                            })
                            .catch(error => {
                                console.error(`❌ Failed to load RDKit structure for index ${i}:`, error);
                            });
                            rdkitImagePromises.push(rdkitPromise);
                        } else {
                            console.log(`⚠️ No RDKit image for index ${i} (rdkit_image_id: ${result.rdkit_image_id})`);
                        }
                    }

                    // 等待所有RDKit化学结构图像生成完成
                    await Promise.all(rdkitImagePromises);

                    // 添加数据行
                    window.smilesResultsData.forEach((result, index) => {
                        const extractionDate = result.extraction_date ?
                            new Date(result.extraction_date).toLocaleString() : 'N/A';

                        // 使用替换后的SMILES(如果有的话),否则使用原始SMILES
                        const finalSmiles = result.smiles_replaced || result.smiles || 'N/A';

                        const rowData = {
                            index: index + 1,
                            image: '', // 原始化学图片将单独添加
                            smiles: finalSmiles,
                            rdkitStructure: '', // RDKit化学结构图像将单独添加
                            imageFileId: result.image_file_id || 'N/A',
                            databaseId: result._id || 'N/A',
                            confidence: result.confidence !== undefined && result.confidence !== null ? result.confidence : 'N/A',
                            sourceFile: result.source_file || 'N/A',
                            method: result.extraction_method || 'DECIMER',
                            date: extractionDate,
                            imageType: result.image_type || 'N/A',
                            filename: result.filename || 'N/A'
                        };

                        const row = worksheet.addRow(rowData);

                        // 设置行高以容纳图片
                        row.height = 80;

                        // 如果有原始化学图片数据，添加到工作表 (Chemical Image列，索引1)
                        if (imageDataMap.has(index)) {
                            try {
                                const imageId = workbook.addImage({
                                    buffer: imageDataMap.get(index),
                                    extension: 'png'
                                });

                                worksheet.addImage(imageId, {
                                    tl: { col: 1, row: index + 1 },
                                    ext: { width: 120, height: 60 }
                                });
                            } catch (error) {
                                console.error(`Failed to add image to Excel:`, error);
                                row.getCell('image').value = 'Image load failed';
                            }
                        } else {
                            row.getCell('image').value = result.image_file_id ? 'Image not available' : 'No image';
                        }

                        // 如果有RDKit化学结构图像，添加到工作表 (Structure (RDKit)列，索引3)
                        if (rdkitImageMap.has(index)) {
                            try {
                                console.log(`📝 Adding RDKit image to Excel for row ${index + 1}`);
                                const rdkitImageBuffer = rdkitImageMap.get(index);
                                console.log(`   Buffer size: ${rdkitImageBuffer.byteLength} bytes`);

                                const rdkitImageId = workbook.addImage({
                                    buffer: rdkitImageBuffer,
                                    extension: 'png'
                                });

                                worksheet.addImage(rdkitImageId, {
                                    tl: { col: 3, row: index + 1 },
                                    ext: { width: 120, height: 60 }
                                });
                                console.log(`✅ RDKit image added to Excel for row ${index + 1}`);
                            } catch (error) {
                                console.error(`❌ Failed to add RDKit structure to Excel for row ${index + 1}:`, error);
                                row.getCell('rdkitStructure').value = 'RDKit failed';
                            }
                        } else {
                            console.log(`⚠️ No RDKit image in map for index ${index}`);
                            row.getCell('rdkitStructure').value = 'N/A';
                        }
                    });

                    // 设置数据行样式
                    worksheet.eachRow((row, rowNumber) => {
                        if (rowNumber > 1) { // 跳过表头
                            row.eachCell((cell, colNumber) => {
                                if (colNumber === 2) { // Chemical Image列
                                    cell.alignment = { horizontal: 'center', vertical: 'middle' };
                                } else {
                                    cell.alignment = { horizontal: 'left', vertical: 'middle' };
                                }
                                cell.border = {
                                    top: { style: 'thin' },
                                    left: { style: 'thin' },
                                    bottom: { style: 'thin' },
                                    right: { style: 'thin' }
                                };
                            });
                        }
                    });

                    // 生成文件名（包含时间戳）
                    const timestamp = new Date().toISOString().replace(/[:.]/g, '-').slice(0, 19);
                    const filename = `smiles_results_with_images_${timestamp}.xlsx`;

                    // 生成Excel文件并下载
                    const buffer = await workbook.xlsx.writeBuffer();
                    const blob = new Blob([buffer], {
                        type: 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    });

                    // 创建下载链接
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = filename;
                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                    window.URL.revokeObjectURL(url);

                    // 隐藏加载提示
                    if (loadingMessage && loadingMessage.parentNode) {
                        loadingMessage.parentNode.removeChild(loadingMessage);
                    }

                    // 显示成功消息
                    showDownloadSuccess();

                } catch (error) {
                    console.error('Error generating XLSX file:', error);
                    alert('Failed to generate XLSX file. Please try again.');

                    // 隐藏加载提示
                    const loadingMessage = document.querySelector('.loading-message');
                    if (loadingMessage && loadingMessage.parentNode) {
                        loadingMessage.parentNode.removeChild(loadingMessage);
                    }
                }
            }

            // 加载图片为ArrayBuffer格式（用于ExcelJS）
            async function loadImageAsArrayBuffer(imageFileId) {
                try {
                    const response = await fetch(`/api/images/${imageFileId}/view`);
                    if (!response.ok) {
                        throw new Error(`HTTP error! status: ${response.status}`);
                    }

                    const arrayBuffer = await response.arrayBuffer();
                    return arrayBuffer;
                } catch (error) {
                    console.error('Error loading image:', error);
                    return null;
                }
            }

            // 加载图片为Base64格式（保留用于其他用途）
            async function loadImageAsBase64(imageFileId) {
                try {
                    const response = await fetch(`/api/images/${imageFileId}/view`);
                    if (!response.ok) {
                        throw new Error(`HTTP error! status: ${response.status}`);
                    }

                    const blob = await response.blob();

                    return new Promise((resolve, reject) => {
                        const reader = new FileReader();
                        reader.onload = function() {
                            // 返回base64数据（包含data:image/...前缀）
                            resolve(reader.result);
                        };
                        reader.onerror = function() {
                            reject(new Error('Failed to read image as base64'));
                        };
                        reader.readAsDataURL(blob);
                    });
                } catch (error) {
                    console.error('Error loading image:', error);
                    return null;
                }
            }

            // 显示加载消息
            function showLoadingMessage(text) {
                const message = document.createElement('div');
                message.className = 'loading-message';
                message.style.cssText = `
                    position: fixed;
                    top: 50%;
                    left: 50%;
                    transform: translate(-50%, -50%);
                    background: rgba(0, 0, 0, 0.8);
                    color: white;
                    padding: 20px 30px;
                    border-radius: 8px;
                    box-shadow: 0 4px 12px rgba(0,0,0,0.3);
                    z-index: 10001;
                    font-size: 16px;
                    display: flex;
                    align-items: center;
                    backdrop-filter: blur(5px);
                `;

                message.innerHTML = `
                    <div style="display: flex; align-items: center;">
                        <div style="width: 20px; height: 20px; border: 2px solid #ffffff; border-top: 2px solid transparent; border-radius: 50%; animation: spin 1s linear infinite; margin-right: 15px;"></div>
                        ${text}
                    </div>
                `;

                // 添加旋转动画样式
                if (!document.querySelector('#loading-spinner-style')) {
                    const style = document.createElement('style');
                    style.id = 'loading-spinner-style';
                    style.textContent = `
                        @keyframes spin {
                            0% { transform: rotate(0deg); }
                            100% { transform: rotate(360deg); }
                        }
                    `;
                    document.head.appendChild(style);
                }

                document.body.appendChild(message);
                return message;
            }

            // 显示下载成功消息
            function showDownloadSuccess(customMessage) {
                const message = document.createElement('div');
                message.style.cssText = `
                    position: fixed;
                    top: 20px;
                    right: 20px;
                    background: linear-gradient(45deg, var(--success), var(--accent));
                    color: white;
                    padding: 15px 20px;
                    border-radius: 8px;
                    box-shadow: 0 4px 12px rgba(0,0,0,0.3);
                    z-index: 10000;
                    font-size: 14px;
                    display: flex;
                    align-items: center;
                    animation: slideInRight 0.3s ease-out;
                `;

                message.innerHTML = `
                    <i class="fas fa-check-circle" style="margin-right: 10px; font-size: 16px;"></i>
                    ${customMessage || 'XLSX file with embedded images downloaded successfully!'}
                `;

                document.body.appendChild(message);

                // 3秒后自动移除消息
                setTimeout(() => {
                    message.style.animation = 'slideOutRight 0.3s ease-in';
                    setTimeout(() => {
                        if (document.body.contains(message)) {
                            document.body.removeChild(message);
                        }
                    }, 300);
                }, 3000);
            }

            async function loadSegmentedImages() {
                const filenameFilter = document.getElementById('smiles-filename-filter').value.trim();
                const resultDiv = document.getElementById('segmented-images-result');

                resultDiv.innerHTML = '<div style="text-align: center; padding: 20px;"><i class="fas fa-spinner fa-spin" style="font-size: 2em; color: var(--warning);"></i><br><br>Loading segmented images...</div>';
                resultDiv.classList.remove('hidden');

                try {
                    let url = '/api/images/segmented';
                    if (filenameFilter) {
                        url += `?filename=${encodeURIComponent(filenameFilter)}`;
                    }

                    const response = await fetch(url);
                    const data = await response.json();

                    if (data.success && data.images.length > 0) {
                        let resultHtml = `
                            <div style="color: var(--success); padding: 15px; text-align: center; margin-bottom: 20px;">
                                <i class="fas fa-check-circle" style="margin-right: 10px;"></i>
                                Found ${data.images.length} segmented images!
                            </div>
                            <div style="max-height: 500px; overflow-y: auto; display: grid; grid-template-columns: repeat(auto-fill, minmax(250px, 1fr)); gap: 15px;">
                        `;

                        data.images.forEach((image, index) => {
                            resultHtml += `
                                <div style="border: 1px solid var(--border); border-radius: 8px; padding: 15px; background: var(--card-bg);">
                                    <h4 style="color: var(--primary); margin-bottom: 10px; font-size: 14px;">Image ${index + 1}</h4>
                                    <img src="data:image/png;base64,${image.image_data}"
                                         style="width: 100%; max-width: 200px; height: auto; border-radius: 4px; margin-bottom: 10px;"
                                         alt="Chemical Structure ${index + 1}">
                                    <p style="font-size: 12px;"><strong>File ID:</strong> ${image.file_id}</p>
                                    <p style="font-size: 12px;"><strong>Filename:</strong> ${image.filename}</p>
                                    <p style="font-size: 12px;"><strong>Upload Date:</strong> ${image.upload_date || 'N/A'}</p>

                                    <!-- SMILES提取按钮 -->
                                    <button class="neon-button"
                                            onclick="extractSmilesFromSingleImage('${image.file_id}', ${index})"
                                            style="width: 100%; margin-top: 10px; padding: 8px; font-size: 12px; background: linear-gradient(45deg, var(--secondary), var(--accent));">
                                        <i class="fas fa-atom" style="margin-right: 5px;"></i>
                                        Extract SMILES
                                    </button>

                                    <!-- SMILES结果显示区域 -->
                                    <div id="smiles-result-${index}" style="margin-top: 10px; padding: 8px; background: var(--glass); border-radius: 4px; display: none;">
                                        <div id="smiles-content-${index}"></div>
                                    </div>
                                </div>
                            `;
                        });

                        resultHtml += '</div>';
                        resultDiv.innerHTML = resultHtml;
                    } else {
                        resultDiv.innerHTML = `
                            <div style="color: var(--warning); padding: 15px; text-align: center;">
                                <i class="fas fa-exclamation-triangle" style="margin-right: 10px;"></i>
                                ${data.error || 'No segmented images found'}
                            </div>
                        `;
                    }
                } catch (error) {
                    resultDiv.innerHTML = `
                        <div style="color: var(--error); padding: 15px; text-align: center;">
                            <i class="fas fa-times-circle" style="margin-right: 10px;"></i>
                            Error loading segmented images: ${error.message}
                        </div>
                    `;
                }
            }

            // 从单个图片提取SMILES
            async function extractSmilesFromSingleImage(fileId, imageIndex) {
                const resultDiv = document.getElementById(`smiles-result-${imageIndex}`);
                const contentDiv = document.getElementById(`smiles-content-${imageIndex}`);

                // 显示加载状态
                resultDiv.style.display = 'block';
                contentDiv.innerHTML = `
                    <div style="text-align: center; color: var(--warning);">
                        <i class="fas fa-spinner fa-spin" style="margin-right: 5px;"></i>
                        Extracting SMILES...
                    </div>
                `;

                try {
                    // 获取当前选中的模型，如果没有选择则默认为 modelA
                    const modelType = currentSelectedModel || 'modelA';
                    console.log('🔍 extractSmilesFromSingleImage - Using model:', modelType);

                    // 如果选择了 Model C (YOLO)，需要先用YOLO处理，然后让用户选择Model A或B
                    if (modelType === 'modelC') {
                        contentDiv.innerHTML = `
                            <div style="padding: 15px; background: linear-gradient(135deg, rgba(155, 89, 182, 0.1), rgba(155, 89, 182, 0.05)); border: 1px solid #9b59b6; border-radius: 8px;">
                                <div style="color: #9b59b6; margin-bottom: 10px;">
                                    <i class="fas fa-eye" style="margin-right: 5px;"></i>
                                    <strong>YOLO Preprocessing Mode</strong>
                                </div>
                                <p style="margin-bottom: 10px;">Please select which model to use for SMILES extraction after YOLO preprocessing:</p>
                                <div style="display: flex; gap: 10px; justify-content: center;">
                                    <button onclick="extractWithYOLOAndModel('${fileId}', ${imageIndex}, 'modelA')"
                                            class="neon-button"
                                            style="background: linear-gradient(45deg, var(--primary), var(--accent)); padding: 8px 16px;">
                                        <i class="fas fa-robot" style="margin-right: 5px;"></i>
                                        Use Model A (DECIMER)
                                    </button>
                                    <button onclick="extractWithYOLOAndModel('${fileId}', ${imageIndex}, 'modelB')"
                                            class="neon-button"
                                            style="background: linear-gradient(45deg, var(--secondary), var(--warning)); padding: 8px 16px;">
                                        <i class="fas fa-cog" style="margin-right: 5px;"></i>
                                        Use Model B (AIChemist)
                                    </button>
                                </div>
                            </div>
                        `;
                        return;
                    }

                    const response = await fetch('/api/smiles/extract-single', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            file_id: fileId,
                            model_type: modelType,
                            use_yolo: false
                        })
                    });

                    const data = await response.json();

                    if (data.success && data.smiles) {
                        // 构建数据库存储状态显示
                        let dbStatusHtml = '';
                        if (data.stored_to_db) {
                            dbStatusHtml = `
                                <div style="color: var(--success); font-size: 11px; margin-top: 5px;">
                                    <i class="fas fa-database" style="margin-right: 5px;"></i>
                                    Stored to database (ID: ${data.smiles_db_id || 'N/A'})
                                </div>
                            `;
                        } else if (data.stored_to_db === false && data.message) {
                            dbStatusHtml = `
                                <div style="color: var(--warning); font-size: 11px; margin-top: 5px;">
                                    <i class="fas fa-info-circle" style="margin-right: 5px;"></i>
                                    ${data.message}
                                </div>
                            `;
                        }

                        // 获取方法名称
                        const methodName = data.method || (modelType === 'modelB' ? 'AIChemist API' : 'DECIMER Model');

                        contentDiv.innerHTML = `
                            <div style="color: var(--success); margin-bottom: 8px;">
                                <i class="fas fa-check-circle" style="margin-right: 5px;"></i>
                                <strong>SMILES Extracted Successfully!</strong>
                            </div>
                            <div style="background: var(--card-bg); padding: 8px; border-radius: 4px; margin-bottom: 8px;">
                                <strong>SMILES:</strong>
                                <span style="font-family: monospace; color: var(--accent); word-break: break-all;">${data.smiles}</span>
                            </div>
                            <div style="font-size: 11px; color: var(--text-secondary);">
                                <div><strong>Method:</strong> <span style="color: var(--accent); font-weight: bold;">${methodName}</span></div>
                                <div><strong>Confidence:</strong> ${data.confidence !== undefined && data.confidence !== null ? data.confidence : 'N/A'}</div>
                                <div><strong>File ID:</strong> ${data.file_id}</div>
                            </div>
                            ${dbStatusHtml}
                        `;
                    } else {
                        contentDiv.innerHTML = `
                            <div style="color: var(--warning); text-align: center;">
                                <i class="fas fa-exclamation-triangle" style="margin-right: 5px;"></i>
                                ${data.error || 'SMILES extraction failed'}
                            </div>
                        `;
                    }
                } catch (error) {
                    contentDiv.innerHTML = `
                        <div style="color: var(--error); text-align: center;">
                            <i class="fas fa-times-circle" style="margin-right: 5px;"></i>
                            Error: ${error.message}
                        </div>
                    `;
                }
            }

            // 使用YOLO预处理后再用指定模型提取SMILES
            async function extractWithYOLOAndModel(fileId, imageIndex, extractionModel) {
                const resultDiv = document.getElementById(`smiles-result-${imageIndex}`);
                const contentDiv = document.getElementById(`smiles-content-${imageIndex}`);

                // 显示加载状态
                contentDiv.innerHTML = `
                    <div style="text-align: center; color: var(--warning);">
                        <i class="fas fa-spinner fa-spin" style="margin-right: 5px;"></i>
                        Step 1/2: Processing with YOLO...
                    </div>
                `;

                try {
                    // 第一步：使用YOLO处理图像
                    console.log('🔍 Step 1: Processing with YOLO for file:', fileId);
                    const yoloResponse = await fetch('/api/yolo/process', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            file_id: fileId
                        })
                    });

                    const yoloData = await yoloResponse.json();

                    if (!yoloData.success) {
                        throw new Error(yoloData.error || 'YOLO processing failed');
                    }

                    console.log('✅ YOLO processing completed, detected structures:', yoloData.count);

                    // 更新状态显示
                    contentDiv.innerHTML = `
                        <div style="text-align: center; color: var(--warning);">
                            <i class="fas fa-spinner fa-spin" style="margin-right: 5px;"></i>
                            Step 2/2: Extracting SMILES with ${extractionModel === 'modelA' ? 'DECIMER' : 'AIChemist'}...
                        </div>
                    `;

                    // 第二步：使用选定的模型提取SMILES（使用填充后的图像）
                    console.log('🔍 Step 2: Extracting SMILES with model:', extractionModel);
                    console.log('📊 YOLO filled image available:', !!yoloData.filled_image_base64);
                    const smilesResponse = await fetch('/api/smiles/extract-single', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            file_id: fileId,
                            model_type: extractionModel,
                            use_yolo: true,
                            filled_image_base64: yoloData.filled_image_base64
                        })
                    });

                    const smilesData = await smilesResponse.json();

                    if (smilesData.success && smilesData.smiles) {
                        // 构建数据库存储状态显示
                        let dbStatusHtml = '';
                        if (smilesData.stored_to_db) {
                            dbStatusHtml = `
                                <div style="color: var(--success); font-size: 11px; margin-top: 5px;">
                                    <i class="fas fa-database" style="margin-right: 5px;"></i>
                                    Stored to database (ID: ${smilesData.smiles_db_id || 'N/A'})
                                </div>
                            `;
                        }

                        // 获取方法名称
                        const methodName = `YOLO + ${extractionModel === 'modelB' ? 'AIChemist API' : 'DECIMER Model'}`;

                        contentDiv.innerHTML = `
                            <div style="color: var(--success); margin-bottom: 8px;">
                                <i class="fas fa-check-circle" style="margin-right: 5px;"></i>
                                <strong>SMILES Extracted Successfully!</strong>
                            </div>
                            <div style="background: var(--card-bg); padding: 8px; border-radius: 4px; margin-bottom: 8px;">
                                <strong>SMILES:</strong>
                                <span style="font-family: monospace; color: var(--accent); word-break: break-all;">${smilesData.smiles}</span>
                            </div>
                            <div style="font-size: 11px; color: var(--text-secondary);">
                                <div><strong>Method:</strong> <span style="color: #9b59b6; font-weight: bold;">${methodName}</span></div>
                                <div><strong>YOLO Structures Detected:</strong> ${yoloData.count}</div>
                                <div><strong>Confidence:</strong> ${smilesData.confidence !== undefined && smilesData.confidence !== null ? smilesData.confidence : 'N/A'}</div>
                                <div><strong>File ID:</strong> ${smilesData.file_id}</div>
                            </div>
                            ${dbStatusHtml}
                        `;
                    } else {
                        contentDiv.innerHTML = `
                            <div style="color: var(--warning); text-align: center;">
                                <i class="fas fa-exclamation-triangle" style="margin-right: 5px;"></i>
                                ${smilesData.error || 'SMILES extraction failed'}
                            </div>
                        `;
                    }
                } catch (error) {
                    console.error('❌ Error in YOLO + Model extraction:', error);
                    contentDiv.innerHTML = `
                        <div style="color: var(--error); text-align: center;">
                            <i class="fas fa-times-circle" style="margin-right: 5px;"></i>
                            Error: ${error.message}
                        </div>
                    `;
                }
            }

            async function loadImageFilenames() {
                const selectElement = document.getElementById('smiles-filename-filter');

                try {
                    const response = await fetch('/api/images/filenames');
                    const data = await response.json();

                    // 清空现有选项
                    selectElement.innerHTML = '<option value="">All Images</option>';

                    if (data.success && data.filenames.length > 0) {
                        data.filenames.forEach(filenameInfo => {
                            const option = document.createElement('option');
                            option.value = filenameInfo.filename;
                            option.textContent = `${filenameInfo.filename} (${filenameInfo.count} images)`;
                            selectElement.appendChild(option);
                        });

                        // 显示成功消息
                        const resultDiv = document.getElementById('smiles-result');
                        resultDiv.innerHTML = `
                            <div style="color: var(--success); padding: 15px; text-align: center;">
                                <i class="fas fa-check-circle" style="margin-right: 10px;"></i>
                                Loaded ${data.filenames.length} filenames successfully!
                            </div>
                        `;
                        resultDiv.classList.remove('hidden');

                        // 3秒后隐藏消息
                        setTimeout(() => {
                            resultDiv.classList.add('hidden');
                        }, 3000);
                    } else {
                        alert('No filenames found in database');
                    }
                } catch (error) {
                    alert('Error loading filenames: ' + error.message);
                }
            }

            function clearFilenameFilter() {
                const selectElement = document.getElementById('smiles-filename-filter');
                selectElement.value = '';

                // 隐藏结果区域
                const resultDiv = document.getElementById('smiles-result');
                resultDiv.classList.add('hidden');
            }
                // 添加缺少的函数定义





// 用户相关函数
async function loadUserInfo() {
    try {
        const response = await fetch('/api/user/info');
        const data = await response.json();

        if (data.success && data.user) {
            const user = data.user;
            const userDetails = document.getElementById('userDetails');
            const adminBtn = document.getElementById('adminBtn');

            // 显示用户信息
            let usageText = '';
            if (user.max_usage > 0) {
                usageText = ` (${user.usage_count}/${user.max_usage})`;
            } else {
                usageText = ` (${user.usage_count}/∞)`;
            }

            userDetails.innerHTML = `
                <i class="fas fa-user"></i> ${user.username}
                <span style="color: var(--accent);">${user.role === 'admin' ? '管理员' : '临时用户'}</span>
                ${usageText}
            `;

            // 如果是管理员，显示管理面板按钮
            if (user.role === 'admin') {
                adminBtn.style.display = 'flex';
            }
        }
    } catch (error) {
        console.error('加载用户信息失败:', error);
        // 如果获取用户信息失败，可能是未登录，跳转到登录页
        window.location.href = '/login';
    }
}

async function logout() {
    if (!confirm('确定要退出登录吗？')) {
        return;
    }

    try {
        const response = await fetch('/api/logout', {
            method: 'POST',
            headers: {
                'Content-Type': 'application/json',
            }
        });

        const data = await response.json();

        if (data.success) {
            window.location.href = data.redirect || '/login';
        } else {
            alert('退出登录失败，请重试');
        }
    } catch (error) {
        console.error('退出登录失败:', error);
        // 即使出错也跳转到登录页
        window.location.href = '/login';
    }
}

// 处理API请求错误，检查是否需要重新登录
function handleApiError(response) {
    if (response.status === 401) {
        alert('登录已过期或账号已达到使用限制，请重新登录');
        window.location.href = '/login';
        return true;
    }
    return false;
}

// 重写fetch函数以自动处理认证错误
const originalFetch = window.fetch;
window.fetch = async function(...args) {
    const response = await originalFetch.apply(this, args);

    if (response.status === 401) {
        const data = await response.json().catch(() => ({}));
        if (data.redirect) {
            alert('登录已过期或账号已达到使用限制，请重新登录');
            window.location.href = data.redirect;
        }
    }

    return response;
};
        </script>
    '''

    # 完整的HTML模板
    HTML_TEMPLATE = HTML_TEMPLATE_PART1 + HTML_TEMPLATE_PART2 + JAVASCRIPT_PART + '''
        </body>
        </html>
    '''

    # 登录页面模板
    LOGIN_TEMPLATE = '''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>长鑫 Research Platform - 登录</title>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                min-height: 100vh;
                display: flex;
                align-items: center;
                justify-content: center;
                overflow: hidden;
            }

            /* 视频背景 */
            .video-background {
                position: fixed;
                top: 0;
                left: 0;
                width: 100%;
                height: 100%;
                object-fit: cover;
                z-index: -2;
            }

            /* 背景遮罩层 */
            .background-overlay {
                position: fixed;
                top: 0;
                left: 0;
                width: 100%;
                height: 100%;
                background: rgba(0, 0, 0, 0.4);
                z-index: -1;
            }

            .login-container {
                background: rgba(255, 255, 255, 0.95);
                padding: 40px;
                border-radius: 20px;
                box-shadow: 0 15px 35px rgba(0, 0, 0, 0.3);
                width: 100%;
                max-width: 400px;
                backdrop-filter: blur(10px);
                z-index: 1;
            }

            .login-header {
                text-align: center;
                margin-bottom: 30px;
            }

            .login-header h1 {
                color: #333;
                font-size: 28px;
                margin-bottom: 10px;
            }

            .login-header p {
                color: #666;
                font-size: 14px;
            }

            .form-group {
                margin-bottom: 20px;
            }

            .form-group label {
                display: block;
                margin-bottom: 8px;
                color: #333;
                font-weight: 500;
            }

            .form-group input {
                width: 100%;
                padding: 12px 15px;
                border: 2px solid #e1e5e9;
                border-radius: 10px;
                font-size: 16px;
                transition: border-color 0.3s;
            }

            .form-group input:focus {
                outline: none;
                border-color: #667eea;
            }

            .login-btn {
                width: 100%;
                padding: 12px;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                border-radius: 10px;
                font-size: 16px;
                font-weight: 600;
                cursor: pointer;
                transition: transform 0.2s;
            }

            .login-btn:hover {
                transform: translateY(-2px);
            }

            .login-btn:disabled {
                opacity: 0.6;
                cursor: not-allowed;
                transform: none;
            }

            .message {
                margin-top: 15px;
                padding: 10px;
                border-radius: 8px;
                text-align: center;
                font-size: 14px;
            }

            .message.error {
                background-color: #fee;
                color: #c33;
                border: 1px solid #fcc;
            }

            .message.success {
                background-color: #efe;
                color: #363;
                border: 1px solid #cfc;
            }

            .admin-info {
                margin-top: 20px;
                padding: 15px;
                background-color: #f8f9fa;
                border-radius: 10px;
                border-left: 4px solid #667eea;
            }

            .admin-info h3 {
                color: #333;
                margin-bottom: 10px;
                font-size: 16px;
            }

            .admin-info p {
                color: #666;
                font-size: 14px;
                margin-bottom: 5px;
            }
        </style>
    </head>
    <body>
        <!-- 视频背景 -->
        <video class="video-background" autoplay muted loop playsinline>
            <source src="/static/video/长鑫存储.mp4" type="video/mp4">
            您的浏览器不支持视频播放
        </video>

        <!-- 背景遮罩层 -->
        <div class="background-overlay"></div>

        <div class="login-container">
            <div class="login-header">
                <h1>🔬 长鑫 Platform</h1>
                <p>物理材料AI研究平台</p>
            </div>

            <form id="loginForm">
                <div class="form-group">
                    <label for="username">用户名</label>
                    <input type="text" id="username" name="username" required>
                </div>

                <div class="form-group">
                    <label for="password">密码</label>
                    <input type="password" id="password" name="password" required>
                </div>

                <button type="submit" class="login-btn" id="loginBtn">登录</button>
            </form>

            <div id="message" class="message" style="display: none;"></div>


        </div>

        <script>
            document.getElementById('loginForm').addEventListener('submit', async function(e) {
                e.preventDefault();

                const username = document.getElementById('username').value.trim();
                const password = document.getElementById('password').value.trim();
                const loginBtn = document.getElementById('loginBtn');
                const messageDiv = document.getElementById('message');

                if (!username || !password) {
                    showMessage('请输入用户名和密码', 'error');
                    return;
                }

                loginBtn.disabled = true;
                loginBtn.textContent = '登录中...';

                try {
                    const response = await fetch('/api/login', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ username, password })
                    });

                    const data = await response.json();

                    if (data.success) {
                        showMessage('登录成功，正在跳转...', 'success');
                        setTimeout(() => {
                            window.location.href = data.redirect || '/';
                        }, 1000);
                    } else {
                        showMessage(data.message || '登录失败', 'error');
                    }
                } catch (error) {
                    showMessage('网络错误，请重试', 'error');
                } finally {
                    loginBtn.disabled = false;
                    loginBtn.textContent = '登录';
                }
            });

            function showMessage(text, type) {
                const messageDiv = document.getElementById('message');
                messageDiv.textContent = text;
                messageDiv.className = 'message ' + type;
                messageDiv.style.display = 'block';

                if (type === 'success') {
                    setTimeout(() => {
                        messageDiv.style.display = 'none';
                    }, 3000);
                }
            }
        </script>
    </body>
    </html>
    '''

    # 管理员页面模板
    ADMIN_TEMPLATE = '''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>长鑫 Platform - 管理员面板</title>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background: #f5f7fa;
                min-height: 100vh;
            }

            .header {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 20px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }

            .header h1 {
                margin-bottom: 10px;
            }

            .header-actions {
                display: flex;
                gap: 15px;
                align-items: center;
            }

            .btn {
                padding: 8px 16px;
                border: none;
                border-radius: 6px;
                cursor: pointer;
                font-size: 14px;
                transition: all 0.3s;
            }

            .btn-primary {
                background: rgba(255,255,255,0.2);
                color: black;
                border: 1px solid rgba(255,255,255,0.3);
            }

            .btn-primary:hover {
                background: rgba(255,255,255,0.3);
            }

            .btn-success {
                background: #28a745;
                color: white;
            }

            .btn-danger {
                background: #dc3545;
                color: white;
            }

            .btn-warning {
                background: #ffc107;
                color: #212529;
            }

            .container {
                max-width: 1200px;
                margin: 0 auto;
                padding: 30px 20px;
            }

            .card {
                background: white;
                border-radius: 12px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                margin-bottom: 30px;
                overflow: hidden;
            }

            .card-header {
                background: #f8f9fa;
                padding: 20px;
                border-bottom: 1px solid #e9ecef;
            }

            .card-header h2 {
                color: #333;
                margin-bottom: 5px;
            }

            .card-body {
                padding: 20px;
            }

            .form-group {
                margin-bottom: 20px;
            }

            .form-group label {
                display: block;
                margin-bottom: 8px;
                color: #333;
                font-weight: 500;
            }

            .form-group input {
                width: 100%;
                max-width: 300px;
                padding: 10px 12px;
                border: 2px solid #e1e5e9;
                border-radius: 6px;
                font-size: 14px;
            }

            .accounts-table {
                width: 100%;
                border-collapse: collapse;
                margin-top: 20px;
            }

            .accounts-table th,
            .accounts-table td {
                padding: 12px;
                text-align: left;
                border-bottom: 1px solid #e9ecef;
            }

            .accounts-table th {
                background: #f8f9fa;
                font-weight: 600;
                color: #333;
            }

            .status-badge {
                padding: 4px 8px;
                border-radius: 4px;
                font-size: 12px;
                font-weight: 500;
            }

            .status-active {
                background: #d4edda;
                color: #155724;
            }

            .status-inactive {
                background: #f8d7da;
                color: #721c24;
            }

            .message {
                margin: 15px 0;
                padding: 12px;
                border-radius: 6px;
                font-size: 14px;
            }

            .message.success {
                background: #d4edda;
                color: #155724;
                border: 1px solid #c3e6cb;
            }

            .message.error {
                background: #f8d7da;
                color: #721c24;
                border: 1px solid #f5c6cb;
            }

            .loading {
                opacity: 0.6;
                pointer-events: none;
            }
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🛠️ 管理员面板</h1>
            <div class="header-actions">
                <button class="btn btn-primary" onclick="window.location.href='/'">返回主页</button>
                <button class="btn btn-primary" onclick="logout()">退出登录</button>
            </div>
        </div>

        <div class="container">
            <!-- 生成临时账号 -->
            <div class="card">
                <div class="card-header">
                    <h2>🔑 生成临时账号</h2>
                    <p>为用户创建限制使用次数的临时账号</p>
                </div>
                <div class="card-body">
                    <form id="generateForm">
                        <div class="form-group">
                            <label for="maxUsage">使用次数限制</label>
                            <input type="number" id="maxUsage" name="maxUsage" value="10" min="1" max="100">
                        </div>
                        <button type="submit" class="btn btn-success">生成账号</button>
                    </form>
                    <div id="generateMessage" class="message" style="display: none;"></div>
                    <div id="newAccount" style="display: none; margin-top: 20px; padding: 15px; background: #e7f3ff; border-radius: 6px;">
                        <h3>🎉 新账号已生成</h3>
                        <p><strong>用户名:</strong> <span id="newUsername"></span></p>
                        <p><strong>密码:</strong> <span id="newPassword"></span></p>
                        <p><strong>使用次数:</strong> <span id="newMaxUsage"></span></p>
                        <p><strong>过期时间:</strong> <span id="newExpires"></span></p>
                    </div>
                </div>
            </div>

            <!-- 临时账号管理 -->
            <div class="card">
                <div class="card-header">
                    <h2>👥 临时账号管理</h2>
                    <p>查看和管理所有临时账号</p>
                </div>
                <div class="card-body">
                    <button class="btn btn-primary" onclick="loadAccounts()">刷新列表</button>
                    <div id="accountsContainer">
                        <table class="accounts-table" id="accountsTable">
                            <thead>
                                <tr>
                                    <th>用户名</th>
                                    <th>密码</th>
                                    <th>创建时间</th>
                                    <th>使用次数</th>
                                    <th>最大使用次数</th>
                                    <th>状态</th>
                                    <th>过期时间</th>
                                    <th>操作</th>
                                </tr>
                            </thead>
                            <tbody id="accountsTableBody">
                                <tr>
                                    <td colspan="8" style="text-align: center; color: #666;">加载中...</td>
                                </tr>
                            </tbody>
                        </table>
                    </div>
                </div>
            </div>
        </div>

        <script>
            // 页面加载时获取账号列表
            document.addEventListener('DOMContentLoaded', function() {
                loadAccounts();
            });

            // 生成临时账号
            document.getElementById('generateForm').addEventListener('submit', async function(e) {
                e.preventDefault();

                const maxUsage = document.getElementById('maxUsage').value;
                const messageDiv = document.getElementById('generateMessage');
                const form = e.target;

                form.classList.add('loading');

                try {
                    const response = await fetch('/api/admin/generate_temp_account', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ max_usage: parseInt(maxUsage) })
                    });

                    const data = await response.json();

                    if (data.success) {
                        showMessage(messageDiv, '账号生成成功！', 'success');
                        showNewAccount(data.account);
                        loadAccounts(); // 刷新账号列表
                    } else {
                        showMessage(messageDiv, data.message || '生成失败', 'error');
                    }
                } catch (error) {
                    showMessage(messageDiv, '网络错误，请重试', 'error');
                } finally {
                    form.classList.remove('loading');
                }
            });

            // 显示新生成的账号信息
            function showNewAccount(account) {
                document.getElementById('newUsername').textContent = account.username;
                document.getElementById('newPassword').textContent = account.password;
                document.getElementById('newMaxUsage').textContent = account.max_usage;
                document.getElementById('newExpires').textContent = new Date(account.expires_at).toLocaleString();
                document.getElementById('newAccount').style.display = 'block';
            }

            // 加载账号列表
            async function loadAccounts() {
                const tableBody = document.getElementById('accountsTableBody');
                tableBody.innerHTML = '<tr><td colspan="7" style="text-align: center; color: #666;">加载中...</td></tr>';

                try {
                    const response = await fetch('/api/admin/temp_accounts');
                    const data = await response.json();

                    if (data.success) {
                        displayAccounts(data.accounts);
                    } else {
                        tableBody.innerHTML = '<tr><td colspan="7" style="text-align: center; color: #dc3545;">加载失败</td></tr>';
                    }
                } catch (error) {
                    tableBody.innerHTML = '<tr><td colspan="7" style="text-align: center; color: #dc3545;">网络错误</td></tr>';
                }
            }

            // 显示账号列表
            function displayAccounts(accounts) {
                const tableBody = document.getElementById('accountsTableBody');

                if (accounts.length === 0) {
                    tableBody.innerHTML = '<tr><td colspan="8" style="text-align: center; color: #666;">暂无临时账号</td></tr>';
                    return;
                }

                tableBody.innerHTML = accounts.map(account => `
                    <tr>
                        <td>${account.username}</td>
                        <td style="font-family: monospace; color: var(--accent); font-weight: bold;">
                            ${account.plain_password || '密码已加密'}
                        </td>
                        <td>${new Date(account.created_at).toLocaleString()}</td>
                        <td>${account.usage_count}</td>
                        <td>${account.max_usage}</td>
                        <td>
                            <span class="status-badge ${account.is_active ? 'status-active' : 'status-inactive'}">
                                ${account.is_active ? '激活' : '停用'}
                            </span>
                        </td>
                        <td>${new Date(account.expires_at).toLocaleString()}</td>
                        <td>
                            ${account.is_active ?
                                `<button class="btn btn-warning" onclick="deactivateAccount('${account.username}')">停用</button>` :
                                ''
                            }
                            <button class="btn btn-danger" onclick="deleteAccount('${account.username}')">删除</button>
                        </td>
                    </tr>
                `).join('');
            }

            // 停用账号
            async function deactivateAccount(username) {
                if (!confirm(`确定要停用账号 ${username} 吗？`)) {
                    return;
                }

                try {
                    const response = await fetch('/api/admin/deactivate_account', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ username })
                    });

                    const data = await response.json();

                    if (data.success) {
                        alert('账号已停用');
                        loadAccounts();
                    } else {
                        alert(data.message || '操作失败');
                    }
                } catch (error) {
                    alert('网络错误，请重试');
                }
            }

            // 删除账号
            async function deleteAccount(username) {
                if (!confirm(`确定要删除账号 ${username} 吗？此操作不可恢复！`)) {
                    return;
                }

                try {
                    const response = await fetch('/api/admin/delete_account', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ username })
                    });

                    const data = await response.json();

                    if (data.success) {
                        alert('账号已删除');
                        loadAccounts();
                    } else {
                        alert(data.message || '操作失败');
                    }
                } catch (error) {
                    alert('网络错误，请重试');
                }
            }

            // 退出登录
            async function logout() {
                try {
                    const response = await fetch('/api/logout', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        }
                    });

                    const data = await response.json();

                    if (data.success) {
                        window.location.href = data.redirect || '/login';
                    }
                } catch (error) {
                    window.location.href = '/login';
                }
            }

            // 显示消息
            function showMessage(element, text, type) {
                element.textContent = text;
                element.className = 'message ' + type;
                element.style.display = 'block';

                setTimeout(() => {
                    element.style.display = 'none';
                }, 5000);
            }
        </script>
    </body>
    </html>
    '''

    # 静态文件路由 - 提供视频文件
    @app.route('/static/video/<filename>')
    def serve_video(filename):
        """提供视频文件"""
        from flask import send_from_directory
        video_path = os.path.join(STATIC_FOLDER, 'video')
        return send_from_directory(video_path, filename)

    # 登录相关路由
    @app.route('/login')
    def login():
        return render_template_string(LOGIN_TEMPLATE)

    @app.route('/api/login', methods=['POST'])
    def api_login():
        data = request.get_json()
        username = data.get('username', '').strip()
        password = data.get('password', '').strip()

        if not username or not password:
            return jsonify({"success": False, "message": "用户名和密码不能为空"}), 400

        user = user_manager.authenticate(username, password)
        if user:
            session['user'] = user
            session.permanent = True
            return jsonify({
                "success": True,
                "message": "登录成功",
                "user": user,
                "redirect": "/"
            })
        else:
            return jsonify({"success": False, "message": "用户名或密码错误，或账号已过期/达到使用限制"}), 401

    @app.route('/api/logout', methods=['POST'])
    def api_logout():
        session.pop('user', None)
        return jsonify({"success": True, "message": "已退出登录", "redirect": "/login"})

    @app.route('/api/user/info')
    @login_required
    def get_user_info():
        return jsonify({"success": True, "user": session['user']})

    # 管理员功能路由
    @app.route('/admin')
    @admin_required
    def admin_panel():
        return render_template_string(ADMIN_TEMPLATE)

    @app.route('/api/admin/generate_temp_account', methods=['POST'])
    @admin_required
    def generate_temp_account():
        data = request.get_json()
        max_usage = data.get('max_usage', 10)

        try:
            max_usage = int(max_usage)
            if max_usage <= 0:
                max_usage = 10
        except:
            max_usage = 10

        account = user_manager.generate_temp_account(max_usage)
        return jsonify({"success": True, "account": account})

    @app.route('/api/admin/temp_accounts')
    @admin_required
    def list_temp_accounts():
        accounts = user_manager.list_temp_accounts_with_passwords()
        return jsonify({"success": True, "accounts": accounts})

    @app.route('/api/admin/deactivate_account', methods=['POST'])
    @admin_required
    def deactivate_account():
        data = request.get_json()
        username = data.get('username', '').strip()

        if not username:
            return jsonify({"success": False, "message": "用户名不能为空"}), 400

        if user_manager.deactivate_account(username):
            return jsonify({"success": True, "message": f"账号 {username} 已停用"})
        else:
            return jsonify({"success": False, "message": "停用失败，账号不存在或无法停用"}), 400

    @app.route('/api/admin/delete_account', methods=['POST'])
    @admin_required
    def delete_account():
        data = request.get_json()
        username = data.get('username', '').strip()

        if not username:
            return jsonify({"success": False, "message": "用户名不能为空"}), 400

        if user_manager.delete_account(username):
            return jsonify({"success": True, "message": f"账号 {username} 已删除"})
        else:
            return jsonify({"success": False, "message": "删除失败，账号不存在或无法删除"}), 400

    # API路由
    @app.route('/')
    @login_required
    @track_usage
    def index():
        return render_template_string(HTML_TEMPLATE)

    @app.route('/api/database/stats')
    def get_database_stats():
        filename = request.args.get('filename')  # 获取文件筛选参数
        if filename:
            stats = db_manager.get_database_stats_by_file(filename)
        else:
            stats = db_manager.get_database_stats()
        return jsonify(stats)

    @app.route('/api/documents/<document_id>', methods=['DELETE'])
    def delete_document(document_id):
        """删除单个文档"""
        try:
            success = db_manager.delete_document(document_id)
            if success:
                return jsonify({'message': '文档删除成功', 'success': True})
            else:
                return jsonify({'error': '文档删除失败', 'success': False}), 500
        except Exception as e:
            print(f"❌ Error deleting document: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/database/clear', methods=['POST'])
    def clear_database():
        """清空所有数据库数据"""
        try:
            success = db_manager.clear_all_data()
            if success:
                return jsonify({'message': '数据库已清空', 'success': True})
            else:
                return jsonify({'error': '清空数据库失败', 'success': False}), 500
        except Exception as e:
            print(f"❌ Error clearing database: {e}")
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/documents')
    @login_required
    def get_documents():
        """获取所有文档和片段文件，支持文件名筛选"""
        try:
            # 检查是否有文件名筛选参数
            filename_filter = request.args.get('filename')

            if filename_filter:
                # 筛选特定文件：只返回该文件及其切割后的片段
                documents = db_manager.get_all_documents()

                # 找到匹配的原始文档
                target_doc = None
                for doc in documents:
                    if doc.get('filename') == filename_filter and not doc.get('is_segment_file', False):
                        target_doc = doc
                        break

                if not target_doc:
                    return jsonify([])  # 如果没找到匹配的文档，返回空列表

                # 获取该文档的片段文件
                expanded_documents = [target_doc]  # 先添加原始文档

                # 添加该文档的片段文件
                segment_files = db_manager.get_segment_files(target_doc['id'])
                for segment_file in segment_files:
                    # 根据文件类型确定显示的文件类型
                    display_file_type = get_display_file_type(segment_file['file_type'])

                    file_entry = {
                        'id': segment_file['id'],
                        'filename': segment_file['virtual_filename'],
                        'file_type': display_file_type,
                        'upload_time': segment_file['created_time'],
                        'file_size': 'Generated',
                        'page_count': 'N/A',
                        'language': 'Generated',
                        'is_segment_file': True,
                        'parent_document_id': target_doc['id'],
                        'parent_filename': target_doc['filename'],
                        'segment_file_type': segment_file['file_type']
                    }
                    expanded_documents.append(file_entry)

            else:
                # 显示所有文档和片段文件
                documents = db_manager.get_all_documents()
                expanded_documents = []

                for doc in documents:
                    # 添加原始文档
                    expanded_documents.append(doc)

                    # 获取该文档的片段文件
                    segment_files = db_manager.get_segment_files(doc['id'])

                    # 将每个片段文件作为独立条目添加
                    for segment_file in segment_files:
                        # 根据文件类型确定显示的文件类型
                        display_file_type = get_display_file_type(segment_file['file_type'])

                        file_entry = {
                            'id': segment_file['id'],
                            'filename': segment_file['virtual_filename'],
                            'file_type': display_file_type,
                            'upload_time': segment_file['created_time'],
                            'file_size': 'Generated',
                            'page_count': 'N/A',
                            'language': 'Generated',
                            'is_segment_file': True,
                            'parent_document_id': doc['id'],
                            'parent_filename': doc['filename'],
                            'segment_file_type': segment_file['file_type']
                        }
                        expanded_documents.append(file_entry)

            return jsonify(expanded_documents)
        except Exception as e:
            print(f"Error in get_documents: {e}")
            # 如果出错，返回基本文档列表
            documents = db_manager.get_all_documents()
            return jsonify(documents)





    @app.route('/api/physicochemical/extract', methods=['POST'])
    def extract_physicochemical_properties():
        """提取理化性质API端点"""
        try:
            data = request.get_json()
            document_ids = data.get('document_ids', [])

            if not document_ids:
                return jsonify({
                    'success': False,
                    'error': 'No document IDs provided'
                }), 400

            # 获取理化性质提取器
            from physicochemical_extractor import create_physicochemical_extractor
            extractor = create_physicochemical_extractor(db_manager)

            results = []
            total_molecules = 0
            total_stored = 0

            for doc_id in document_ids:
                doc = db_manager.get_document_by_id(doc_id)
                if doc:
                    text_content = doc.get('text_content', '')
                    if text_content:
                        result = extractor.extract_physicochemical_properties(text_content, doc_id)
                        results.append({
                            'document_id': doc_id,
                            'filename': doc.get('filename', ''),
                            'result': result
                        })

                        if result.get('status') == 'success':
                            total_molecules += result.get('molecules_extracted', 0)
                            total_stored += result.get('molecules_stored', 0)

            return jsonify({
                'success': True,
                'total_molecules_extracted': total_molecules,
                'total_molecules_stored': total_stored,
                'documents_processed': len(results),
                'results': results
            })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500

    @app.route('/api/documents/upload', methods=['POST'])
    @login_required
    @track_usage
    def upload_documents():
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'error': 'No files selected'}), 400

        results = []
        total_molecules = 0

        for file in files:
            if file.filename == '':
                continue

            # 模拟文档处理
            filename = file.filename
            file_type = filename.split('.')[-1].upper()
            file_size = len(file.read())
            file.seek(0)  # 重置文件指针

            # 模拟文本内容
            text_content = f"Processed content from {filename}. This document contains pharmaceutical research data, molecular structures, and bioactivity information relevant to drug discovery and development."

            # 存储到数据库
            document_id = db_manager.store_document(filename, file_type, text_content, file_size, 1)

            # 获取相关分子数量
            molecules = db_manager.get_molecules_by_document(document_id)
            total_molecules += len(molecules)

            results.append({
                'document_id': document_id,
                'filename': filename,
                'molecules_count': len(molecules)
            })

        return jsonify({
            'status': 'success',
            'processed_files': len(results),
            'total_molecules': total_molecules,
            'results': results
        })





    @app.route('/api/workflow/execute', methods=['POST'])
    def execute_workflow_legacy():
        data = request.get_json()
        workflow_id = data.get('workflow_id')
        document_ids = data.get('document_ids', [])

        if not workflow_id:
            return jsonify({'error': 'No workflow selected'}), 400

        if not document_ids:
            return jsonify({'error': 'No documents selected'}), 400

        # 获取文档信息
        documents = []
        for doc_id in document_ids:
            document = db_manager.get_document_by_id(int(doc_id))
            if document:
                documents.append(document)

        # 执行工作流
        result = workflow_manager.execute_workflow(workflow_id, documents)

        return jsonify(result)

    @app.route('/api/documents/process', methods=['POST'])
    @login_required
    @track_usage
    def process_documents():
        """处理文档上传和PDF分割的统一接口"""
        try:
            if 'files' not in request.files:
                return jsonify({'success': False, 'message': 'No files provided'}), 400

            files = request.files.getlist('files')
            if not files or all(f.filename == '' for f in files):
                return jsonify({'success': False, 'message': 'No valid files provided'}), 400

            # 获取PDF处理选项
            enable_ocr = request.form.get('enable_ocr', 'false').lower() == 'true'
            enable_formula = request.form.get('enable_formula', 'false').lower() == 'true'
            enable_table = request.form.get('enable_table', 'false').lower() == 'true'

            processed_count = 0
            pdf_results = []
            total_chemical_structures = 0

            # 检查文件大小
            max_file_size = 500 * 1024 * 1024  # 500MB
            for file in files:
                if file.filename == '':
                    continue

                # 检查文件大小
                file.seek(0, 2)  # 移动到文件末尾
                file_size = file.tell()
                file.seek(0)  # 重置到开头

                if file_size > max_file_size:
                    return jsonify({
                        'success': False,
                        'message': f'File {file.filename} is too large ({file_size / 1024 / 1024:.1f}MB). Maximum size is 500MB.'
                    }), 413

            for file in files:
                if file.filename == '':
                    continue

                # 保存文件 - 使用支持中文的安全文件名处理
                filename = safe_filename_with_chinese(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                # 检查是否为PDF文件
                if filename.lower().endswith('.pdf'):
                    # PDF文件使用分割处理
                    try:
                        # 这里调用PDF分割逻辑
                        segments,image_dir = process_pdf_file(file_path, enable_ocr, enable_formula, enable_table)

                        # 生成处理结果文件
                        try:
                            with open(file_path, 'rb') as f:
                                file_size = len(f.read())
                        except:
                            file_size = 0

                        # 尝试加载Magic PDF生成的完整结果
                        magic_pdf_results = load_magic_pdf_results(filename)

                        if magic_pdf_results and 'markdown' in magic_pdf_results:
                            # 使用Magic PDF生成的完整Markdown内容
                            print(f"🎯 Using Magic PDF results for {filename}")
                            processing_results = {
                                'markdown': magic_pdf_results['markdown'],
                                'content_list': magic_pdf_results.get('content_list', {}),
                                'middle_json': magic_pdf_results.get('middle_json', {})
                            }
                        else:
                            # 回退到基于segments的生成
                            print(f"📄 Using segments-based generation for {filename}")
                            processing_results = generate_document_processing_results(filename, segments)

                        # 存储主文档和处理结果到数据库
                        document_id = db_manager.store_processed_document(
                            filename=filename,
                            file_type='PDF',
                            file_size=file_size,
                            page_count=len(set(seg.get('page', 1) for seg in segments)),
                            segments=segments,
                            markdown_content=processing_results['markdown'],
                            content_list=processing_results['content_list'],
                            middle_json=processing_results['middle_json']
                        )

                        # 处理化学结构提取
                        chemical_count = 0
                        images = os.listdir(image_dir)
                        try:
                            document_name = os.path.splitext(filename)[0]  # 去掉扩展名
                            for image in images:
                                chemical_count += chemical_processor.process_pdf_for_chemical_structures(
                                    os.path.join(image_dir,image), str(document_id), document_name
                            )
                        except Exception as e:
                            print(f"❌ Chemical structure processing failed for {filename}: {e}")

                        pdf_results.append({
                            'filename': filename,
                            'segments': len(segments),
                            'chemical_structures': chemical_count,
                            'document_id': str(document_id)
                        })

                    except Exception as e:
                        print(f"Error processing PDF {filename}: {str(e)}")
                        # 如果PDF处理失败，作为普通文档处理
                        try:
                            with open(file_path, 'rb') as f:
                                file_size = len(f.read())
                        except:
                            file_size = 0

                        document_id = db_manager.store_document(
                            filename=filename,
                            file_type='PDF',
                            text_content=f"PDF file: {filename} (processing failed)",
                            file_size=file_size,
                            page_count=1
                        )

                else:
                    # 非PDF文件的常规处理
                    try:
                        if filename.lower().endswith(('.txt', '.md')):
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                            file_type = 'TXT'
                        elif filename.lower().endswith('.docx'):
                            # 这里可以添加DOCX处理逻辑
                            content = f"DOCX file: {filename}"
                            file_type = 'DOCX'
                        else:
                            content = f"File: {filename}"
                            file_type = 'OTHER'

                        # 获取文件大小
                        try:
                            with open(file_path, 'rb') as f:
                                file_size = len(f.read())
                        except:
                            file_size = 0

                        document_id = db_manager.store_document(
                            filename=filename,
                            file_type=file_type,
                            text_content=content,
                            file_size=file_size,
                            page_count=1
                        )

                    except Exception as e:
                        print(f"Error processing file {filename}: {str(e)}")
                        try:
                            with open(file_path, 'rb') as f:
                                file_size = len(f.read())
                        except:
                            file_size = 0

                        document_id = db_manager.store_document(
                            filename=filename,
                            file_type='OTHER',
                            text_content=f"File: {filename} (processing failed)",
                            file_size=file_size,
                            page_count=1
                        )

                processed_count += 1

            # 计算总的化学结构数量
            total_chemical_structures = sum(result.get('chemical_structures', 0) for result in pdf_results)

            response_data = {
                'success': True,
                'processed_count': processed_count,
                'total_chemical_structures': total_chemical_structures,
                'message': f'Successfully processed {processed_count} files'
            }

            if pdf_results:
                response_data['pdf_results'] = pdf_results
                if total_chemical_structures > 0:
                    response_data['message'] += f' and extracted {total_chemical_structures} chemical structures'

            return jsonify(response_data)

        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"❌ Document processing error: {error_details}")

            # 根据错误类型返回不同的状态码
            if "413" in str(e) or "too large" in str(e).lower():
                return jsonify({
                    'success': False,
                    'message': f'File too large: {str(e)}'
                }), 413
            elif "timeout" in str(e).lower():
                return jsonify({
                    'success': False,
                    'message': f'Processing timeout: {str(e)}'
                }), 408
            else:
                return jsonify({
                    'success': False,
                    'message': f'Error processing documents: {str(e)}',
                    'error_type': type(e).__name__
                }), 500

    def process_pdf_file(file_path, enable_ocr=True, enable_formula=True, enable_table=False):
        """处理PDF文件并返回分割结果 - 使用GPU加速的PDF处理器"""
        try:
            # 首先尝试使用优化的PDF处理器（支持GPU加速）
            try:
                print(f"🚀 Using GPU-accelerated PDF processor for {file_path}")

                # 创建进度回调函数
                def progress_callback(progress, message):
                    print(f"📊 Progress: {progress}% - {message}")

                # 使用优化的PDF处理器
                result = pdf_processor.auto_seg(file_path, progress_callback=progress_callback)

                if result.get('success'):
                    print(f"✅ GPU-accelerated processing completed: {result.get('processing_mode')}")
                    print(f"⏱️ Processing time: {result.get('processing_time')}")

                    # 从处理结果中提取segments
                    # 这里需要解析markdown或content_list来生成segments
                    segments = []

                    # 如果有content_list，解析它
                    if 'content_list' in result.get('output_files', {}):
                        try:
                            import json
                            content_list_path = os.path.join(result.get('output_dir', ''), result['output_files']['content_list'])
                            if os.path.exists(content_list_path):
                                with open(content_list_path, 'r', encoding='utf-8') as f:
                                    content_list = json.load(f)

                                # 解析content_list生成segments
                                for item in content_list:
                                    if isinstance(item, dict):
                                        segments.append({
                                            'text': item.get('text', ''),
                                            'type': item.get('type', 'text'),
                                            'page': item.get('page', 1),
                                            'bbox': item.get('bbox', []),
                                            'confidence': item.get('confidence', 0.95),
                                            'block_id': item.get('block_id', f"gpu_block_{len(segments)}")
                                        })
                        except Exception as e:
                            print(f"⚠️ Error parsing content_list: {e}")

                    # 如果没有segments，从markdown生成基本segments
                    if not segments and 'markdown' in result.get('output_files', {}):
                        try:
                            md_path = os.path.join(result.get('output_dir', ''), result['output_files']['markdown'])
                            if os.path.exists(md_path):
                                with open(md_path, 'r', encoding='utf-8') as f:
                                    markdown_content = f.read()

                                # 简单解析markdown生成segments
                                lines = markdown_content.split('\n')
                                page_num = 1
                                for i, line in enumerate(lines):
                                    if line.strip():
                                        text_type = "text"
                                        if line.startswith('# '):
                                            text_type = "title"
                                        elif line.startswith('## '):
                                            text_type = "heading"
                                        elif line.startswith('### '):
                                            text_type = "subheading"

                                        segments.append({
                                            'text': line.strip(),
                                            'type': text_type,
                                            'page': page_num,
                                            'bbox': [],
                                            'confidence': 0.95,
                                            'block_id': f"gpu_md_block_{i}"
                                        })
                        except Exception as e:
                            print(f"⚠️ Error parsing markdown: {e}")

                    if segments:
                        print(f"✅ GPU processing extracted {len(segments)} segments")
                        return segments,result['image_dir']
                    else:
                        print("⚠️ No segments extracted from GPU processing, falling back to PyMuPDF")

            except Exception as e:
                error_msg = str(e)
                print(f"⚠️ GPU-accelerated processing failed: {error_msg}")

                # 检查是否是PDF页数为0的错误
                if "total_page=0" in error_msg or "页数为0" in error_msg or "meta_scan need_drop" in error_msg:
                    print("❌ PDF文件无效或损坏，无法处理")
                    return [{
                        'text': f"PDF文件处理失败：文件无效或损坏\n\n错误详情：{error_msg}\n\n可能原因：\n1. PDF文件损坏或不完整\n2. PDF文件为空（0页）\n3. PDF文件格式不被支持\n\n建议：\n- 检查PDF文件是否能正常打开\n- 尝试重新下载或获取PDF文件\n- 确认文件格式正确",
                        'page': 1,
                        'type': 'error',
                        'confidence': 0.0,
                        'error_type': 'invalid_pdf',
                        'error_details': error_msg
                    }], "uploads/images"

                print("🔄 Falling back to PyMuPDF...")

                # 备用方案：使用PyMuPDF获取更详细的信息


            return [], "uploads/images"

        except Exception as e:
            print(f"Error in process_pdf_file: {str(e)}")
            return [{
                'text': f"PDF processing failed: {str(e)}\nFile path: {file_path}\nPlease check if the file is a valid PDF document.",
                'page': 1,
                'type': 'error',
                'confidence': 0.0,
                'error_details': str(e)
            }], "uploads/images"

    @app.route('/api/documents/<int:document_id>/images')
    def get_document_images(document_id):
        """获取文档中的分子图片"""
        document = db_manager.get_document_by_id(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404

        # 模拟从文档中提取的分子图片
        import random
        image_count = random.randint(2, 6)

        # 生成模拟的分子结构图片URL（使用占位符）
        images = []
        for i in range(image_count):
            images.append({
                'url': f'https://via.placeholder.com/200x150/1A1A2E/00D4FF?text=Molecule+{i+1}',
                'type': f'Molecular Structure {i+1}',
                'index': i
            })

        return jsonify({
            'document_id': document_id,
            'document_name': document['filename'],
            'images': images,
            'total_images': len(images)
        })

    @app.route('/api/images/segmented', methods=['GET'])
    def get_segmented_images():
        """获取切割后的化学结构图像"""
        try:
            source_file = request.args.get('source_file')
            filename = request.args.get('filename')

            segmentation_processor = get_segmentation_processor()
            images = segmentation_processor.get_stored_structures(source_file, filename=filename)

            # 获取图像数据并转换为base64
            image_list = []
            for img_info in images:
                try:
                    image_data = segmentation_processor.get_image_by_id(img_info['file_id'])
                    if image_data:
                        import base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        image_list.append({
                            'file_id': img_info['file_id'],
                            'filename': img_info['filename'],
                            'metadata': img_info['metadata'],
                            'upload_date': img_info['upload_date'].isoformat() if img_info['upload_date'] else None,
                            'image_data': image_base64
                        })
                except Exception as e:
                    print(f"Error processing image {img_info['file_id']}: {e}")
                    continue

            return jsonify({
                'success': True,
                'images': image_list
            })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Failed to retrieve segmented images: {str(e)}',
                'images': []
            }), 500

    @app.route('/api/images/filenames', methods=['GET'])
    def get_image_filenames():
        """获取数据库中的图像文件名列表"""
        try:
            segmentation_processor = get_segmentation_processor()

            # 获取所有图像的文件名
            images = segmentation_processor.get_stored_structures()

            # 统计每个文件名的图像数量
            filename_counts = {}
            for img_info in images:
                filename = img_info.get('filename', 'Unknown')
                if filename in filename_counts:
                    filename_counts[filename] += 1
                else:
                    filename_counts[filename] = 1

            # 转换为列表格式
            filenames = [
                {'filename': filename, 'count': count}
                for filename, count in filename_counts.items()
            ]

            return jsonify({
                'success': True,
                'filenames': filenames
            })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Failed to retrieve image filenames: {str(e)}',
                'filenames': []
            }), 500
        """获取文档中的分子图片"""
        document = db_manager.get_document_by_id(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404

        # 模拟从文档中提取的分子图片
        import random
        image_count = random.randint(2, 6)

        # 生成模拟的分子结构图片URL（使用占位符）
        images = []
        for i in range(image_count):
            images.append({
                'url': f'https://via.placeholder.com/200x150/1A1A2E/00D4FF?text=Molecule+{i+1}',
                'type': f'Molecular Structure {i+1}',
                'index': i
            })

        return jsonify({
            'document_id': document_id,
            'document_name': document['filename'],
            'images': images,
            'total_images': len(images)
        })

    @app.route('/api/smiles/extract', methods=['POST'])
    def extract_smiles():
        """从数据库图像中批量提取SMILES并存储到数据库"""
        try:
            filename = request.form.get('filename')

            # 如果filename为空字符串，转换为None以处理"All Images"的情况
            if filename == '':
                filename = None

            smiles_extractor = get_smiles_extractor()

            # 使用MongoDB批量提取方法，自动存储到数据库
            result = smiles_extractor.extract_smiles_from_mongodb_images(filename=filename)

            return jsonify(result)

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'SMILES extraction failed: {str(e)}',
                'results': []
            }), 500
            filename = request.args.get('filename')

            segmentation_processor = get_segmentation_processor()
            images = segmentation_processor.get_stored_structures(source_file, filename=filename)

            # 获取图像数据并转换为base64
            image_list = []
            for img_info in images:
                try:
                    image_data = segmentation_processor.get_image_by_id(img_info['file_id'])
                    if image_data:
                        import base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        image_list.append({
                            'file_id': img_info['file_id'],
                            'filename': img_info['filename'],
                            'metadata': img_info['metadata'],
                            'upload_date': img_info['upload_date'].isoformat() if img_info['upload_date'] else None,
                            'image_data': image_base64
                        })
                except Exception as e:
                    print(f"Error processing image {img_info['file_id']}: {e}")
                    continue

            return jsonify({
                'success': True,
                'images': image_list
            })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Failed to retrieve segmented images: {str(e)}',
                'images': []
            }), 500



    @app.route('/api/images/<file_id>', methods=['GET'])
    def get_image_by_id(file_id):
        """通过文件ID获取图像数据"""
        try:
            segmentation_processor = get_segmentation_processor()
            image_data = segmentation_processor.get_image_by_id(file_id)

            if image_data:
                import base64
                image_base64 = base64.b64encode(image_data).decode('utf-8')
                return jsonify({
                    'success': True,
                    'image_data': image_base64,
                    'file_id': file_id
                })
            else:
                return jsonify({
                    'success': False,
                    'error': f'Image not found for ID: {file_id}'
                }), 404

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Failed to retrieve image: {str(e)}'
            }), 500

    @app.route('/api/images/<file_id>/view', methods=['GET'])
    def view_image_by_id(file_id):
        """通过文件ID直接查看图像（返回图像文件而非JSON）"""
        try:
            segmentation_processor = get_segmentation_processor()
            image_data = segmentation_processor.get_image_by_id(file_id)

            if image_data:
                from flask import Response
                import io

                # 创建图像响应
                response = Response(
                    image_data,
                    mimetype='image/png',  # 假设是PNG格式
                    headers={
                        'Content-Disposition': f'inline; filename="chemical_structure_{file_id}.png"',
                        'Cache-Control': 'public, max-age=3600'  # 缓存1小时
                    }
                )
                return response
            else:
                return "Image not found", 404

        except Exception as e:
            return f"Failed to retrieve image: {str(e)}", 500


    @app.route('/api/yolo/process', methods=['POST'])
    def process_with_yolo():
        """使用YOLO模型处理图像，检测化学结构"""
        try:
            data = request.get_json()
            file_id = data.get('file_id')

            if not file_id:
                return jsonify({
                    'success': False,
                    'error': 'file_id is required'
                }), 400

            print(f"🔍 YOLO processing - file_id: {file_id}")

            # 导入YOLO处理器
            try:
                from yolo_processor import get_yolo_processor
                # 强制重新加载以确保使用最新配置
                yolo_processor = get_yolo_processor(force_reload=True)
            except Exception as e:
                print(f"❌ Failed to load YOLO processor: {e}")
                return jsonify({
                    'success': False,
                    'error': f'Failed to load YOLO processor: {str(e)}'
                }), 500

            # 使用YOLO处理器处理图像（填充模式）
            result = yolo_processor.process_from_file_id(file_id, db_manager, fill_with_ce=True)

            if result['success']:
                print(f"✅ YOLO processing successful, detected {result['count']} structures")

                # 返回填充后的图像
                return jsonify({
                    'success': True,
                    'count': result['count'],
                    'filled_image_base64': result.get('filled_image_base64'),
                    'detections': [
                        {
                            'index': d['index'],
                            'confidence': d['confidence'],
                            'bbox': d['xyxy'],
                            'class': d['class']
                        }
                        for d in result.get('detections', [])
                    ],
                    'pad_info': result.get('pad_info')
                })
            else:
                print(f"❌ YOLO processing failed: {result.get('error')}")
                return jsonify({
                    'success': False,
                    'error': result.get('error', 'YOLO processing failed')
                }), 500

        except Exception as e:
            print(f"❌ Error in YOLO processing endpoint: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': str(e)
            }), 500


    @app.route('/api/smiles/extract-single', methods=['POST'])
    def extract_smiles_single():
        """从单个图像中提取SMILES并存储到数据库"""
        try:
            data = request.get_json()
            file_id = data.get('file_id')
            model_type = data.get('model_type', 'modelA')  # 获取模型类型，默认为 modelA
            use_yolo = data.get('use_yolo', False)  # 是否使用YOLO预处理
            yolo_structures = data.get('yolo_structures', None)  # YOLO检测到的结构

            if not file_id:
                return jsonify({
                    'success': False,
                    'error': 'file_id is required'
                }), 400

            print(f"🔍 extract_smiles_single - file_id: {file_id}, model_type: {model_type}, use_yolo: {use_yolo}")

            # 如果使用YOLO预处理，使用填充后的图像
            filled_image_base64 = data.get('filled_image_base64', None)

            if use_yolo and filled_image_base64:
                print(f"🔍 Using YOLO filled image (with Ce markers)")
                # 将base64转换为图像数据
                import base64
                image_data = base64.b64decode(filled_image_base64)
                print(f"✅ Using YOLO filled image, size: {len(image_data)} bytes")
            else:
                # 如果YOLO没有检测到结构，回退到正常流程
                if use_yolo:
                    print(f"⚠️ YOLO enabled but no filled image provided, falling back to original image")
                    use_yolo = False  # 重置use_yolo标志，使用原始图像
                # 正常流程：从数据库获取图像
                image_data = None

            # 如果选择的是 Model B (AIChemist)
            if model_type == 'modelB':
                print("🔍 Using AIChemist API for SMILES extraction")
                import requests
                import base64
                import urllib3
                from bson import ObjectId

                # 获取smiles_extractor实例(用于保存RDKit图像)
                smiles_extractor = get_smiles_extractor()

                # 禁用 SSL 警告（因为 AIChemist API 的 SSL 证书已过期）
                urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

                try:
                    # 如果没有使用YOLO预处理，从 GridFS 获取图像数据和文件信息
                    if not use_yolo or image_data is None:
                        print(f"📸 Getting image data for file_id: {file_id}")
                        image_data = db_manager.get_image_by_id(file_id)
                        if not image_data:
                            print(f"❌ Image not found for file_id: {file_id}")
                            return jsonify({
                                'success': False,
                                'error': 'Image not found'
                            }), 404

                        print(f"✅ Image data retrieved, size: {len(image_data)} bytes")

                    # 获取文件元数据
                    source_file = None
                    filename = None
                    segment_number = None
                    image_type = None

                    if use_yolo:
                        # YOLO预处理的图像使用特殊标记
                        filename = f"yolo_processed_{file_id}"
                        source_file = f"yolo_processed_{file_id}"
                        image_type = 'yolo_chemical_structure'
                        print(f"📄 Using YOLO preprocessed image metadata")
                    else:
                        try:
                            grid_file = db_manager.medicinal_image_fs.get(ObjectId(file_id))
                            filename = grid_file.filename

                            # 从元数据中获取信息
                            if grid_file.metadata:
                                source_file = grid_file.metadata.get('source_file')
                                segment_number = grid_file.metadata.get('segment_number')
                                image_type = grid_file.metadata.get('image_type', 'chemical_structure')
                                print(f"📄 File metadata - filename: {filename}, source_file: {source_file}, segment: {segment_number}")
                            else:
                                image_type = 'chemical_structure'
                                print(f"📄 File metadata - filename: {filename}, no metadata found")

                            # 如果 source_file 为 None，尝试从文件名推导
                            if not source_file and filename:
                                # 从文件名中提取源文件信息（例如：document_name_chemical_structure_1.png）
                                source_file = filename
                                print(f"📄 source_file derived from filename: {source_file}")
                        except Exception as e:
                            print(f"⚠️ Could not retrieve file metadata: {e}")
                            filename = f"image_{file_id}"
                            image_type = 'chemical_structure'
                            # 如果获取失败，使用 file_id 作为 source_file
                            source_file = f"image_{file_id}"

                    # 将图像转换为 Base64
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
                    print(f"✅ Image converted to Base64, length: {len(image_base64)}")

                    # 调用 AIChemist API
                    aichemist_url = 'https://api-ocsr.alchemist.iresearch.net.cn/ocsr/'
                    headers = {
                        'X-API-Version': '1.0',
                        'Content-Type': 'application/json'
                    }
                    payload = {
                        'imageBase64': image_base64
                    }

                    print(f"🔗 Calling AIChemist API at {aichemist_url}")
                    # 禁用 SSL 证书验证，因为 AIChemist API 的 SSL 证书已过期
                    response = requests.post(aichemist_url, json=payload, headers=headers, timeout=30, verify=False)
                    print(f"📡 AIChemist API response status: {response.status_code}")

                    if response.status_code == 200:
                        api_result = response.json()
                        print(f"📊 AIChemist API result: {api_result}")
                        print(f"📊 API result type: {type(api_result)}")
                        print(f"📊 API result keys: {api_result.keys() if isinstance(api_result, dict) else 'Not a dict'}")

                        # 尝试多种可能的字段名称
                        smiles = api_result.get('smiles', '') or api_result.get('SMILES', '') or api_result.get('smi', '')
                        molblock = api_result.get('molblock', '') or api_result.get('mol', '') or api_result.get('MOL', '')
                        confidence = api_result.get('confidence', 0) or api_result.get('score', 0) or api_result.get('confidence_score', 0)

                        # 如果 API 返回的是嵌套结构，尝试提取
                        if isinstance(api_result, dict) and 'data' in api_result:
                            data = api_result.get('data', {})
                            if isinstance(data, dict):
                                smiles = smiles or data.get('smiles', '') or data.get('SMILES', '')
                                molblock = molblock or data.get('molblock', '') or data.get('mol', '')
                                confidence = confidence or data.get('confidence', 0)

                        print(f"📋 SMILES: {smiles}, Molblock length: {len(molblock) if molblock else 0}, Confidence: {confidence}")

                        # 如果返回的是 molblock 格式，使用 RDKit 转换为 SMILES
                        if molblock and not smiles:
                            try:
                                print(f"🔄 Converting molblock to SMILES...")
                                print(f"🔄 Molblock preview: {molblock[:200] if len(molblock) > 200 else molblock}")
                                from rdkit import Chem
                                mol = Chem.MolFromMolBlock(molblock)
                                if mol is not None:
                                    # 使用 kekuleSmiles=True 参数，避免芳香化
                                    smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                    print(f"✅ Converted molblock to SMILES (kekuleSmiles=True): {smiles}")
                                else:
                                    print(f"❌ Failed to parse molblock")
                                    # 尝试使用 MolFromMolBlock 的其他参数
                                    try:
                                        mol = Chem.MolFromMolBlock(molblock, removeHs=False)
                                        if mol is not None:
                                            # 使用 kekuleSmiles=True 参数，避免芳香化
                                            smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                            print(f"✅ Converted molblock to SMILES (removeHs=False, kekuleSmiles=True): {smiles}")
                                    except Exception as e2:
                                        print(f"⚠️ Error with removeHs=False: {str(e2)}")
                            except Exception as e:
                                print(f"⚠️ Error converting molblock to SMILES: {str(e)}")
                                import traceback
                                traceback.print_exc()

                        if smiles:
                            print(f"✅ SMILES extracted: {smiles}")

                            # 🔄 如果API直接返回SMILES，也需要转换为Kekule形式
                            original_smiles = smiles
                            try:
                                from rdkit import Chem
                                mol = Chem.MolFromSmiles(original_smiles)
                                if mol is not None:
                                    # 检查是否包含芳香原子
                                    has_aromatic = any(atom.GetIsAromatic() for atom in mol.GetAtoms())
                                    if has_aromatic:
                                        print(f"🔄 Converting aromatic SMILES to Kekule form: {original_smiles}")
                                        try:
                                            # 先调用Kekulize()来去芳香化分子对象
                                            Chem.Kekulize(mol, clearAromaticFlags=True)
                                            smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                            print(f"✅ Converted to Kekule SMILES: {smiles}")
                                        except Exception as kekulize_error:
                                            print(f"⚠️ Kekulize with clearAromaticFlags failed: {kekulize_error}")
                                            try:
                                                Chem.Kekulize(mol, clearAromaticFlags=False)
                                                smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                                print(f"✅ Converted to Kekule SMILES (without clearing flags): {smiles}")
                                            except Exception as e2:
                                                print(f"⚠️ Kekulize failed: {e2}, using kekuleSmiles parameter only")
                                                smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                            except Exception as e:
                                print(f"⚠️ RDKit conversion failed: {e}, using original SMILES")
                                smiles = original_smiles

                            # 🔄 标记替换处理：将SMILES中的标记原子替换为异丙烯基
                            # 支持的标记: Ce, [Ce], [Co], [Ge], .[Ce]
                            replaced_smiles = None
                            replacement_error = None
                            markers = ['.[Ce]', '[Ce]', '[Co]', '[Ge]', 'Ce']
                            has_marker = any(marker in smiles for marker in markers)

                            if has_marker:
                                try:
                                    detected_markers = [m for m in markers if m in smiles]
                                    print(f"🔄 Detected markers {detected_markers} in SMILES: {smiles}")
                                    print(f"🔄 Replacing with isopropenyl group C=C(C)")
                                    replaced_smiles = replace_ce_with_allyl(smiles)
                                    print(f"✅ Marker replaced successfully: {replaced_smiles}")
                                except Exception as e:
                                    print(f"⚠️ Error replacing marker: {str(e)}")
                                    replacement_error = str(e)
                                    replaced_smiles = None
                            else:
                                print(f"ℹ️ No marker found in SMILES, skipping replacement")

                            # 生成RDKit化学结构图像并保存到数据库(只为被替换过的SMILES生成)
                            rdkit_image_id = None
                            if replaced_smiles:
                                try:
                                    from rdkit import Chem
                                    from rdkit.Chem import Draw
                                    import io
                                    from gridfs import GridFS

                                    mol = Chem.MolFromSmiles(replaced_smiles)
                                    if mol is not None:
                                        # 生成图像
                                        img = Draw.MolToImage(mol, size=(300, 300))
                                        img_byte_arr = io.BytesIO()
                                        img.save(img_byte_arr, format='PNG')
                                        img_bytes = img_byte_arr.getvalue()

                                        # 保存到GridFS (使用smiles_extractor的数据库连接)
                                        fs = GridFS(smiles_extractor.db, collection='rdkit_structures')
                                        rdkit_image_id = fs.put(
                                            img_bytes,
                                            filename=f'rdkit_{replaced_smiles[:50]}.png',
                                            content_type='image/png',
                                            smiles=replaced_smiles,
                                            created_at=datetime.datetime.utcnow()
                                        )
                                        print(f"✅ RDKit structure image saved to database: {rdkit_image_id}")
                                    else:
                                        print(f"⚠️ Could not generate RDKit structure from SMILES: {replaced_smiles}")
                                except Exception as rdkit_error:
                                    print(f"⚠️ Error generating RDKit structure: {str(rdkit_error)}")

                            # 处理 confidence 值
                            confidence_value = 1  # 默认值改为 1
                            if confidence:
                                try:
                                    confidence_float = float(confidence)
                                    # 如果 confidence 在 0-1 之间，转换为百分比
                                    if 0 <= confidence_float <= 1:
                                        confidence_value = int(confidence_float * 100)
                                        # 如果转换后为 0，改为 1
                                        if confidence_value == 0:
                                            confidence_value = 1
                                    else:
                                        confidence_value = int(confidence_float)
                                        # 如果为 0，改为 1
                                        if confidence_value == 0:
                                            confidence_value = 1
                                except (ValueError, TypeError):
                                    confidence_value = 1  # 异常情况也改为 1

                            print(f"📊 Confidence value: {confidence_value}")

                            # 决定存储哪个SMILES作为主要结果
                            final_smiles = replaced_smiles if replaced_smiles else smiles

                            # 存储到数据库
                            smiles_record = {
                                'smiles': final_smiles,  # 主要存储替换后的SMILES(如果有替换的话)
                                'smiles_original': smiles,  # 保留原始SMILES
                                'smiles_replaced': replaced_smiles,  # 替换后的SMILES
                                'has_ce_replacement': replaced_smiles is not None,  # 是否进行了替换
                                'replacement_error': replacement_error,  # 记录替换错误（如果有）
                                'rdkit_image_id': str(rdkit_image_id) if rdkit_image_id else None,  # RDKit图像ID
                                'confidence': confidence_value,
                                'extraction_method': 'AIChemist API',
                                'image_file_id': file_id,
                                'source_file': source_file,
                                'extraction_date': datetime.datetime.utcnow(),
                                'image_type': image_type if image_type else 'chemical_structure',
                                'filename': filename,
                                'segment_number': segment_number
                            }

                            try:
                                result = db_manager.db['smiles_results'].insert_one(smiles_record)
                                print(f"✅ SMILES stored to database with ID: {result.inserted_id}")

                                # 决定返回哪个SMILES作为主要结果
                                final_smiles = replaced_smiles if replaced_smiles else smiles

                                response_data = {
                                    'success': True,
                                    'smiles': final_smiles,  # 主要返回替换后的SMILES(如果有替换的话)
                                    'smiles_original': smiles,  # 原始SMILES
                                    'smiles_replaced': replaced_smiles,  # 替换后的SMILES
                                    'has_ce_replacement': replaced_smiles is not None,  # 是否进行了Ce替换
                                    'replacement_error': replacement_error,  # 替换错误信息
                                    'confidence': smiles_record['confidence'],
                                    'method': 'AIChemist API',
                                    'file_id': file_id,
                                    'stored_to_db': True,
                                    'smiles_db_id': str(result.inserted_id),
                                    'message': 'SMILES extracted successfully using AIChemist API'
                                }

                                if replaced_smiles:
                                    response_data['message'] += ' (Marker replaced with isopropenyl group)'

                                return jsonify(response_data)
                            except Exception as e:
                                print(f"⚠️ Error storing to database: {str(e)}")

                                # 决定返回哪个SMILES作为主要结果
                                final_smiles = replaced_smiles if replaced_smiles else smiles

                                response_data = {
                                    'success': True,
                                    'smiles': final_smiles,  # 主要返回替换后的SMILES(如果有替换的话)
                                    'smiles_original': smiles,
                                    'smiles_replaced': replaced_smiles,
                                    'has_ce_replacement': replaced_smiles is not None,
                                    'replacement_error': replacement_error,
                                    'confidence': confidence_value,
                                    'method': 'AIChemist API',
                                    'file_id': file_id,
                                    'stored_to_db': False,
                                    'message': 'SMILES extracted but not stored to database'
                                }

                                if replaced_smiles:
                                    response_data['message'] += ' (Marker replaced with isopropenyl group)'

                                return jsonify(response_data)
                        else:
                            print(f"❌ No SMILES extracted from AIChemist API")
                            print(f"❌ API response was: {api_result}")
                            print(f"❌ Tried to extract from keys: smiles, SMILES, smi, molblock, mol, MOL")

                            # 最后的尝试：检查是否有任何包含 SMILES 或 MOL 的字段
                            for key, value in api_result.items():
                                print(f"❌ Available key: {key} = {str(value)[:100]}")

                            return jsonify({
                                'success': False,
                                'error': 'No SMILES extracted from AIChemist API',
                                'api_response': str(api_result)[:500]
                            }), 400
                    else:
                        print(f"❌ AIChemist API error: {response.status_code}")
                        print(f"Response text: {response.text}")
                        return jsonify({
                            'success': False,
                            'error': f'AIChemist API error: {response.status_code}'
                        }), 400

                except Exception as e:
                    print(f"❌ Error calling AIChemist API: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    return jsonify({
                        'success': False,
                        'error': f'AIChemist API error: {str(e)}'
                    }), 500

            else:
                # Model A: 使用 DECIMER
                print("🔍 Using DECIMER for SMILES extraction")
                smiles_extractor = get_smiles_extractor()

                # 如果使用YOLO预处理，需要先将图像保存到临时位置
                if use_yolo and image_data is not None:
                    print("🔍 Processing YOLO preprocessed image with DECIMER")
                    import tempfile
                    import os
                    from PIL import Image
                    import io

                    # 将图像数据转换为PIL Image
                    image_pil = Image.open(io.BytesIO(image_data))

                    # 保存到临时文件
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                        tmp_path = tmp_file.name
                        image_pil.save(tmp_path)

                    try:
                        # 使用DECIMER处理临时文件
                        result = smiles_extractor.extract_smiles_from_image(tmp_path)

                        # 存储到数据库(即使是YOLO预处理的图像也需要存储)
                        if result.get('success'):
                            try:
                                # 获取原始图像的元数据
                                from bson import ObjectId
                                file_doc = smiles_extractor.medicinal_image_fs.get(ObjectId(file_id))

                                smiles_doc = {
                                    'image_file_id': file_id,
                                    'source_file': file_doc.metadata.get('source_file'),
                                    'segment_number': file_doc.metadata.get('segment_number'),
                                    'image_type': file_doc.metadata.get('image_type'),
                                    'smiles': result['smiles'],
                                    'confidence': result.get('confidence', 1.0),
                                    'extraction_method': 'YOLO + DECIMER',
                                    'extraction_date': datetime.datetime.utcnow(),
                                    'filename': file_doc.filename
                                }

                                # 插入到数据库
                                smiles_collection = db_manager.db['smiles_results']
                                smiles_id = smiles_collection.insert_one(smiles_doc).inserted_id
                                result['smiles_db_id'] = str(smiles_id)
                                result['stored_to_db'] = True
                                print(f"✅ DECIMER SMILES stored to database with ID: {smiles_id}")
                            except Exception as db_error:
                                print(f"⚠️ Failed to store DECIMER SMILES to database: {db_error}")
                                result['stored_to_db'] = False
                    finally:
                        # 清理临时文件
                        if os.path.exists(tmp_path):
                            os.remove(tmp_path)
                else:
                    # 直接使用file_id提取SMILES并存储到数据库
                    result = smiles_extractor.extract_smiles_from_file_id(file_id, store_to_db=True)

                if result.get('success'):
                    smiles = result.get('smiles')

                    # 🔄 标记替换处理：将SMILES中的标记原子替换为异丙烯基
                    # 支持的标记: Ce, [Ce], [Co], [Ge], .[Ce]
                    replaced_smiles = None
                    replacement_error = None
                    markers = ['.[Ce]', '[Ce]', '[Co]', '[Ge]', 'Ce']
                    has_marker = smiles and any(marker in smiles for marker in markers)

                    if has_marker:
                        try:
                            detected_markers = [m for m in markers if m in smiles]
                            print(f"🔄 Detected markers {detected_markers} in SMILES: {smiles}")
                            print(f"🔄 Replacing with isopropenyl group C=C(C)")
                            replaced_smiles = replace_ce_with_allyl(smiles)
                            print(f"✅ Marker replaced successfully: {replaced_smiles}")

                            # 生成RDKit化学结构图像并保存到数据库
                            rdkit_image_id = None
                            try:
                                from rdkit import Chem
                                from rdkit.Chem import Draw
                                import io
                                from gridfs import GridFS

                                mol = Chem.MolFromSmiles(replaced_smiles)
                                if mol is not None:
                                    # 生成图像
                                    img = Draw.MolToImage(mol, size=(300, 300))
                                    img_byte_arr = io.BytesIO()
                                    img.save(img_byte_arr, format='PNG')
                                    img_bytes = img_byte_arr.getvalue()

                                    # 保存到GridFS (使用smiles_extractor的数据库连接)
                                    fs = GridFS(smiles_extractor.db, collection='rdkit_structures')
                                    rdkit_image_id = fs.put(
                                        img_bytes,
                                        filename=f'rdkit_{replaced_smiles[:50]}.png',
                                        content_type='image/png',
                                        smiles=replaced_smiles,
                                        created_at=datetime.datetime.utcnow()
                                    )
                                    print(f"✅ RDKit structure image saved to database: {rdkit_image_id}")
                                else:
                                    print(f"⚠️ Could not generate RDKit structure from SMILES: {replaced_smiles}")
                            except Exception as rdkit_error:
                                print(f"⚠️ Error generating RDKit structure: {str(rdkit_error)}")

                            # 如果存储到了数据库,更新数据库记录
                            if result.get('stored_to_db') and result.get('smiles_db_id'):
                                try:
                                    from bson import ObjectId
                                    update_data = {
                                        'smiles': replaced_smiles,  # 更新主要的smiles字段为替换后的值
                                        'smiles_original': smiles,  # 保存原始SMILES
                                        'smiles_replaced': replaced_smiles,
                                        'has_ce_replacement': True,
                                        'replacement_error': None
                                    }
                                    if rdkit_image_id:
                                        update_data['rdkit_image_id'] = str(rdkit_image_id)

                                    db_manager.db['smiles_results'].update_one(
                                        {'_id': ObjectId(result.get('smiles_db_id'))},
                                        {'$set': update_data}
                                    )
                                    print(f"✅ Updated database record with replaced SMILES and RDKit image")
                                except Exception as e:
                                    print(f"⚠️ Error updating database with replaced SMILES: {str(e)}")
                        except Exception as e:
                            print(f"⚠️ Error replacing marker: {str(e)}")
                            replacement_error = str(e)
                            replaced_smiles = None

                            # 如果存储到了数据库,记录错误(但保持原始SMILES)
                            if result.get('stored_to_db') and result.get('smiles_db_id'):
                                try:
                                    from bson import ObjectId
                                    db_manager.db['smiles_results'].update_one(
                                        {'_id': ObjectId(result.get('smiles_db_id'))},
                                        {'$set': {
                                            'smiles_original': smiles,  # 保存原始SMILES
                                            'smiles_replaced': None,
                                            'has_ce_replacement': False,
                                            'replacement_error': replacement_error
                                        }}
                                    )
                                except Exception as e2:
                                    print(f"⚠️ Error updating database with replacement error: {str(e2)}")
                    else:
                        if smiles:
                            print(f"ℹ️ No marker found in SMILES, skipping replacement")
                            # 即使没有标记,也要更新数据库记录标明没有替换
                            if result.get('stored_to_db') and result.get('smiles_db_id'):
                                try:
                                    from bson import ObjectId
                                    db_manager.db['smiles_results'].update_one(
                                        {'_id': ObjectId(result.get('smiles_db_id'))},
                                        {'$set': {
                                            'smiles_original': smiles,
                                            'smiles_replaced': None,
                                            'has_ce_replacement': False,
                                            'replacement_error': None
                                        }}
                                    )
                                except Exception as e:
                                    print(f"⚠️ Error updating database: {str(e)}")

                    method_name = 'YOLO + DECIMER Model' if use_yolo else 'DECIMER Model'

                    # 决定返回哪个SMILES作为主要结果
                    final_smiles = replaced_smiles if replaced_smiles else smiles

                    response_data = {
                        'success': True,
                        'smiles': final_smiles,  # 主要返回替换后的SMILES(如果有替换的话)
                        'smiles_original': smiles,  # 原始SMILES
                        'smiles_replaced': replaced_smiles,  # 替换后的SMILES
                        'has_ce_replacement': replaced_smiles is not None,  # 是否进行了Ce替换
                        'replacement_error': replacement_error,  # 替换错误信息
                        'confidence': result.get('confidence'),
                        'method': method_name,
                        'file_id': file_id,
                        'stored_to_db': result.get('stored_to_db', False),
                        'smiles_db_id': result.get('smiles_db_id'),
                        'message': result.get('message', 'SMILES extracted successfully')
                    }

                    if replaced_smiles:
                        response_data['message'] += ' (Marker replaced with isopropenyl group)'

                    return jsonify(response_data)
                else:
                    return jsonify({
                    'success': False,
                    'error': result.get('error', 'SMILES extraction failed')
                })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'SMILES extraction failed: {str(e)}'
            }), 500

    @app.route('/api/smiles/results', methods=['GET'])
    def get_smiles_results():
        """获取SMILES提取结果 (支持无限滚动)"""
        try:
            # 获取分页参数
            offset = request.args.get('offset', 0, type=int)
            limit = request.args.get('limit', 50, type=int)

            # 验证参数
            if offset < 0:
                offset = 0
            if limit < 1 or limit > 500:
                limit = 50

            smiles_extractor = get_smiles_extractor()
            all_results = smiles_extractor.get_all_smiles_results()

            # 计算总数
            total_count = len(all_results)

            # 获取当前批次的数据
            page_results = all_results[offset:offset + limit]

            # 判断是否还有更多数据
            has_more = (offset + limit) < total_count

            print(f"📄 SMILES Results Infinite Scroll: offset={offset}, limit={limit}, total={total_count}, has_more={has_more}")

            return jsonify({
                'success': True,
                'results': page_results,
                'total_count': total_count,
                'offset': offset,
                'limit': limit,
                'has_more': has_more
            })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Failed to retrieve SMILES results: {str(e)}',
                'results': [],
                'total_count': 0,
                'has_more': False
            }), 500

    @app.route('/api/smiles/export/all/excel', methods=['GET'])
    def export_all_smiles_to_excel():
        """导出所有SMILES结果到Excel（包含所有图像）- 直接从数据库导出"""
        print("🚀 Starting SMILES export to Excel (ALL records)...")
        try:
            import io
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.drawing.image import Image as XLImage
            from PIL import Image as PILImage
            from datetime import datetime as dt
            import io as iolib

            # 获取所有SMILES结果（不分页）
            smiles_extractor = get_smiles_extractor()
            all_results = smiles_extractor.get_all_smiles_results()

            total_records = len(all_results)
            print(f"📊 Total SMILES records to export: {total_records}")

            if not all_results:
                return jsonify({
                    'success': False,
                    'error': 'No SMILES results found in database'
                }), 404

            # 创建Excel工作簿
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = 'SMILES Results'

            # 定义列 (增加图像列宽度以适应更大的图片)
            columns = [
                ('Index', 8),
                ('Chemical Image', 30),  # 从25增加到30
                ('SMILES', 50),
                ('Structure (RDKit)', 30),  # 从25增加到30
                ('Image File ID', 25),
                ('Database ID', 25),
                ('Confidence', 12),
                ('Source File', 30),
                ('Extraction Method', 15),
                ('Extraction Date', 20),
                ('Image Type', 15),
                ('Filename', 30)
            ]

            # 设置列头
            for col_idx, (header_name, width) in enumerate(columns, 1):
                cell = worksheet.cell(row=1, column=col_idx)
                cell.value = header_name
                cell.font = Font(bold=True, color='FFFFFFFF')
                cell.fill = PatternFill(start_color='FF4472C4', end_color='FF4472C4', fill_type='solid')
                cell.alignment = Alignment(horizontal='center', vertical='center')
                worksheet.column_dimensions[chr(64 + col_idx)].width = width

            # 设置行高以容纳图片
            worksheet.row_dimensions[1].height = 20

            # 统计信息
            success_count = 0
            fail_count = 0
            image_load_errors = []

            # 添加数据行
            print(f"📝 Processing {total_records} SMILES records...")
            for row_idx, result in enumerate(all_results, 2):
                try:
                    # 提取日期
                    extraction_date = result.get('extraction_date')
                    if extraction_date:
                        if isinstance(extraction_date, str):
                            extraction_date = extraction_date
                        else:
                            extraction_date = extraction_date.isoformat() if hasattr(extraction_date, 'isoformat') else str(extraction_date)
                    else:
                        extraction_date = 'N/A'

                    # 使用替换后的SMILES（如果有），否则使用原始SMILES
                    final_smiles = result.get('smiles_replaced') or result.get('smiles') or 'N/A'

                    # 添加行数据
                    row_data = [
                        row_idx - 1,  # Index
                        '',  # Chemical Image (将单独添加)
                        final_smiles,  # SMILES
                        '',  # Structure (RDKit) (将单独添加)
                        result.get('image_file_id', 'N/A'),  # Image File ID
                        result.get('_id', 'N/A'),  # Database ID
                        result.get('confidence', 'N/A'),  # Confidence
                        result.get('source_file', 'N/A'),  # Source File
                        result.get('extraction_method', 'DECIMER'),  # Method
                        extraction_date,  # Date
                        result.get('image_type', 'N/A'),  # Image Type
                        result.get('filename', 'N/A')  # Filename
                    ]

                    # 添加行到工作表
                    for col_idx, value in enumerate(row_data, 1):
                        cell = worksheet.cell(row=row_idx, column=col_idx)
                        cell.value = value
                        cell.border = Border(
                            left=Side(style='thin'),
                            right=Side(style='thin'),
                            top=Side(style='thin'),
                            bottom=Side(style='thin')
                        )
                        cell.alignment = Alignment(horizontal='left', vertical='center')

                    # 设置行高以容纳更大的图片 (从80增加到120)
                    worksheet.row_dimensions[row_idx].height = 120

                    # 异步加载和添加化学图像
                    if result.get('image_file_id'):
                        try:
                            image_data = smiles_extractor.get_image_by_id(result['image_file_id'])
                            if image_data:
                                try:
                                    # 将图像数据转换为PIL Image
                                    img = PILImage.open(iolib.BytesIO(image_data))

                                    # 不压缩图片，保持原始质量
                                    # 只在Excel中调整显示尺寸

                                    # 保存到BytesIO (保持原始尺寸)
                                    img_bytes = iolib.BytesIO()
                                    img.save(img_bytes, format='PNG')
                                    img_bytes.seek(0)

                                    # 添加图像到Excel (openpyxl正确用法)
                                    img_obj = XLImage(img_bytes)
                                    # 设置较大的显示尺寸以提高清晰度
                                    img_obj.width = 200  # 从120增加到200
                                    img_obj.height = 150  # 从60增加到150，保持比例
                                    worksheet.add_image(img_obj, f'B{row_idx}')
                                    print(f"✅ Image added for row {row_idx}: {result['image_file_id'][:8]}...")
                                except Exception as img_error:
                                    print(f"⚠️ Failed to process image for row {row_idx}: {str(img_error)}")
                                    worksheet.cell(row=row_idx, column=2).value = f'Image processing error'
                                    image_load_errors.append((row_idx, str(img_error)))
                            else:
                                worksheet.cell(row=row_idx, column=2).value = 'Image not available'
                                print(f"⚠️ Image data not found for row {row_idx}: {result['image_file_id']}")
                        except Exception as img_error:
                            print(f"⚠️ Failed to load image for row {row_idx}: {str(img_error)}")
                            worksheet.cell(row=row_idx, column=2).value = 'Failed to load'
                            image_load_errors.append((row_idx, str(img_error)))

                    # 加载RDKit结构图像（如果有）
                    if result.get('rdkit_image_id') and result['rdkit_image_id'] != 'None':
                        try:
                            from bson import ObjectId
                            from gridfs import GridFS

                            fs = GridFS(smiles_extractor.db, collection='rdkit_structures')
                            try:
                                grid_out = fs.get(ObjectId(result['rdkit_image_id']))
                                rdkit_image_data = grid_out.read()

                                # 处理RDKit图像 - 保持原始质量
                                rdkit_img = PILImage.open(iolib.BytesIO(rdkit_image_data))

                                # 不压缩图片，保持原始质量
                                # 只在Excel中调整显示尺寸

                                # 保存到BytesIO (保持原始尺寸)
                                rdkit_img_bytes = iolib.BytesIO()
                                rdkit_img.save(rdkit_img_bytes, format='PNG')
                                rdkit_img_bytes.seek(0)

                                # 添加图像到Excel (openpyxl正确用法)
                                rdkit_img_obj = XLImage(rdkit_img_bytes)
                                # 设置较大的显示尺寸以提高清晰度
                                rdkit_img_obj.width = 200  # 从120增加到200
                                rdkit_img_obj.height = 150  # 从60增加到150，保持比例
                                worksheet.add_image(rdkit_img_obj, f'D{row_idx}')
                                print(f"✅ RDKit image added for row {row_idx}")
                            except Exception as rdkit_error:
                                print(f"⚠️ RDKit image not found for row {row_idx}: {str(rdkit_error)}")
                                worksheet.cell(row=row_idx, column=4).value = 'N/A'
                        except Exception as rdkit_error:
                            print(f"⚠️ Failed to load RDKit image for row {row_idx}: {str(rdkit_error)}")
                            worksheet.cell(row=row_idx, column=4).value = 'N/A'

                    if row_idx % 100 == 0:
                        print(f"📄 Processed {row_idx - 1} records...")

                    success_count += 1

                except Exception as row_error:
                    print(f"❌ Error processing row {row_idx}: {str(row_error)}")
                    fail_count += 1
                    continue

            # 生成Excel文件
            print(f"✅ Successfully processed {success_count} records (with {fail_count} failures)")

            # 生成文件
            excel_bytes = iolib.BytesIO()
            workbook.save(excel_bytes)
            excel_bytes.seek(0)

            # 生成响应
            timestamp = dt.now().strftime('%Y%m%d_%H%M%S')
            filename = f'SMILES_Results_AllData_{timestamp}.xlsx'

            print(f"📦 Excel file generated: {filename}")

            return send_file(
                excel_bytes,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=filename
            )

        except Exception as e:
            print(f"❌ Error exporting SMILES to Excel: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': f'Failed to export SMILES to Excel: {str(e)}'
            }), 500

    @app.route('/api/smiles/generate-structure-image', methods=['POST'])
    def generate_structure_image():
        """使用RDKit从SMILES生成化学结构图像"""
        try:
            from rdkit import Chem
            from rdkit.Chem import Draw
            import io

            data = request.get_json()
            smiles = data.get('smiles')

            if not smiles:
                return jsonify({'error': 'SMILES is required'}), 400

            # 使用RDKit生成分子对象
            mol = Chem.MolFromSmiles(smiles)
            if mol is None:
                return jsonify({'error': 'Invalid SMILES'}), 400

            # 生成分子图像
            img = Draw.MolToImage(mol, size=(300, 300))

            # 将图像转换为PNG字节流
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)

            # 返回图像
            return send_file(
                img_byte_arr,
                mimetype='image/png',
                as_attachment=False
            )

        except Exception as e:
            print(f"❌ Error generating structure image: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Failed to generate structure image: {str(e)}'}), 500

    @app.route('/api/rdkit-images/<image_id>', methods=['GET'])
    def get_rdkit_image(image_id):
        """从数据库获取RDKit化学结构图像"""
        print(f"🔍 API request for RDKit image: {image_id}")
        try:
            from bson import ObjectId
            from gridfs import GridFS
            import io

            # 获取smiles_extractor实例
            print(f"🔍 Getting smiles_extractor instance...")
            extractor = get_smiles_extractor()
            print(f"✅ Got smiles_extractor, database: {extractor.db.name}")

            # 从GridFS获取图像 (使用smiles_extractor的数据库连接)
            fs = GridFS(extractor.db, collection='rdkit_structures')
            print(f"✅ GridFS initialized for collection: rdkit_structures")

            try:
                print(f"🔍 Attempting to retrieve image with ID: {image_id}")
                grid_out = fs.get(ObjectId(image_id))
                image_data = grid_out.read()

                print(f"✅ RDKit image retrieved successfully: {image_id}, size: {len(image_data)} bytes")

                # 返回图像
                return send_file(
                    io.BytesIO(image_data),
                    mimetype='image/png',
                    as_attachment=False
                )
            except Exception as e:
                print(f"❌ RDKit image not found in GridFS: {image_id}")
                print(f"   Error type: {type(e).__name__}")
                print(f"   Error message: {str(e)}")
                import traceback
                traceback.print_exc()
                return jsonify({'error': 'Image not found'}), 404

        except Exception as e:
            print(f"❌ Error retrieving RDKit image: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Failed to retrieve image: {str(e)}'}), 500

    @app.route('/api/molecules/process', methods=['POST'])
    @login_required
    @track_usage
    def process_molecules():
        """处理分子结构识别"""
        data = request.get_json()
        document_id = data.get('document_id')
        model_type = data.get('model_type')

        if not document_id or not model_type:
            return jsonify({'error': 'Missing document_id or model_type'}), 400

        document = db_manager.get_document_by_id(int(document_id))
        if not document:
            return jsonify({'error': 'Document not found'}), 404

        import random
        import time

        # 模拟处理时间
        time.sleep(1)

        if model_type == 'modelA':
            # Model A: 使用 DECIMER 模型
            try:
                chemical_images = db_manager.get_chemical_structure_images(str(document_id))

                if not chemical_images:
                    return jsonify({
                        'status': 'success',
                        'model_type': 'modelA',
                        'processed_count': 0,
                        'smiles_count': 0,
                        'confidence': 0,
                        'results': [],
                        'message': 'No chemical structure images found'
                    })

                results = []
                success_count = 0

                # 使用 DECIMER 处理每个化学结构图像
                for image_info in chemical_images[:10]:
                    try:
                        file_id = image_info.get('file_id')
                        image_data = db_manager.get_image_by_id(file_id)

                        if not image_data:
                            continue

                        # 这里可以调用实际的 DECIMER 模型
                        # 目前使用模拟数据
                        sample_smiles = [
                            'CCO', 'C1=CC=CC=C1', 'CC(=O)O', 'CCN(CC)CC', 'C1CCCCC1',
                            'CC(C)NC(=O)C1=CC=CC=C1', 'CCN(CC)C(=O)C1=CC=C(C=C1)O'
                        ]
                        smiles = random.choice(sample_smiles)

                        # 🔄 确保SMILES是Kekule形式（去芳香化）
                        try:
                            from smiles_converter import ensure_kekule_smiles, has_aromatic_smiles
                            if has_aromatic_smiles(smiles):
                                print(f"🔄 Converting aromatic SMILES to Kekule form: {smiles}")
                                smiles = ensure_kekule_smiles(smiles, verbose=True)
                                print(f"✅ Converted to Kekule SMILES: {smiles}")
                        except Exception as e:
                            print(f"⚠️ SMILES conversion failed: {e}")

                        results.append({
                            'smiles': smiles,
                            'confidence': random.randint(75, 95),
                            'molecular_weight': round(random.uniform(150, 400), 2),
                            'molecular_formula': f'C{random.randint(6,20)}H{random.randint(8,30)}N{random.randint(0,3)}O{random.randint(0,4)}',
                            'structure_image': f'https://via.placeholder.com/200x150/FFFFFF/000000?text={smiles.replace("=", "%3D")}',
                            'source': 'DECIMER',
                            'file_id': file_id
                        })
                        success_count += 1
                    except Exception as e:
                        print(f"Error processing image {file_id} with DECIMER: {str(e)}")
                        continue

                return jsonify({
                    'status': 'success',
                    'model_type': 'modelA',
                    'processed_count': len(chemical_images),
                    'smiles_count': success_count,
                    'confidence': int(sum([r['confidence'] for r in results]) / len(results)) if results else 0,
                    'results': results,
                    'source': 'DECIMER'
                })

            except Exception as e:
                print(f"DECIMER processing error: {str(e)}")
                return jsonify({
                    'error': f'DECIMER processing failed: {str(e)}'
                }), 500

        elif model_type == 'modelB':
            # Model B: 使用 AIChemist API (已在 /api/molecules/process-aichemist 中实现)
            # 这里保留向后兼容性
            result_count = random.randint(2, 5)
            results = []

            for i in range(result_count):
                sample_smiles = [
                    'CCO', 'C1=CC=CC=C1', 'CC(=O)O', 'CCN(CC)CC', 'C1CCCCC1',
                    'CC(C)NC(=O)C1=CC=CC=C1', 'CCN(CC)C(=O)C1=CC=C(C=C1)O'
                ]
                smiles = random.choice(sample_smiles)

                # 🔄 确保SMILES是Kekule形式（去芳香化）
                try:
                    from smiles_converter import ensure_kekule_smiles, has_aromatic_smiles
                    if has_aromatic_smiles(smiles):
                        print(f"🔄 Converting aromatic SMILES to Kekule form: {smiles}")
                        smiles = ensure_kekule_smiles(smiles, verbose=True)
                        print(f"✅ Converted to Kekule SMILES: {smiles}")
                except Exception as e:
                    print(f"⚠️ SMILES conversion failed: {e}")

                results.append({
                    'smiles': smiles,
                    'confidence': random.randint(75, 95),
                    'molecular_weight': round(random.uniform(150, 400), 2),
                    'molecular_formula': f'C{random.randint(6,20)}H{random.randint(8,30)}N{random.randint(0,3)}O{random.randint(0,4)}',
                    'structure_image': f'https://via.placeholder.com/200x150/FFFFFF/000000?text={smiles.replace("=", "%3D")}',
                    'source': 'AIChemist'
                })

            return jsonify({
                'status': 'success',
                'model_type': model_type,
                'processed_count': result_count,
                'smiles_count': result_count,
                'confidence': random.randint(80, 95),
                'results': results,
                'source': 'AIChemist'
            })

        elif model_type == 'ensemble':
            # Ensemble模式处理
            processed_count = random.randint(3, 6)
            disagreement_count = random.randint(0, 2)  # 随机决定是否有不一致

            if disagreement_count == 0:
                # 所有结果一致
                consensus_results = []
                for i in range(processed_count):
                    sample_smiles = [
                        'CCO', 'C1=CC=CC=C1', 'CC(=O)O', 'CCN(CC)CC'
                    ]
                    smiles = random.choice(sample_smiles)

                    consensus_results.append({
                        'smiles': smiles,
                        'confidence': random.randint(85, 98),
                        'molecular_weight': round(random.uniform(150, 400), 2),
                        'molecular_formula': f'C{random.randint(6,20)}H{random.randint(8,30)}N{random.randint(0,3)}O{random.randint(0,4)}',
                        'structure_image': f'https://via.placeholder.com/200x150/FFFFFF/000000?text={smiles.replace("=", "%3D")}'
                    })

                return jsonify({
                    'status': 'success',
                    'model_type': 'ensemble',
                    'processed_count': processed_count,
                    'smiles_count': processed_count,
                    'confidence': random.randint(90, 98),
                    'disagreements': [],
                    'consensus_results': consensus_results
                })
            else:
                # 有不一致的结果
                disagreements = []
                for i in range(disagreement_count):
                    # 生成两个不同的SMILES
                    smiles_options = ['CCO', 'C1=CC=CC=C1', 'CC(=O)O', 'CCN(CC)CC', 'C1CCCCC1']
                    modelA_smiles = random.choice(smiles_options)
                    modelB_smiles = random.choice([s for s in smiles_options if s != modelA_smiles])

                    disagreements.append({
                        'original_image': f'https://via.placeholder.com/200x150/1A1A2E/FFFFFF?text=Original+{i+1}',
                        'modelA_smiles': modelA_smiles,
                        'modelA_confidence': random.randint(70, 90),
                        'modelA_image': f'https://via.placeholder.com/200x150/00D4FF/FFFFFF?text={modelA_smiles.replace("=", "%3D")}',
                        'modelB_smiles': modelB_smiles,
                        'modelB_confidence': random.randint(70, 90),
                        'modelB_image': f'https://via.placeholder.com/200x150/FF6B35/FFFFFF?text={modelB_smiles.replace("=", "%3D")}'
                    })

                # 生成模型A和B的完整结果
                modelA_results = {
                    'results': [
                        {
                            'smiles': d['modelA_smiles'],
                            'confidence': d['modelA_confidence'],
                            'molecular_weight': round(random.uniform(150, 400), 2),
                            'molecular_formula': f'C{random.randint(6,20)}H{random.randint(8,30)}N{random.randint(0,3)}O{random.randint(0,4)}',
                            'structure_image': d['modelA_image']
                        } for d in disagreements
                    ]
                }

                modelB_results = {
                    'results': [
                        {
                            'smiles': d['modelB_smiles'],
                            'confidence': d['modelB_confidence'],
                            'molecular_weight': round(random.uniform(150, 400), 2),
                            'molecular_formula': f'C{random.randint(6,20)}H{random.randint(8,30)}N{random.randint(0,3)}O{random.randint(0,4)}',
                            'structure_image': d['modelB_image']
                        } for d in disagreements
                    ]
                }

                return jsonify({
                    'status': 'success',
                    'model_type': 'ensemble',
                    'processed_count': processed_count,
                    'smiles_count': processed_count - disagreement_count,
                    'confidence': random.randint(75, 90),
                    'disagreements': disagreements,
                    'modelA_results': modelA_results,
                    'modelB_results': modelB_results
                })

        return jsonify({'error': 'Invalid model_type'}), 400

    @app.route('/api/molecules/process-aichemist', methods=['POST'])
    @login_required
    @track_usage
    def process_molecules_aichemist():
        """使用 AIChemist API 处理分子结构识别"""
        import requests
        import base64
        from io import BytesIO
        from PIL import Image

        data = request.get_json()
        document_id = data.get('document_id')

        if not document_id:
            return jsonify({'error': 'Missing document_id'}), 400

        try:
            # 导入 RDKit 用于 molblock 转 SMILES
            try:
                from rdkit import Chem
                rdkit_available = True
            except ImportError:
                rdkit_available = False
                print("⚠️ RDKit not available, molblock conversion will be skipped")

            # 获取文档信息
            document = db_manager.get_document_by_id(int(document_id))
            if not document:
                return jsonify({'error': 'Document not found'}), 404

            # 获取化学结构图像
            chemical_images = db_manager.get_chemical_structure_images(str(document_id))

            if not chemical_images:
                return jsonify({
                    'status': 'success',
                    'model_type': 'modelB',
                    'processed_count': 0,
                    'smiles_count': 0,
                    'confidence': 0,
                    'results': [],
                    'message': 'No chemical structure images found'
                })

            # AIChemist API 配置
            aichemist_url = 'https://api-ocsr.alchemist.iresearch.net.cn/ocsr/'
            headers = {
                'X-API-Version': '1.0',
                'Content-Type': 'application/json'
            }

            results = []
            success_count = 0

            # 处理每个化学结构图像
            for image_info in chemical_images[:10]:  # 限制处理前10个图像
                try:
                    file_id = image_info.get('file_id')

                    # 从 GridFS 获取图像数据
                    image_data = db_manager.get_image_by_id(file_id)
                    if not image_data:
                        continue

                    # 将图像转换为 Base64
                    image_base64 = base64.b64encode(image_data).decode('utf-8')

                    # 调用 AIChemist API
                    payload = {
                        'imageBase64': image_base64
                    }

                    response = requests.post(aichemist_url, json=payload, headers=headers, timeout=30)

                    if response.status_code == 200:
                        api_result = response.json()

                        # 提取 SMILES 和其他信息
                        smiles = api_result.get('smiles', '')
                        molblock = api_result.get('molblock', '')
                        confidence = api_result.get('confidence', 0)

                        # 如果返回的是 molblock 格式，使用 RDKit 转换为 SMILES
                        if molblock and not smiles and rdkit_available:
                            try:
                                mol = Chem.MolFromMolBlock(molblock)
                                if mol is not None:
                                    # 使用 kekuleSmiles=True 参数，避免芳香化
                                    smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                    print(f"✅ Converted molblock to SMILES (kekuleSmiles=True): {smiles}")
                            except Exception as e:
                                print(f"⚠️ Error converting molblock to SMILES: {str(e)}")

                        # 🔄 如果API直接返回SMILES，也需要转换为Kekule形式
                        if smiles and rdkit_available:
                            original_smiles = smiles
                            try:
                                mol = Chem.MolFromSmiles(original_smiles)
                                if mol is not None:
                                    # 检查是否包含芳香原子
                                    has_aromatic = any(atom.GetIsAromatic() for atom in mol.GetAtoms())
                                    if has_aromatic:
                                        print(f"🔄 Converting aromatic SMILES to Kekule form: {original_smiles}")
                                        try:
                                            # 先调用Kekulize()来去芳香化分子对象
                                            Chem.Kekulize(mol, clearAromaticFlags=True)
                                            smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                            print(f"✅ Converted to Kekule SMILES: {smiles}")
                                        except Exception as kekulize_error:
                                            print(f"⚠️ Kekulize with clearAromaticFlags failed: {kekulize_error}")
                                            try:
                                                Chem.Kekulize(mol, clearAromaticFlags=False)
                                                smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                                                print(f"✅ Converted to Kekule SMILES (without clearing flags): {smiles}")
                                            except Exception as e2:
                                                print(f"⚠️ Kekulize failed: {e2}, using kekuleSmiles parameter only")
                                                smiles = Chem.MolToSmiles(mol, kekuleSmiles=True)
                            except Exception as e:
                                print(f"⚠️ RDKit conversion failed: {e}, using original SMILES")
                                smiles = original_smiles

                        if smiles:
                            results.append({
                                'smiles': smiles,
                                'confidence': int(confidence * 100) if confidence <= 1 else int(confidence),
                                'molecular_weight': api_result.get('molecular_weight', 0),
                                'molecular_formula': api_result.get('molecular_formula', ''),
                                'structure_image': f'data:image/png;base64,{image_base64[:100]}...',
                                'source': 'AIChemist',
                                'file_id': file_id
                            })
                            success_count += 1
                    else:
                        print(f"AIChemist API error: {response.status_code}")

                except Exception as e:
                    print(f"Error processing image {file_id}: {str(e)}")
                    continue

            return jsonify({
                'status': 'success',
                'model_type': 'modelB',
                'processed_count': len(chemical_images),
                'smiles_count': success_count,
                'confidence': int(sum([r['confidence'] for r in results]) / len(results)) if results else 0,
                'results': results,
                'source': 'AIChemist API'
            })

        except Exception as e:
            print(f"AIChemist processing error: {str(e)}")
            return jsonify({
                'error': f'AIChemist processing failed: {str(e)}'
            }), 500

    # PDF 分割处理相关路由
    @app.route('/api/pdf/segment', methods=['POST'])
    def segment_pdf():
        """处理 PDF 分割请求"""
        try:
            if 'pdf_files' not in request.files:
                return jsonify({'success': False, 'message': 'No PDF files provided'}), 400

            files = request.files.getlist('pdf_files')
            enable_ocr = request.form.get('enable_ocr', 'true').lower() == 'true'
            enable_formula = request.form.get('enable_formula', 'true').lower() == 'true'
            enable_table = request.form.get('enable_table', 'false').lower() == 'true'

            if not files:
                return jsonify({'success': False, 'message': 'No files selected'}), 400

            # 生成任务ID
            task_id = f"pdf_task_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{random.randint(1000, 9999)}"

            # 保存文件到临时目录
            temp_dir = "/tmp/pdf_processing"
            os.makedirs(temp_dir, exist_ok=True)

            saved_files = []
            for file in files:
                if file.filename.lower().endswith('.pdf'):
                    filename = f"{task_id}_{file.filename}"
                    filepath = os.path.join(temp_dir, filename)
                    file.save(filepath)
                    saved_files.append({
                        'original_name': file.filename,
                        'saved_path': filepath,
                        'filename': filename
                    })

            if not saved_files:
                return jsonify({'success': False, 'message': 'No valid PDF files found'}), 400

            # 启动后台处理任务（这里使用模拟，实际应该使用异步任务队列）
            # 在实际应用中，你可以使用 Celery 或其他任务队列
            import threading

            def process_pdfs_background():
                try:
                    output_dir = f"/root/local-disk/医药分子/picture_yiyao_output/{task_id}"
                    os.makedirs(output_dir, exist_ok=True)

                    results = []
                    for file_info in saved_files:
                        try:
                            # 使用 PDF 处理器处理文件
                            result = pdf_processor.auto_seg(file_info['saved_path'], output_dir)

                            if result['success']:
                                # 存储到数据库 - 保留完整的Markdown内容
                                full_markdown_content = result['content']['markdown']
                                print(f"📄 Storing complete Markdown content: {len(full_markdown_content)} characters")

                                document_id = db_manager.store_document(
                                    file_info['original_name'],
                                    'PDF',
                                    full_markdown_content,  # 保存完整内容，不截断
                                    os.path.getsize(file_info['saved_path']),
                                    1  # user_id
                                )

                                results.append({
                                    'filename': file_info['original_name'],
                                    'document_id': document_id,
                                    'pages': result['pages'],
                                    'mode': result['processing_mode'],
                                    'processing_time': result['processing_time'],
                                    'output_dir': output_dir
                                })
                            else:
                                print(f"Failed to process {file_info['original_name']}: {result.get('error', 'Unknown error')}")

                        except Exception as e:
                            print(f"Error processing {file_info['original_name']}: {e}")
                            continue

                    # 保存任务结果
                    task_result_file = os.path.join(temp_dir, f"{task_id}_result.json")
                    with open(task_result_file, 'w', encoding='utf-8') as f:
                        json.dump({
                            'completed': True,
                            'results': results,
                            'task_id': task_id
                        }, f, ensure_ascii=False, indent=2)

                except Exception as e:
                    print(f"Background processing error: {e}")
                    # 保存错误结果
                    task_result_file = os.path.join(temp_dir, f"{task_id}_result.json")
                    with open(task_result_file, 'w', encoding='utf-8') as f:
                        json.dump({
                            'completed': True,
                            'error': str(e),
                            'results': []
                        }, f, ensure_ascii=False, indent=2)

            # 启动后台线程
            thread = threading.Thread(target=process_pdfs_background)
            thread.daemon = True
            thread.start()

            return jsonify({
                'success': True,
                'task_id': task_id,
                'message': f'Started processing {len(saved_files)} PDF files',
                'files_count': len(saved_files)
            })

        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/pdf/status/<task_id>')
    def get_pdf_status(task_id):
        """获取 PDF 处理状态"""
        try:
            temp_dir = "/tmp/pdf_processing"
            task_result_file = os.path.join(temp_dir, f"{task_id}_result.json")

            if os.path.exists(task_result_file):
                with open(task_result_file, 'r', encoding='utf-8') as f:
                    result = json.load(f)
                return jsonify(result)
            else:
                return jsonify({
                    'completed': False,
                    'message': 'Processing in progress...'
                })
        except Exception as e:
            return jsonify({'completed': False, 'error': str(e)}), 500

    @app.route('/api/pdf/processed')
    def get_processed_pdfs():
        """获取已处理的 PDF 文档列表"""
        try:
            # 从数据库获取 PDF 类型的文档
            documents = db_manager.get_all_documents()
            pdf_documents = [doc for doc in documents if doc.get('file_type', '').upper() == 'PDF']

            # 添加处理时间信息
            for doc in pdf_documents:
                doc['processed_time'] = doc.get('upload_time', datetime.now().isoformat())

            return jsonify({
                'success': True,
                'documents': pdf_documents
            })
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/pdf/view/<document_id>/<format>')
    def view_pdf_document(document_id, format):
        """查看 PDF 处理结果"""
        try:
            # 从数据库获取文档信息
            documents = db_manager.get_all_documents()
            document = None
            for doc in documents:
                if str(doc.get('id')) == str(document_id):
                    document = doc
                    break

            if not document:
                return jsonify({'success': False, 'message': 'Document not found'}), 404

            # 模拟返回处理结果
            if format == 'markdown':
                # 生成模拟的 Markdown 内容
                content = f"""# {document['filename']} - Processing Results

## Document Information
- **Filename**: {document['filename']}
- **File Type**: {document['file_type']}
- **Upload Time**: {document['upload_time']}
- **File Size**: {document.get('file_size', 'Unknown')} bytes

## Extracted Content

### Abstract
This document contains pharmaceutical research data with molecular structures and bioactivity information relevant to drug discovery and development.

### Key Findings
1. **Molecular Structures**: Multiple chemical compounds identified
2. **Bioactivity Data**: IC50 values and efficacy measurements
3. **Safety Profile**: Toxicity assessments and side effect analysis
4. **Pharmacokinetics**: ADME properties and metabolic pathways

### Chemical Compounds
- **Compound A**: C20H25N3O4 (MW: 371.43)
- **Compound B**: C18H22N2O3 (MW: 314.38)
- **Compound C**: C22H28N4O5 (MW: 428.48)

### Experimental Data
| Compound | IC50 (nM) | Selectivity | Toxicity |
|----------|-----------|-------------|----------|
| A        | 12.5      | High        | Low      |
| B        | 8.3       | Medium      | Medium   |
| C        | 15.7      | High        | Low      |

### Conclusions
The analyzed compounds show promising bioactivity profiles with acceptable safety margins for further development.
"""
            elif format == 'json':
                # 生成模拟的 JSON 内容
                content = json.dumps({
                    "document_info": {
                        "filename": document['filename'],
                        "file_type": document['file_type'],
                        "upload_time": document['upload_time'],
                        "file_size": document.get('file_size', 0)
                    },
                    "processing_results": {
                        "pages_processed": random.randint(10, 150),
                        "processing_mode": "OCR",
                        "extraction_method": "magic-pdf",
                        "confidence_score": random.uniform(0.85, 0.98)
                    },
                    "extracted_content": {
                        "sections": [
                            {
                                "type": "title",
                                "content": document['filename'],
                                "confidence": 0.95
                            },
                            {
                                "type": "abstract",
                                "content": "This document contains pharmaceutical research data...",
                                "confidence": 0.89
                            },
                            {
                                "type": "table",
                                "content": {
                                    "headers": ["Compound", "IC50 (nM)", "Selectivity", "Toxicity"],
                                    "rows": [
                                        ["A", "12.5", "High", "Low"],
                                        ["B", "8.3", "Medium", "Medium"],
                                        ["C", "15.7", "High", "Low"]
                                    ]
                                },
                                "confidence": 0.92
                            }
                        ]
                    },
                    "molecular_structures": [
                        {
                            "compound_id": "A",
                            "smiles": "CC(C)CC1=CC=C(C=C1)C(C)C(=O)N2CCN(CC2)C3=CC=CC=C3",
                            "molecular_formula": "C20H25N3O4",
                            "molecular_weight": 371.43
                        },
                        {
                            "compound_id": "B",
                            "smiles": "CC1=CC=C(C=C1)C(=O)NC2=CC=C(C=C2)N3CCCC3",
                            "molecular_formula": "C18H22N2O3",
                            "molecular_weight": 314.38
                        }
                    ]
                }, indent=2)
            else:
                return jsonify({'success': False, 'message': 'Invalid format'}), 400

            return jsonify({
                'success': True,
                'content': content,
                'format': format
            })

        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/pdf/download/<document_id>')
    def download_pdf_result(document_id):
        """下载 PDF 处理结果"""
        try:
            # 从数据库获取文档信息
            documents = db_manager.get_all_documents()
            document = None
            for doc in documents:
                if str(doc.get('id')) == str(document_id):
                    document = doc
                    break

            if not document:
                return jsonify({'success': False, 'message': 'Document not found'}), 404

            # 获取实际的segment files
            from flask import Response
            import zipfile
            import io

            # 获取文档的segment files
            segment_files = db_manager.get_segment_files(document_id)

            if not segment_files:
                return jsonify({'success': False, 'message': 'No processed files found'}), 404

            # 创建ZIP文件包含所有处理结果
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

                for file_info in segment_files:
                    file_type = file_info.get('file_type', 'unknown')
                    filename = file_info.get('virtual_filename', f'file_{file_type}')

                    # 获取文件内容
                    content = db_manager.get_segment_file_content(document_id, file_type)

                    if content:
                        if isinstance(content, bytes):
                            # 二进制文件（如PDF）
                            zip_file.writestr(filename, content)
                        else:
                            # 文本文件
                            zip_file.writestr(filename, str(content))

            zip_buffer.seek(0)

            return Response(
                zip_buffer.getvalue(),
                mimetype='application/zip',
                headers={
                    'Content-Disposition': f'attachment; filename={document["filename"]}_processed.zip'
                }
            )

        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/<document_id>/segments', methods=['GET'])
    def get_document_segments(document_id):
        """获取文档的分割片段"""
        try:
            segments = db_manager.get_document_segments(document_id)
            return jsonify({
                'success': True,
                'document_id': document_id,
                'segments': segments,
                'total_segments': len(segments)
            })
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/database/info', methods=['GET'])
    def get_database_info():
        """获取数据库连接信息"""
        try:
            stats = db_manager.get_database_stats()
            return jsonify({
                'success': True,
                'database_type': stats.get('database_type', 'Unknown'),
                'connection_status': 'Connected' if db_manager.use_mongodb else 'SQLite Fallback',
                'stats': stats
            })
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/<document_id>/processing-results', methods=['GET'])
    def get_document_processing_results(document_id):
        """获取文档的处理结果（Markdown、content_list.json、middle.json）"""
        try:
            results = db_manager.get_processing_results(document_id)
            if results:
                return jsonify({
                    'success': True,
                    'document_id': document_id,
                    'results': results
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'No processing results found for this document'
                }), 404
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/<document_id>/download/<file_type>', methods=['GET'])
    def download_processing_result(document_id, file_type):
        """下载处理结果文件"""
        try:
            results = db_manager.get_processing_results(document_id)
            if not results:
                return jsonify({'success': False, 'message': 'No processing results found'}), 404

            from flask import Response
            import json

            if file_type == 'markdown':
                content = results.get('markdown_content', '')
                mimetype = 'text/markdown'
                filename = f"{results.get('filename', 'document')}_processed.md"
            elif file_type == 'content_list':
                content = json.dumps(results.get('content_list', {}), indent=2, ensure_ascii=False)
                mimetype = 'application/json'
                filename = f"{results.get('filename', 'document')}_content_list.json"
            elif file_type == 'middle_json':
                content = json.dumps(results.get('middle_json', {}), indent=2, ensure_ascii=False)
                mimetype = 'application/json'
                filename = f"{results.get('filename', 'document')}_middle.json"
            else:
                return jsonify({'success': False, 'message': 'Invalid file type'}), 400

            return Response(
                content,
                mimetype=mimetype,
                headers={
                    'Content-Disposition': f'attachment; filename={filename}'
                }
            )

        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/<document_id>/segment-files', methods=['GET'])
    def get_document_segment_files(document_id):
        """获取文档的片段文件列表"""
        try:
            segment_files = db_manager.get_segment_files(document_id)
            return jsonify({
                'success': True,
                'document_id': document_id,
                'segment_files': segment_files
            })
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/<document_id>/segment-files/<file_type>/download', methods=['GET'])
    def download_segment_file(document_id, file_type):
        """下载特定类型的片段文件"""
        try:
            # 获取文件内容
            content = db_manager.get_segment_file_content(document_id, file_type)

            if content is None:
                return jsonify({'success': False, 'message': 'File not found'}), 404

            # 获取文件信息
            segment_files = db_manager.get_segment_files(document_id)
            target_file = None
            for file_info in segment_files:
                if file_info['file_type'] == file_type:
                    target_file = file_info
                    break

            if not target_file:
                return jsonify({'success': False, 'message': 'File metadata not found'}), 404

            from flask import Response
            import json

            # 根据文件类型生成内容和MIME类型
            filename = target_file["virtual_filename"]

            if file_type.endswith('_pdf'):
                # PDF格式文件
                if isinstance(content, bytes):
                    # 真正的PDF二进制内容
                    mimetype = 'application/pdf'
                else:
                    # 文本内容，作为PDF下载
                    if isinstance(content, str):
                        content = content.encode('utf-8')
                    mimetype = 'application/pdf'
            elif file_type == 'markdown':
                # Markdown文件
                if isinstance(content, bytes):
                    content = content.decode('utf-8')
                mimetype = 'text/markdown'
            elif file_type.endswith('_json') or file_type in ['content_list', 'middle_json', 'content_list_json', 'middle']:
                # JSON文件
                if isinstance(content, bytes):
                    content = content.decode('utf-8')
                elif isinstance(content, dict):
                    content = json.dumps(content, indent=2, ensure_ascii=False)
                mimetype = 'application/json'
            else:
                # 默认作为文本处理
                if isinstance(content, bytes):
                    content = content.decode('utf-8')
                elif isinstance(content, dict):
                    content = json.dumps(content, indent=2, ensure_ascii=False)
                    mimetype = 'application/json'
                else:
                    content = str(content)
                    mimetype = 'text/plain'

            return Response(
                content,
                mimetype=mimetype,
                headers={
                    'Content-Disposition': f'attachment; filename={filename}'
                }
            )

        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/use-magic-pdf', methods=['POST'])
    def use_magic_pdf_content():
        """更新文档使用Magic PDF生成的完整内容"""
        try:
            data = request.get_json()
            document_id = data.get('document_id')

            if not document_id:
                return jsonify({'success': False, 'message': 'Document ID is required'}), 400

            # 获取文档信息
            document = db_manager.get_document_by_id(document_id)
            if not document:
                return jsonify({'success': False, 'message': 'Document not found'}), 404

            filename = document.get('filename', '')

            # 尝试加载Magic PDF结果
            magic_pdf_results = load_magic_pdf_results(filename)

            if not magic_pdf_results or 'markdown' not in magic_pdf_results:
                return jsonify({'success': False, 'message': 'Magic PDF results not found for this document'}), 404

            # 更新数据库中的处理结果
            if db_manager.use_mongodb:
                from bson import ObjectId
                if isinstance(document_id, str):
                    document_id = ObjectId(document_id)

                # 更新processing_results集合
                update_result = db_manager.db.processing_results.update_one(
                    {"document_id": document_id},
                    {"$set": {
                        "markdown_content": magic_pdf_results['markdown'],
                        "content_list": magic_pdf_results.get('content_list', {}),
                        "middle_json": magic_pdf_results.get('middle_json', {}),
                        "updated_time": datetime.now(),
                        "source": "magic_pdf"
                    }},
                    upsert=True
                )

                # 更新segment_files集合中的markdown文件
                base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
                db_manager.db.segment_files.update_one(
                    {"document_id": document_id, "file_type": "markdown"},
                    {"$set": {
                        "content": magic_pdf_results['markdown'],
                        "updated_time": datetime.now(),
                        "source": "magic_pdf"
                    }}
                )

                return jsonify({
                    'success': True,
                    'message': f'Document updated with Magic PDF content ({len(magic_pdf_results["markdown"])} characters)',
                    'content_length': len(magic_pdf_results['markdown'])
                })
            else:
                return jsonify({'success': False, 'message': 'SQLite mode not supported for this operation'}), 400

        except Exception as e:
            print(f"Error updating document with Magic PDF content: {e}")
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/documents/batch-use-magic-pdf', methods=['POST'])
    def batch_use_magic_pdf_content():
        """批量更新所有可用的文档使用Magic PDF内容"""
        try:
            updated_count = 0
            failed_count = 0
            results = []

            # 获取所有文档
            documents = db_manager.get_all_documents()

            for document in documents:
                try:
                    filename = document.get('filename', '')
                    document_id = document.get('id')

                    # 尝试加载Magic PDF结果
                    magic_pdf_results = load_magic_pdf_results(filename)

                    if magic_pdf_results and 'markdown' in magic_pdf_results:
                        # 更新文档
                        if db_manager.use_mongodb:
                            from bson import ObjectId
                            if isinstance(document_id, str):
                                doc_id = ObjectId(document_id)
                            else:
                                doc_id = document_id

                            # 更新processing_results
                            db_manager.db.processing_results.update_one(
                                {"document_id": doc_id},
                                {"$set": {
                                    "markdown_content": magic_pdf_results['markdown'],
                                    "content_list": magic_pdf_results.get('content_list', {}),
                                    "middle_json": magic_pdf_results.get('middle_json', {}),
                                    "updated_time": datetime.now(),
                                    "source": "magic_pdf"
                                }},
                                upsert=True
                            )

                            # 更新segment_files
                            db_manager.db.segment_files.update_one(
                                {"document_id": doc_id, "file_type": "markdown"},
                                {"$set": {
                                    "content": magic_pdf_results['markdown'],
                                    "updated_time": datetime.now(),
                                    "source": "magic_pdf"
                                }}
                            )

                            updated_count += 1
                            results.append({
                                'filename': filename,
                                'status': 'updated',
                                'content_length': len(magic_pdf_results['markdown'])
                            })
                        else:
                            failed_count += 1
                            results.append({
                                'filename': filename,
                                'status': 'failed',
                                'reason': 'SQLite mode not supported'
                            })
                    else:
                        results.append({
                            'filename': filename,
                            'status': 'skipped',
                            'reason': 'Magic PDF results not found'
                        })

                except Exception as e:
                    failed_count += 1
                    results.append({
                        'filename': filename,
                        'status': 'failed',
                        'reason': str(e)
                    })

            return jsonify({
                'success': True,
                'updated_count': updated_count,
                'failed_count': failed_count,
                'total_documents': len(documents),
                'results': results
            })

        except Exception as e:
            print(f"Error in batch Magic PDF update: {e}")
            return jsonify({'success': False, 'message': str(e)}), 500

    if __name__ == '__main__':
        print("🔬" + "="*80 + "🔬")
        print("🚀 长鑫 Research Platform v2.0 Advanced - Starting...")
        print("🔬 Advanced Pharmaceutical AI Technology System")
        print("💻 Access URL: http://localhost:9588")
        print("🌟 New Advanced Features:")
        print("   📄 Batch Document Processing")
        print("   🎯 Professional Prompt Templates")
        print("   🤖 Automated Workflow System")
        print("   🔍 Multi-Document Analysis")
        print("   📑 PDF Segmentation & Analysis")
        print("   ⚡ AI-Powered Processing Pipeline")
        print("🔬" + "="*80 + "🔬")

        app.run(debug=True, host='0.0.0.0', port=9588,use_reloader=False)

except ImportError as e:
    print(f"Missing dependencies: {e}")
    print("Please install required packages:")
    print("pip install flask")
    print("Then run this script again.")
